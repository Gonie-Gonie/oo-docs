from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from oodocs import (
    Author,
    Chapter,
    Document,
    DocumentSettings,
    ImageData,
    Paragraph,
    TitleMatter,
)
from oodocs.components.cover import CoverPage
from oodocs.positioning import PageItemScope, TextBox
from oodocs.presets.templates import CoverPagePreset
from oodocs.styles.cover import CoverPageStyle


_TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
    b"\x00\x00\x03\x01\x01\x00\xc9\xfe\x92\xef\x00\x00\x00\x00IEND\xaeB`\x82"
)


class _SavefigLogo:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def savefig(self, target: object, **kwargs: object) -> None:
        self.calls.append(dict(kwargs))
        target.write(_TINY_PNG)


def _covered_document(*, logo: object | None = None) -> Document:
    return Document(
        "Caller-Owned Title",
        Chapter("Body", Paragraph("Body starts after the standalone cover.")),
        settings=DocumentSettings(
            title_matter=TitleMatter(
                subtitle="Caller-Owned Subtitle",
                authors=[Author("Example Author")],
                cover=CoverPage(
                    eyebrow="TECHNICAL NOTE",
                    organization="Example Lab",
                    logo=logo,
                    date="2026-07-14",
                    footer="Internal review",
                    note=Paragraph("Distribution note."),
                ),
            )
        ),
    )


def test_cover_model_coercion_and_defaults_are_neutral() -> None:
    empty = CoverPage()

    assert empty.eyebrow is None
    assert empty.organization is None
    assert empty.logo is None
    assert empty.date is None
    assert empty.footer is None
    assert empty.note == ()
    assert empty.extra_top == ()
    assert empty.extra_bottom == ()
    assert empty.style is None
    assert empty.resolved_style().name == "Cover page"
    assert empty.resolved_style().accent_color is None

    cover = CoverPage(
        eyebrow="DRAFT",
        organization="Example Lab",
        date="2026-07-14",
        footer="Review copy",
        note=["Legal notice.", Paragraph("Confidentiality notice.")],
        extra_top=["Top block."],
        extra_bottom=["Bottom block."],
        style="accented",
    )

    assert cover.eyebrow is not None
    assert cover.eyebrow[0].plain_text() == "DRAFT"
    assert cover.organization is not None
    assert cover.organization[0].plain_text() == "Example Lab"
    assert [block.plain_text() for block in cover.note] == [
        "Legal notice.",
        "Confidentiality notice.",
    ]
    assert cover.extra_top[0].plain_text() == "Top block."
    assert cover.extra_bottom[0].plain_text() == "Bottom block."
    assert cover.resolved_style().name == "Accented cover"


def test_generic_cover_presets_keep_content_caller_owned() -> None:
    accented = CoverPagePreset.accented()
    centered = CoverPagePreset.centered_logo(_TINY_PNG)

    assert accented.name == "Accented cover"
    assert accented.cover.resolved_style().name == "Accented cover"
    assert accented.cover.organization is None
    assert accented.cover.date is None
    assert accented.cover.footer is None
    assert centered.name == "Centered logo cover"
    assert centered.cover.resolved_style().name == "Centered logo cover"
    assert centered.cover.resolved_style().text_alignment == "center"
    assert centered.cover.footer is None
    assert not hasattr(CoverPagePreset, "eplus_simple")

    settings = CoverPagePreset.accented(
        organization="Example Lab",
        footer="Review copy",
    ).settings(
        subtitle="Caller subtitle",
        authors=[Author("Caller Author")],
    )
    assert settings.title_matter.cover is not None
    assert settings.title_matter.subtitle is not None
    assert settings.title_matter.subtitle[0].plain_text() == "Caller subtitle"
    assert settings.title_matter.authors[0].name == "Caller Author"


def test_cover_logo_reuses_figure_source_semantics(tmp_path: Path) -> None:
    logo_path = tmp_path / "logo.png"
    logo_path.write_bytes(_TINY_PNG)
    image_data = ImageData(_TINY_PNG)
    savefig_logo = _SavefigLogo()
    sources = (logo_path, _TINY_PNG, image_data, savefig_logo)

    figures = [
        CoverPage(
            logo=source,
            style=CoverPageStyle(logo_max_width=1.25, logo_max_height=0.75),
        ).logo_figure()
        for source in sources
    ]

    assert all(figure is not None for figure in figures)
    assert figures[0].image_source == logo_path
    assert isinstance(figures[1].image_source, ImageData)
    assert figures[1].image_source.data == _TINY_PNG
    assert figures[2].image_source is image_data
    assert figures[3].image_source is savefig_logo
    assert all(figure.width <= 1.25 for figure in figures)
    assert all(figure.height <= 0.75 for figure in figures)
    assert len({round(figure.width / figure.height, 6) for figure in figures}) == 1

    for index, source in enumerate(sources):
        html_path = tmp_path / f"logo-{index}.html"
        document = _covered_document(logo=source)
        assert document.validate().ok
        document.save_html(html_path)
        assert "data:image/png;base64," in html_path.read_text(encoding="utf-8")

    assert savefig_logo.calls
    assert savefig_logo.calls[0]["format"] == "png"


def test_html_cover_consumes_title_matter_once_and_has_print_break(tmp_path: Path) -> None:
    html_path = tmp_path / "cover.html"
    _covered_document().save_html(html_path)
    html = html_path.read_text(encoding="utf-8")

    match = re.search(
        r'<section class="([^"]*\boodocs-cover-page\b[^"]*)"[^>]*>(.*?)</section>',
        html,
        flags=re.DOTALL,
    )
    assert match is not None
    classes, cover_html = match.groups()
    assert "oodocs-page-break-after" in classes
    assert "Caller-Owned Title" in cover_html
    assert "Caller-Owned Subtitle" in cover_html
    assert "Example Author" in cover_html
    assert "Example Lab" in cover_html
    assert html.count('class="oodocs-title"') == 1
    assert ".oodocs-page-break-after" in html
    assert "break-after: page" in html
    assert "page-break-after: always" in html


def test_title_matter_without_cover_keeps_inline_header(tmp_path: Path) -> None:
    document = Document(
        "Inline Title",
        Paragraph("Body follows inline title matter."),
        settings=DocumentSettings(
            title_matter=TitleMatter(
                subtitle="Inline Subtitle",
                authors=[Author("Inline Author")],
            )
        ),
    )
    html_path = tmp_path / "inline-title.html"

    document.save_html(html_path)
    html = html_path.read_text(encoding="utf-8")

    assert '<header class="oodocs-title-matter">' in html
    assert re.search(r'<section class="[^"]*oodocs-cover-page', html) is None
    assert "Inline Title" in html
    assert "Inline Subtitle" in html
    assert "Inline Author" in html


def test_docx_and_pdf_put_body_after_cover(tmp_path: Path) -> None:
    document = _covered_document()
    docx_path = tmp_path / "cover.docx"
    pdf_path = tmp_path / "cover.pdf"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)

    word_document = WordDocument(docx_path)
    word_text = [paragraph.text for paragraph in word_document.paragraphs]
    assert len(word_document.sections) >= 2
    assert word_text.count("Caller-Owned Title") == 1
    assert word_text.count("Caller-Owned Subtitle") == 1
    assert word_text.count("Example Author") == 1

    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    page_text = [page.extract_text() or "" for page in pdf_reader.pages]
    assert len(page_text) >= 2
    assert "Caller-Owned Title" in page_text[0]
    assert "Caller-Owned Subtitle" in page_text[0]
    assert "Example Author" in page_text[0]
    assert "Body starts after the standalone cover." not in page_text[0]
    assert any("Body starts after the standalone cover." in text for text in page_text[1:])


def test_cover_validation_reports_asset_and_scope_codes(tmp_path: Path) -> None:
    missing_logo = _covered_document(logo=tmp_path / "missing-logo.png")
    missing_result = missing_logo.validate()

    missing_issues = [
        issue for issue in missing_result.errors if issue.code == "cover-asset-missing"
    ]
    assert len(missing_issues) == 1
    assert missing_issues[0].path.startswith("document.settings.title_matter.cover.logo")

    unsupported_logo = _covered_document(logo=object()).validate()
    assert "cover-asset-missing" in {
        issue.code for issue in unsupported_logo.errors
    }

    overlay_without_cover = Document(
        "Overlay Contract",
        Paragraph("Body."),
        settings=DocumentSettings(
            overlays=[
                TextBox(
                    "COVER ONLY",
                    x=0.2,
                    y=0.2,
                    width=1.5,
                    height=0.3,
                    scope=PageItemScope.cover(),
                )
            ]
        ),
    )
    scope_result = overlay_without_cover.validate()
    assert "cover-overlay-without-cover" in {
        issue.code for issue in scope_result.errors
    }


def test_cover_reference_documents_generic_contract() -> None:
    reference = Path("docs/reference/cover-page.md").read_text(encoding="utf-8")
    normalized = " ".join(reference.split())

    for phrase in (
        "visible document title still comes from `Document.title`",
        "subtitle and authors still come from `TitleMatter`",
        "`CoverPagePreset.accented(...)`",
        "`CoverPagePreset.centered_logo(...)`",
        "a compatible `savefig(...)` method",
        "`cover-asset-missing`",
        "`cover-overlay-without-cover`",
        "`oodocs-cover-page`",
        "`oodocs-page-break-after`",
    ):
        assert phrase in normalized
