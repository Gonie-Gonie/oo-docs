from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module():
    module_path = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "config_reference_example"
        / "main.py"
    )
    spec = importlib.util.spec_from_file_location("config_reference_example_main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_config_reference_example_builds_outputs(tmp_path: Path) -> None:
    example = _load_example_module()
    reference = example.load_config_reference()
    outputs = example.build(tmp_path)

    assert reference.title == "Documentation build config"
    assert len(reference.fields) == 6
    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])

    word_document = WordDocument(outputs["docx"])
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages)
    html_text = _normalized_html_text(outputs["html"])

    assert "Configuration Reference Example" in paragraph_texts
    assert "1 Configuration Overview" in paragraph_texts
    assert "2 Defaults and Examples" in paragraph_texts
    assert "build.output_dir" in table_text
    assert "OODOCS_OUTPUT_DIR" in table_text
    assert "environment.cache_dir" in table_text
    assert any("Configuration field reference." in text for text in paragraph_texts)
    assert any("Environment variable overrides documented by the schema." in text for text in paragraph_texts)
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=(
            "Configuration Reference Example",
            "1 Configuration Overview",
            "2 Defaults and Examples",
        ),
        min_tables=3,
    )
    assert "Configuration Reference Example" in pdf_text
    assert "Required fields" in pdf_text
    assert "Environment variables" in pdf_text
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=("Configuration Reference Example", "build.output_dir", "OODOCS_STRICT"),
        min_pages=1,
    )
    assert "Configuration Reference Example" in html_text
    assert "build.strict = true" in html_text
    assert_html_internal_links_resolve(outputs["html"])


def test_config_reference_example_supports_common_cli(tmp_path: Path, capsys) -> None:
    example = _load_example_module()
    output_dir = tmp_path / "cli"

    outputs = example.build(output_dir / "programmatic", output_formats=("html",))
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()

    example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "config-reference.html").exists()
