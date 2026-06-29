from __future__ import annotations

import importlib.util
import sys
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_template_module(name: str):
    example_dir = Path(__file__).resolve().parents[1] / "examples" / "template_presets"
    module_path = example_dir / f"{name}.py"
    spec = importlib.util.spec_from_file_location(f"examples.template_presets.{name}", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    sys.path.insert(0, str(example_dir))
    try:
        spec.loader.exec_module(module)
    finally:
        sys.path.remove(str(example_dir))
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_template_preset_examples_build_all_outputs(tmp_path: Path) -> None:
    build_all = _load_template_module("build_all")
    outputs = build_all.build_all(tmp_path)

    assert set(outputs) == {
        "journal_article_template",
    }

    expected_titles = {
        "journal-article-template": "Content-First Journal Article Template",
    }
    for stem, title in expected_titles.items():
        docx_path = tmp_path / f"{stem}.docx"
        pdf_path = tmp_path / f"{stem}.pdf"
        html_path = tmp_path / f"{stem}.html"

        assert_rendered_bundle(docx_path, pdf_path, html_path)

        word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
        word_document = WordDocument(docx_path)
        table_text = "\n".join(
            cell.text
            for table in word_document.tables
            for row in table.rows
            for cell in row.cells
        )
        pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
        normalized_pdf_text = " ".join(pdf_text.split())
        html_text = _normalized_html_text(html_path)

        assert title in word_text
        assert "Introduction" in word_text
        assert "Content-first journal templates separate manuscript content from repeated document assembly." in word_text
        assert "JournalArticleTemplate.build(...)" in word_text
        assert "Minimal Article" in word_text
        assert "Data Availability" in word_text
        assert "Acknowledgements" in word_text
        assert "Template preset catalog for article, report, manual, and book-shaped documents." in word_text
        assert "JournalArticleTemplate.build(...) input schema." in word_text
        assert "Template responsibility" in table_text
        assert "TechnicalReportTemplate" in table_text
        assert "SoftwareManualTemplate" in table_text
        assert "BookTemplate" in table_text
        assert "Template-first authoring compared with direct manuscript assembly." in word_text
        assert "journal_paper_example" in table_text
        assert "acknowledgements" in table_text
        assert len(word_document.tables) >= 4
        assert "References" in word_text
        assert_docx_structure(
            docx_path,
            required_paragraphs=(title, "1 Introduction", "References"),
            min_tables=1,
        )
        assert title in normalized_pdf_text
        assert "Methods" in normalized_pdf_text
        assert "Template preset catalog" in normalized_pdf_text
        assert "input schema" in normalized_pdf_text
        assert "Data Availability" in normalized_pdf_text
        assert "References" in normalized_pdf_text
        assert_pdf_text_and_pages(
            pdf_path,
            required_text=(title, "Methods", "References"),
            min_pages=2,
        )
        assert title in html_text
        assert "content-first template" in html_text
        assert "Template preset catalog" in html_text
        assert "journal_paper_example" in html_text
        assert "Minimal Article" in html_text
        assert_html_internal_links_resolve(
            html_path,
            required_text=(title, "content-first template"),
        )


def test_template_preset_minimal_article_omits_optional_sections(tmp_path: Path) -> None:
    journal_article_template = _load_template_module("journal_article_template")
    document = journal_article_template.build_minimal_document()

    assert document.validate().ok
    outputs = document.save_all(
        tmp_path,
        stem="minimal-article",
        formats=("html",),
        verbose=False,
    )
    html_text = _normalized_html_text(outputs["html"])

    assert "Minimal Article" in html_text
    assert "Body" in html_text
    assert "Data Availability" not in html_text
    assert "Acknowledgements" not in html_text


def test_template_preset_examples_support_common_cli(tmp_path: Path, capsys) -> None:
    build_all = _load_template_module("build_all")
    journal_article_template = _load_template_module("journal_article_template")
    output_dir = tmp_path / "cli"

    all_outputs = build_all.build_all(
        output_dir / "programmatic-all",
        output_formats=("html",),
    )
    assert set(all_outputs) == {"journal_article_template"}
    assert set(all_outputs["journal_article_template"].keys()) == {"html"}
    assert all_outputs["journal_article_template"]["html"].exists()

    template_outputs = journal_article_template.build(
        output_dir / "programmatic-one",
        output_formats=("html",),
    )
    assert set(template_outputs.keys()) == {"html"}
    assert template_outputs["html"].exists()

    build_all.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "journal-article-template.html").exists()
