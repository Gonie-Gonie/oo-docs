from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_release_notes_digest_example_builds_outputs(tmp_path: Path) -> None:
    release_notes_example = _load_example_module("release_notes_digest")
    files = release_notes_example.release_note_files()
    release_dates = release_notes_example.release_dates_from_git()
    latest_release = files[0]
    latest_version = latest_release.stem
    latest_release_date = release_notes_example.release_date_for_version(
        latest_version,
        release_dates,
    )
    latest_release_path = f"release-notes/{latest_release.name}"

    expected_count = len(list(Path(release_notes_example.RELEASE_NOTES_DIR).glob("*.md")))
    assert len(files) == expected_count
    assert files == sorted(
        files,
        key=release_notes_example.version_parts_from_filename,
        reverse=True,
    )
    assert latest_release.name == max(
        files,
        key=release_notes_example.version_parts_from_filename,
    ).name
    assert release_notes_example.release_type_from_version((1, 0, 0)) == "Major"
    assert release_notes_example.release_type_from_version((0, 10, 0)) == "Minor"
    assert release_notes_example.release_type_from_version((0, 9, 1)) == "Patch"
    assert release_notes_example.release_type_from_version((0, 9, 0)) == "Minor"
    assert all(
        release_notes_example.release_date_for_version(path.stem, release_dates)
        for path in files
    )
    assert latest_release_date == release_notes_example.PENDING_RELEASE_DATE or re.fullmatch(
        r"\d{4}-\d{2}-\d{2}",
        latest_release_date,
    )

    outputs = release_notes_example.build_release_notes(tmp_path)
    docx_path = outputs["docx"]
    pdf_path = outputs["pdf"]
    html_path = outputs["html"]

    assert_rendered_bundle(docx_path, pdf_path, html_path)

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    normalized_html_text = _normalized_html_text(html_path)
    html_text = html_path.read_text(encoding="utf-8")

    assert "OODocs Release Notes" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert any("Release Note Index" in text for text in paragraph_texts)
    assert any("Release-note digest workflow" in text for text in paragraph_texts)
    assert any("Version Management" in text for text in paragraph_texts)
    assert "3 Version History" in paragraph_texts
    assert any("Release runbook" in text for text in paragraph_texts)
    assert any("setuptools-scm" in text for text in paragraph_texts + [table_text])
    assert latest_release.name in table_text
    assert latest_release_date in table_text
    assert "latest" in table_text
    assert any(latest_release_path in text for text in paragraph_texts)
    assert f"Release date: {latest_release_date}." in paragraph_texts
    assert latest_version in paragraph_texts
    assert "Highlights" in paragraph_texts
    assert f"3 {latest_version}" not in paragraph_texts
    assert f"3.1 {latest_version}" not in paragraph_texts
    assert f"4 {latest_version}" not in paragraph_texts
    assert not any(
        re.fullmatch(r"\d+(?:\.\d+)* Highlights", text)
        for text in paragraph_texts
    )
    assert any(
        "Version-management rules demonstrated by this example." in text
        for text in paragraph_texts
    )
    assert any(
        "Release note files collected from the repository." in text
        for text in paragraph_texts
    )
    assert len(word_document.tables) == 2
    assert_docx_structure(
        docx_path,
        required_paragraphs=(
            "OODocs Release Notes",
            "Contents",
            "3 Version History",
            latest_version,
        ),
        table_count=2,
    )

    assert "OODocs Release Notes" in pdf_text
    assert "Contents" in pdf_text
    assert "Release Note Index" in pdf_text
    assert "Release-note digest workflow" in pdf_text
    assert "Version Management" in pdf_text
    assert "Version History" in pdf_text
    assert latest_release_path in pdf_text
    assert latest_release_date in pdf_text
    assert "setuptools-scm" in pdf_text
    assert len(pdf_reader.pages) >= 3
    assert_pdf_text_and_pages(
        pdf_path,
        required_text=(
            "OODocs Release Notes",
            "Version Management",
            "Version History",
        ),
        min_pages=3,
    )

    assert "OODocs Release Notes" in normalized_html_text
    assert "Contents" in normalized_html_text
    assert "Release Note Index" in normalized_html_text
    assert "Release-note digest workflow" in normalized_html_text
    assert "Version Management" in normalized_html_text
    assert "Version History" in normalized_html_text
    assert "Release runbook" in normalized_html_text
    assert latest_release_path in normalized_html_text
    assert latest_release_date in normalized_html_text
    assert "vMAJOR.MINOR.PATCH" in normalized_html_text
    assert (
        'class="oodocs-toc-entry oodocs-toc-entry-no-page '
        'oodocs-toc-entry-level-1"'
    ) in html_text
    toc_html = html_text.split("</nav>", 1)[0]
    assert ">3 Version History</a>" in toc_html
    assert f">{latest_version}</a>" in toc_html
    assert ">Highlights</a>" not in toc_html
    assert_html_internal_links_resolve(
        html_path,
        required_text=("OODocs Release Notes", "Version History"),
    )
