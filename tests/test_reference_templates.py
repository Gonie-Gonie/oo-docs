from __future__ import annotations

from html import unescape
from pathlib import Path
import re

import pytest
from docx import Document as WordDocument
from pypdf import PdfReader

from oodocs import Document, DocumentSettings, Paragraph, Section, Table, Theme, ref_range, refs
from oodocs.layout.indexing import build_render_index, resolve_block_reference
from oodocs.references import ReferenceFormat, bracket_ref, page_ref, paren_ref
from oodocs.styles import REFERENCE_KINDS, ReferenceDefaults, ReferenceTemplate


def _html_text(path: Path) -> str:
    markup = path.read_text(encoding="utf-8")
    return " ".join(unescape(re.sub(r"<[^>]+>", "", markup)).split())


def _docx_text(path: Path) -> str:
    document = WordDocument(path)
    table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return " ".join(
        [*(paragraph.text for paragraph in document.paragraphs), *table_text]
    )


def _pdf_text(path: Path) -> str:
    return " ".join(
        (page.extract_text() or "")
        for page in PdfReader(path).pages
    )


def test_reference_template_contract_and_resolved_target_kinds() -> None:
    assert REFERENCE_KINDS == (
        "part",
        "chapter",
        "section",
        "paragraph",
        "table",
        "figure",
        "equation",
        "code_block",
        "box",
        "countable",
    )

    suffix = ReferenceTemplate("절", plural_label="절", template="{value}{label}")
    label_free = ReferenceTemplate("", template="{value}{label}")
    custom_plural = ReferenceTemplate(
        "item",
        plural_label="items",
        template="{label} {value}",
        plural_template="{value} {label}",
    )
    assert suffix.format("2.1") == "2.1절"
    assert label_free.format("7") == "7"
    assert custom_plural.format("1 and 2", label="items", plural=True) == "1 and 2 items"

    chapter = Section("Chapter", level=1)
    section = Section("Section", level=2)
    chapter.add(section)
    document = Document("Kinds", chapter)
    render_index = build_render_index(document)

    resolved_chapter = resolve_block_reference(chapter, document.settings.theme, render_index)
    resolved_section = resolve_block_reference(section, document.settings.theme, render_index)
    assert resolved_chapter.target_kind == "chapter"
    assert resolved_section.target_kind == "section"

    with pytest.raises(ValueError, match="Unsupported reference target kind"):
        ReferenceDefaults({"project_record": ReferenceTemplate("Record")})
    with pytest.raises(ValueError, match="must contain '\\{value\\}'"):
        ReferenceTemplate("Broken", template="{label}")


@pytest.mark.parametrize(
    ("locale", "expected"),
    [
        ("en-US", "See Chapter 1 and Section 1.1."),
        ("ko-KR", "See 1장 and 1.1절."),
    ],
)
@pytest.mark.render
def test_locale_reference_placement_matches_in_docx_pdf_and_html(
    tmp_path: Path,
    locale: str,
    expected: str,
) -> None:
    chapter = Section("Introduction", level=1)
    scope = Section("Scope", level=2)
    chapter.add(scope, Paragraph("See ", chapter.ref(), " and ", scope.ref(), "."))
    document = Document(
        "Localized references",
        chapter,
        settings=DocumentSettings(theme=Theme.from_locale(locale)),
    )

    stem = locale.lower()
    docx_path = document.save_docx(tmp_path / f"{stem}.docx")
    pdf_path = document.save_pdf(tmp_path / f"{stem}.pdf")
    html_path = document.save_html(tmp_path / f"{stem}.html")

    assert expected in _docx_text(docx_path)
    assert expected in _pdf_text(pdf_path)
    assert expected in _html_text(html_path)


def test_custom_reference_templates_and_all_reference_helpers(tmp_path: Path) -> None:
    first = Table(["Value"], [["A"]], caption="First table.")
    second = Table(["Value"], [["B"]], caption="Second table.")
    table_references = ReferenceDefaults(
        {
            "table": ReferenceTemplate(
                "Artifact",
                plural_label="Artifacts",
                template="{label}@{value}",
                plural_template="{label}({value})",
            )
        }
    )
    document = Document(
        "Custom references",
        Paragraph(
            "single= ",
            first.ref(),
            "; plural= ",
            refs([first, second]),
            "; range= ",
            ref_range(first, second),
            "; paren= ",
            paren_ref(first),
            "; bracket= ",
            bracket_ref(second),
            "; page= ",
            page_ref(first),
            "; override= ",
            first.ref(
                reference_format=ReferenceFormat(
                    label="Node",
                    template="{value}-{label}",
                )
            ),
            ".",
        ),
        first,
        second,
        settings=DocumentSettings(theme=Theme(references=table_references)),
    )

    warning_codes = {issue.code for issue in document.validate().warnings}
    assert "page-aware-reference-degrades" in warning_codes

    html_path = document.save_html(tmp_path / "custom-references.html")
    text = _html_text(html_path)

    assert "single= Artifact@1" in text
    assert "plural= Artifacts(1 and 2)" in text
    assert "range= Artifacts(1-2)" in text
    assert "paren= (Artifact@1)" in text
    assert "bracket= [Artifact@2]" in text
    assert "page= Artifact@1" in text
    assert "override= 1-Node" in text
    assert "Table 1. First table." in text
    assert "Table 2. Second table." in text
