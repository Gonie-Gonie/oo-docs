from __future__ import annotations

from io import BytesIO

from docx import Document as WordDocument
from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)
from oodocs.apidoc import (
    ApiCollectConfig,
    ApiCoverageResult,
    ApiPackage,
    check_api_docs,
    collect_api,
)
from pypdf import PdfReader


LOCAL_ABSOLUTE_PATH_MARKERS = ("C:\\Users", "/home/", "/Users/")
STALE_API_REFERENCE_MARKERS = (
    "RenderedOutputs",
    "ParagraphTitleStyle",
    "from_ipynb",
    "parse_ipynb",
    "TextStyle.color",
    "TextBox.align",
    "ImageData.format",
    "build_table_layout",
    "coerce_blocks",
    "coerce_image_source",
    "coerce_inlines",
    "coerce_positioned_items",
    "image_source_to_buffer",
    "image_source_to_bytes",
    "normalize_media_placement",
    "normalize_table_split",
    "resolve_positioned_boxes",
)
STALE_RENDERED_API_REFERENCE_MARKERS = (
    *STALE_API_REFERENCE_MARKERS,
    "ImageData.savefig",
    "coerce_author_layout",
    "coerce_authors",
    "coerce_cell",
    "coerce_citation_library",
    "coerce_list_item",
    "normalize_citation_style",
    "normalize_color",
    "normalize_counter_format",
    "normalize_explicit_names",
    "normalize_generated_list_scope",
    "normalize_length_unit",
    "normalize_output_format",
    "normalize_output_formats",
    "normalize_parameter_columns",
    "normalize_reference_sort",
    "normalize_reference_style",
    "normalize_text_alignment",
    "normalize_vertical_alignment",
)
PRESET_REEXPORT_API_DOC_EXCLUDES = (
    "oodocs.presets.BookTemplate",
    "oodocs.presets.CalloutBox",
    "oodocs.presets.CompactTable",
    "oodocs.presets.CoverPagePreset",
    "oodocs.presets.info_box",
    "oodocs.presets.JournalArticleTemplate",
    "oodocs.presets.KeyValueTable",
    "oodocs.presets.ManuscriptSection",
    "oodocs.presets.Nomenclature",
    "oodocs.presets.note_box",
    "oodocs.presets.option_table",
    "oodocs.presets.SoftwareManualTemplate",
    "oodocs.presets.success_box",
    "oodocs.presets.TechnicalReportTemplate",
    "oodocs.presets.warning_box",
)


def _assert_no_local_absolute_paths(text: str) -> None:
    assert not any(marker in text for marker in LOCAL_ABSOLUTE_PATH_MARKERS)


def _assert_no_stale_api_reference_markers(text: str) -> None:
    assert not any(marker in text for marker in STALE_API_REFERENCE_MARKERS)


def _assert_no_stale_rendered_api_reference_markers(text: str) -> None:
    assert not any(marker in text for marker in STALE_RENDERED_API_REFERENCE_MARKERS)


def _assert_clean_api_reference_text(text: str) -> None:
    _assert_no_local_absolute_paths(text)
    _assert_no_stale_api_reference_markers(text)


def _assert_clean_rendered_api_reference_text(text: str) -> None:
    _assert_no_local_absolute_paths(text)
    _assert_no_stale_rendered_api_reference_markers(text)


def _docx_text(path) -> str:
    document = WordDocument(path)
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                pieces.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(pieces)


def _pdf_text(path) -> str:
    reader = PdfReader(BytesIO(path.read_bytes()))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def test_apidoc_collects_oodocs_public_api_for_self_reference() -> None:
    api = collect_api("oodocs", collector="auto", public_policy="__all__")

    assert api.name == "oodocs"
    assert api.find_object("oodocs.Document") is not None
    assert api.select_public_objects()


def test_apidoc_pyproject_config_hides_compatibility_adapters() -> None:
    config = ApiCollectConfig.from_pyproject(".")
    api = collect_api(".", config=config)

    assert "oodocs.public_api" in config.module_exclude_patterns
    assert "*.ImageData.savefig" in config.object_exclude_patterns
    assert set(PRESET_REEXPORT_API_DOC_EXCLUDES) <= set(config.object_exclude_patterns)
    assert api.find_module("oodocs.public_api") is None
    assert api.find_object("oodocs.ImageData") is not None
    assert api.find_object("oodocs.ImageData.savefig") is None
    assert api.find_object("oodocs.components.media.ImageData.savefig") is None
    assert api.find_object("oodocs.presets.components.CalloutBox") is not None
    assert api.find_object("oodocs.presets.templates.JournalArticleTemplate") is not None
    for qualname in PRESET_REEXPORT_API_DOC_EXCLUDES:
        assert api.find_object(qualname) is None, qualname


def test_apidoc_renders_oodocs_public_api_reference_bundle(tmp_path) -> None:
    config = ApiCollectConfig.from_pyproject(".")
    api = collect_api(".", config=config)
    coverage = check_api_docs(api)
    document = api.to_help_book(
        presentation="compact",
        include_coverage=True,
        max_heading_level=3,
    )

    outputs = document.save_all(
        tmp_path,
        stem="oodocs-api",
        formats=("docx", "pdf", "html"),
    )
    api_path = api.save_json(tmp_path / "oodocs-api.json")
    coverage_json_path = coverage.save_json(tmp_path / "oodocs-api-coverage.json")
    coverage_csv_path = coverage.save_csv(tmp_path / "oodocs-api-coverage.csv")

    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])
    assert api_path.exists()
    assert coverage_json_path.exists()
    assert coverage_csv_path.exists()
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=("oodocs API Reference",),
        min_tables=3,
    )
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=(
            "oodocs API Reference",
            "API Documentation Coverage",
            "oodocs.Document",
        ),
        min_pages=1,
    )
    assert_html_internal_links_resolve(
        outputs["html"],
        required_text=(
            "oodocs API Reference",
            "API Documentation Coverage",
            "oodocs.Document",
        ),
    )
    html = outputs["html"].read_text(encoding="utf-8")
    assert "Uncategorized API" not in html
    assert "Renderer Extension API" not in html
    _assert_clean_rendered_api_reference_text(_docx_text(outputs["docx"]))
    _assert_clean_rendered_api_reference_text(_pdf_text(outputs["pdf"]))
    _assert_clean_rendered_api_reference_text(html)
    _assert_clean_api_reference_text(api_path.read_text(encoding="utf-8"))
    _assert_clean_api_reference_text(coverage_json_path.read_text(encoding="utf-8"))
    _assert_clean_api_reference_text(coverage_csv_path.read_text(encoding="utf-8"))
    assert not ("<th>Label</th>" in html and "See Also" in html)
    assert html.find("API Contents") < html.find("API Documentation Coverage")
    assert ApiPackage.load_json(api_path).find_object("oodocs.Document") is not None
    assert ApiCoverageResult.load_json(coverage_json_path).public_object_count > 0
