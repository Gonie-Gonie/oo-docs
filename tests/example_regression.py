from __future__ import annotations

from html import unescape
from io import BytesIO
from pathlib import Path
import re
from typing import Iterable

from docx import Document as WordDocument
from pypdf import PdfReader


def assert_rendered_bundle(docx_path: Path, pdf_path: Path, html_path: Path) -> None:
    assert docx_path.exists()
    assert pdf_path.exists()
    assert html_path.exists()
    assert docx_path.stat().st_size > 0
    assert pdf_path.stat().st_size > 0
    assert html_path.stat().st_size > 0


def assert_docx_structure(
    docx_path: Path,
    *,
    required_paragraphs: Iterable[str] = (),
    table_count: int | None = None,
    min_tables: int | None = None,
    inline_shape_count: int | None = None,
    min_inline_shapes: int | None = None,
    comment_count: int | None = None,
) -> None:
    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]

    for text in required_paragraphs:
        assert text in paragraph_texts
    if table_count is not None:
        assert len(word_document.tables) == table_count
    if min_tables is not None:
        assert len(word_document.tables) >= min_tables
    if inline_shape_count is not None:
        assert len(word_document.inline_shapes) == inline_shape_count
    if min_inline_shapes is not None:
        assert len(word_document.inline_shapes) >= min_inline_shapes
    if comment_count is not None:
        assert len(word_document.comments) == comment_count


def assert_pdf_text_and_pages(
    pdf_path: Path,
    *,
    required_text: Iterable[str] = (),
    min_pages: int | None = None,
    max_pages: int | None = None,
) -> None:
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)

    for text in required_text:
        assert text in pdf_text
    if min_pages is not None:
        assert len(pdf_reader.pages) >= min_pages
    if max_pages is not None:
        assert len(pdf_reader.pages) <= max_pages


def assert_html_internal_links_resolve(
    html_path: Path,
    *,
    required_hrefs: Iterable[str] = (),
    required_text: Iterable[str] = (),
) -> None:
    html_text = html_path.read_text(encoding="utf-8")
    normalized_text = _normalized_html_text(html_text)
    ids = set(re.findall(r'\sid="([^"]+)"', html_text))
    internal_hrefs = re.findall(r'href="#([^"]+)"', html_text)

    for href in required_hrefs:
        normalized_href = href.removeprefix("#")
        assert normalized_href in internal_hrefs
        assert normalized_href in ids
    for href in internal_hrefs:
        assert href in ids
    for text in required_text:
        assert text in normalized_text


def _normalized_html_text(html_text: str) -> str:
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())
