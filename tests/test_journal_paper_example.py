from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import zipfile

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def _docx_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_journal_paper_example_builds_outputs(tmp_path: Path) -> None:
    paper_example = _load_example_module("journal_paper_example")
    inputs = paper_example.load_inputs()
    assert len(inputs.results) == 4
    assert len(inputs.ablation) == 4
    assert inputs.traceability_diagram.exists()

    outputs = paper_example.build_journal_paper(tmp_path)
    docx_path = outputs["docx"]
    pdf_path = outputs["pdf"]
    html_path = outputs["html"]

    assert_rendered_bundle(docx_path, pdf_path, html_path)
    assert (Path(paper_example.__file__).resolve().parent / "assets" / "benchmark_results.csv").exists()
    assert (Path(paper_example.__file__).resolve().parent / "assets" / "ablation_results.csv").exists()

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    normalized_html_text = _normalized_html_text(html_path)
    html_text = html_path.read_text(encoding="utf-8")

    assert "OODocs Development Philosophy" in paragraph_texts
    assert "Hyeong-Gon Jo [1]*, Codex [2]" in paragraph_texts
    assert "[1] Building Simulation LAB, Seoul National University, Seoul, Republic of Korea" in paragraph_texts
    assert "[2] OpenAI" in paragraph_texts
    assert "* Corresponding author: Hyeong-Gon Jo" in paragraph_texts
    assert any("ORCID 0009-0004-8821-275X" in text for text in paragraph_texts)
    assert "Abstract" in paragraph_texts
    assert "Highlights" in paragraph_texts
    assert "1 Introduction" in paragraph_texts
    assert "2 Workflow Design" in paragraph_texts
    assert "2.1 Evidence Traceability" in paragraph_texts
    assert "3 Study Assets" in paragraph_texts
    assert "4 Results" in paragraph_texts
    assert "4.1 Benchmark Frontier" in paragraph_texts
    assert "4.2 Ablation Signals" in paragraph_texts
    assert "4.3 Late-Revision Cost" in paragraph_texts
    assert "5 Discussion" in paragraph_texts
    assert "6 Conclusion" in paragraph_texts
    assert "Acknowledgements" in paragraph_texts
    assert "References" in paragraph_texts
    assert all(text not in paragraph_texts for text in ("Contents", "List of Tables", "List of Figures"))
    assert any("pandas.read_csv" in text for text in paragraph_texts)
    assert any("matplotlib" in text for text in paragraph_texts)
    assert any("Traceability pipeline used in the study. Source: examples/journal_paper_example/assets/traceability-diagram.png." in text for text in paragraph_texts)
    assert any("Quality-latency frontier generated from examples/journal_paper_example/assets/benchmark_results.csv." in text for text in paragraph_texts)
    assert any("Estimated late-revision synchronization effort generated from the in-script model." in text for text in paragraph_texts)
    assert "Full control over manuscript structure" in table_text
    assert "Fast article skeleton" in table_text
    assert len(word_document.tables) == 4
    assert len(word_document.inline_shapes) == 3
    assert 'w:num="2"' in _docx_document_xml(docx_path)
    assert_docx_structure(
        docx_path,
        required_paragraphs=(
            "OODocs Development Philosophy",
            "Abstract",
            "Highlights",
            "References",
        ),
        table_count=4,
        inline_shape_count=3,
    )

    assert "OODocs Development Philosophy" in pdf_text
    assert "Hyeong-Gon Jo [1]*, Codex [2]" in pdf_text
    assert "Abstract" in pdf_text
    assert "Highlights" in pdf_text
    assert "Introduction" in pdf_text
    assert "Workflow Design" in pdf_text
    assert "Evidence Traceability" in pdf_text
    assert "Study Assets" in pdf_text
    assert "Results" in pdf_text
    assert "Discussion" in pdf_text
    assert "Conclusion" in pdf_text
    assert "References" in pdf_text
    assert "Contents" not in pdf_text
    assert "List of Tables" not in pdf_text
    assert "List of Figures" not in pdf_text
    assert "pandas.read_csv" in pdf_text
    assert "matplotlib" in pdf_text
    assert "Traceability pipeline used in the study" in pdf_text
    assert "Benchmark results loaded from" in pdf_text
    assert "Ablation results loaded from" in pdf_text
    assert "Direct Assembly" in pdf_text
    assert "Template Preset" in pdf_text
    assert "Full control over manuscript structure" in pdf_text
    assert "late-revision" in pdf_text
    assert "synchronization effort" in pdf_text
    assert "in-script model" in pdf_text
    assert "https://doi.org/10.1093/comjnl/27.2.97" in pdf_text
    assert "https://doi.org/10.1198/106186007X178663" in pdf_text
    assert "https://yihui.org/knitr/" in pdf_text
    assert 6 <= len(pdf_reader.pages) <= 8
    assert _pdf_image_draw_count(pdf_path) == 3
    assert_pdf_text_and_pages(
        pdf_path,
        required_text=(
            "OODocs Development Philosophy",
            "Abstract",
            "Results",
            "References",
        ),
        min_pages=6,
        max_pages=8,
    )

    assert "OODocs Development Philosophy" in normalized_html_text
    assert "Hyeong-Gon Jo [1]*, Codex [2]" in normalized_html_text
    assert "Evidence Traceability" in normalized_html_text
    assert "Direct Assembly or Template Preset" in normalized_html_text
    assert "Study Assets" in normalized_html_text
    assert "Benchmark Frontier" in normalized_html_text
    assert "Ablation Signals" in normalized_html_text
    assert "Late-Revision Cost" in normalized_html_text
    assert "The workflow studied here treats the manuscript itself as code." in normalized_html_text
    assert "Full control over manuscript structure" in normalized_html_text
    assert "examples/journal_paper_example/assets/benchmark_results.csv" in normalized_html_text
    assert "https://doi.org/10.1093/comjnl/27.2.97" in normalized_html_text
    assert "https://doi.org/10.1198/106186007X178663" in normalized_html_text
    assert "https://yihui.org/knitr/" in normalized_html_text
    assert html_text.count("data:image/png;base64,") == 3
    assert 'class="oodocs-multi-column-layout"' in html_text
    assert "column-count: 2" in html_text
    assert 'href="#table_2"' in html_text
    assert 'href="#figure_2"' in html_text
    assert_html_internal_links_resolve(
        html_path,
        required_hrefs=("#table_2", "#figure_2"),
        required_text=("OODocs Development Philosophy", "Benchmark Frontier"),
    )


def test_journal_paper_example_supports_common_cli(tmp_path: Path, capsys) -> None:
    paper_example = _load_example_module("journal_paper_example")
    output_dir = tmp_path / "cli"

    outputs = paper_example.build_journal_paper(
        output_dir / "programmatic",
        output_formats=("html",),
    )
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()

    paper_example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "oodocs-development-philosophy.html").exists()
