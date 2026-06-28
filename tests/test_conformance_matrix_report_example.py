from __future__ import annotations

import importlib.util
import json
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module():
    module_path = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "conformance_matrix_report"
        / "main.py"
    )
    spec = importlib.util.spec_from_file_location("conformance_matrix_report_main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_conformance_matrix_report_example_builds_outputs(tmp_path: Path) -> None:
    example = _load_example_module()
    records = example.load_results()
    outputs = example.build(tmp_path)

    assert len(records) == 6
    assert example.status_counts(records)["fail"] == 1
    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])
    assert outputs.full_matrix_json.exists()

    sidecar = json.loads(outputs.full_matrix_json.read_text(encoding="utf-8"))
    assert sidecar["record_count"] == 6
    assert "temperature_error" in sidecar["columns"]
    assert any(record["status"] == "fail" for record in sidecar["records"])

    word_document = WordDocument(outputs["docx"])
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages)
    html_text = _normalized_html_text(outputs["html"])

    assert "Conformance Matrix Report" in paragraph_texts
    assert "1 Claim Boundary" in paragraph_texts
    assert "2 Summary" in paragraph_texts
    assert "3 Failure Detail Appendix" in paragraph_texts
    assert any("PDF excerpt matrix; full matrix is written to the JSON sidecar." in text for text in paragraph_texts)
    assert "case-004" in table_text
    assert "pressure residual exceeds policy" in table_text
    assert "temperature_error" not in table_text
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=(
            "Conformance Matrix Report",
            "1 Claim Boundary",
            "2 Summary",
        ),
        min_tables=4,
    )
    assert "Conformance Matrix Report" in pdf_text
    assert "conformance-matrix-full.json" in pdf_text
    assert "Failure Detail Appendix" in pdf_text
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=("Conformance Matrix Report", "case-004", "full matrix"),
        min_pages=1,
    )
    assert "Conformance Matrix Report" in html_text
    assert "Wide matrix reporting policy." in html_text
    assert_html_internal_links_resolve(outputs["html"])


def test_conformance_matrix_report_example_supports_common_cli(
    tmp_path: Path,
    capsys,
) -> None:
    example = _load_example_module()
    output_dir = tmp_path / "cli"

    outputs = example.build(output_dir / "programmatic", output_formats=("html",))
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()
    assert outputs.full_matrix_json.exists()

    example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "conformance-matrix-report.html").exists()
    assert (output_dir / "conformance-matrix-full.json").exists()
