from __future__ import annotations

import pytest

pytestmark = pytest.mark.examples

from html import unescape
import importlib.util
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module():
    module_path = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "engineering_report_example"
        / "main.py"
    )
    spec = importlib.util.spec_from_file_location("engineering_report_example_main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_engineering_report_example_builds_outputs(tmp_path: Path) -> None:
    example = _load_example_module()
    document = example.build_document()
    outputs = example.build(tmp_path)

    assert document.validate().ok
    assert example.requirement_rows()[0][0] == "R-01"
    assert example.verification_rows()[0][1] == "pass"
    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])

    word_document = WordDocument(outputs["docx"])
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages)
    html_text = _normalized_html_text(outputs["html"])

    assert "Engineering Report Example" in paragraph_texts
    assert "1 Signal Processing Method" in paragraph_texts
    assert any("Quality Gate Algorithm" in text for text in paragraph_texts)
    assert "Algorithm 1" in "\n".join(paragraph_texts)
    assert any("Input:" in text for text in paragraph_texts)
    assert any("Output:" in text for text in paragraph_texts)
    assert "R-01" in table_text
    assert "Drift correction" in table_text
    assert "Signal quality gate" in pdf_text
    assert "Engineering Report Example" in html_text
    assert "Signal quality gate" in html_text
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=(
            "Engineering Report Example",
            "1 Signal Processing Method",
        ),
        min_tables=2,
    )
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=("Engineering Report Example", "Signal quality gate"),
        min_pages=1,
    )
    assert_html_internal_links_resolve(outputs["html"])


def test_engineering_report_example_supports_common_cli(tmp_path: Path, capsys) -> None:
    example = _load_example_module()
    output_dir = tmp_path / "cli"

    outputs = example.build(output_dir / "programmatic", output_formats=("html",))
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()

    example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "engineering-method-report.html").exists()
