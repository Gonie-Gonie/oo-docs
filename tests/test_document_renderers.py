from __future__ import annotations

from base64 import b64decode
from importlib.metadata import version as package_version
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import struct
import zlib
import zipfile

import oodocs
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from pypdf import PdfReader
import pytest

import oodocs.components.generated as generated_components
import oodocs.components.inline as inline_components
import oodocs.review as review_components
from oodocs.components.equations import BASELINE, SUBSCRIPT, SUPERSCRIPT, parse_latex_segments
from oodocs.components.media import build_table_layout
from oodocs.core import length_to_inches
from oodocs.layout.indexing import build_render_index
from oodocs.renderers.pdf import PdfRenderer
from oodocs.review import MarginNote, Todo, margin_note, todo
from oodocs import (
    Affiliation,
    Acronym,
    Algorithm,
    AlignedEquation,
    Assumption,
    Author,
    AuthorLayout,
    Appendix,
    Axiom,
    BlockDefaults,
    BorderStyle,
    Box,
    BoxStyle,
    BulletList,
    CasesEquation,
    CaptionDefaults,
    ChemicalFormula,
    CitationDefaults,
    CitationLibrary,
    CitationSource,
    Chapter,
    Claim,
    Comment,
    CommentList,
    CodeBlock,
    ColumnSpan,
    ColumnSpec,
    Conjecture,
    CountableBlock,
    Corollary,
    CounterStyle,
    CropBox,
    Definition,
    Document,
    DocumentMetadata,
    DocumentSettings,
    DocumentValidationError,
    Divider,
    Equation,
    Example,
    Figure,
    ListOfFigures,
    Footnote,
    FootnoteDefaults,
    FootnoteStyle,
    GeneratedContentDefaults,
    Glossary,
    GlossaryList,
    GlossaryTerm,
    HeaderFooterDefaults,
    HeadingStyle,
    HeadingNumbering,
    ImageBox,
    ImageData,
    InlineChip,
    InlineChipStyle,
    Lemma,
    LinkDefaults,
    ListOfAlgorithms,
    ListStyle,
    Math,
    MultiColumn,
    NumberedList,
    OODocsError,
    OutputBundle,
    PageNumberDefaults,
    PageLayout,
    PageMargins,
    PageSize,
    PageBreak,
    PageItemScope,
    Padding,
    PdfPages,
    Paragraph,
    ParagraphStyle,
    RunInTitleStyle,
    Part,
    Proof,
    Proposition,
    ReactionEquation,
    ReferenceList,
    ReferenceFormat,
    Remark,
    Section,
    Shape,
    StrokeStyle,
    StyleSheet,
    SubFigure,
    SubFigureGroup,
    SubTable,
    SubTableGroup,
    Subsection,
    SubSubsection,
    Table,
    TableCell,
    TableCellStyle,
    TableStyle,
    TableOfContents,
    ListOfTables,
    Text,
    TextStyle,
    TextBox,
    Theorem,
    Theme,
    TitleMatterDefaults,
    TocLevelStyle,
    TypographyDefaults,
    ValidationResult,
    VerticalSpace,
    badge,
    bold,
    chemical_formula,
    inline_code,
    text_color,
    cite,
    comment,
    create_countable_block_type,
    footnote,
    highlight,
    italic,
    keyboard,
    link,
    line_break,
    math,
    markup,
    prescript,
    Ref,
    ref_range,
    refs,
    page_ref,
    paren_ref,
    reference,
    status,
    strikethrough,
    styled,
    subscript,
    superscript,
    tag,
    url,
)
from oodocs.presets.components import (
    CalloutBox,
    info_box,
    KeyValueTable,
    Nomenclature,
    note_box,
    success_box,
    warning_box,
)
from oodocs.presets.templates import (
    BookTemplate,
    CoverPagePreset,
    JournalArticleTemplate,
    ManuscriptSection,
    SoftwareManualTemplate,
    TechnicalReportTemplate,
)
from oodocs.styles import TextStyle

class HighlightedParagraph(Paragraph):
    pass


def _write_sample_image(path: Path) -> None:
    path.write_bytes(_build_sample_png())


def _build_sample_png(width: int = 360, height: int = 220) -> bytes:
    rows: list[bytes] = []
    for y in range(height):
        row = bytearray([0])
        for x in range(width):
            if y < 34:
                pixel = (34, 58, 94)
            elif x < 18 or x >= width - 18 or y < 52 or y >= height - 18:
                pixel = (214, 221, 233)
            elif 26 < x < width - 26 and 70 < y < 102:
                pixel = (205, 121, 62)
            elif (x - 36) // 54 % 2 == 0 and 122 < y < 182:
                pixel = (89, 132, 198)
            else:
                pixel = (247, 249, 252)
            row.extend(pixel)
        rows.append(bytes(row))

    raw_image = b"".join(rows)
    return b"".join(
        (
            b"\x89PNG\r\n\x1a\n",
            _png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)),
            _png_chunk(b"IDAT", zlib.compress(raw_image, level=9)),
            _png_chunk(b"IEND", b""),
        )
    )


def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
    payload = chunk_type + data
    checksum = zlib.crc32(payload) & 0xFFFFFFFF
    return struct.pack(">I", len(data)) + payload + struct.pack(">I", checksum)


class FakeAxis:
    def __init__(self, values: list[object], names: tuple[str | None, ...] = ("",)) -> None:
        self._values = values
        self.names = names
        self.name = names[0] if names else None
        self.nlevels = max((len(value) if isinstance(value, tuple) else 1) for value in values) if values else 1

    def tolist(self) -> list[object]:
        return list(self._values)

    def __iter__(self):
        return iter(self._values)


class FakeDataFrame:
    def __init__(
        self,
        *,
        columns: list[object],
        rows: list[list[object]],
        index: list[object] | None = None,
        index_names: tuple[str | None, ...] = ("",),
    ) -> None:
        self.columns = FakeAxis(columns)
        self._rows = rows
        self.index = FakeAxis(index or list(range(len(rows))), names=index_names)

    def itertuples(self, *, index: bool = False, name: str | None = None):
        for row_index, row in enumerate(self._rows):
            if index:
                yield (self.index.tolist()[row_index], *row)
            else:
                yield tuple(row)


class FakeFigure:
    def __init__(self, image_bytes: bytes) -> None:
        self.image_bytes = image_bytes
        self.calls: list[dict[str, object]] = []

    def savefig(self, target: object, **kwargs: object) -> None:
        self.calls.append(dict(kwargs))
        target.write(self.image_bytes)


def _pdf_font_names(pdf_path: Path) -> set[str]:
    font_names: set[str] = set()
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/Font" not in resources:
            continue
        fonts = resources["/Font"].get_object()
        for font in fonts.values():
            font_object = font.get_object()
            base_font = font_object.get("/BaseFont")
            if base_font is not None:
                font_names.add(str(base_font))
    return font_names


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def _pdf_content_bytes(pdf_path: Path) -> bytes:
    parts: list[bytes] = []
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        contents = page.get_contents()
        if contents is None:
            continue
        if isinstance(contents, list):
            for item in contents:
                parts.append(item.get_data())
        else:
            parts.append(contents.get_data())
    return b"\n".join(parts)


def _pdf_text_context(pdf_path: Path, text: str, window: int = 160) -> bytes:
    content = _pdf_content_bytes(pdf_path)
    needle = f"({text})".encode()
    index = content.find(needle)
    assert index != -1, f"{text!r} not found in PDF content stream"
    start = max(index - window, 0)
    return content[start : index + len(needle) + window]


def _docx_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _docx_word_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


def _docx_relationships_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/_rels/") and name.endswith(".rels")
        )


def _docx_settings_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read("word/settings.xml").decode("utf-8")


def _pdf_uri_targets(pdf_path: Path) -> list[str]:
    targets: list[str] = []
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        for annotation_ref in page.get("/Annots", []):
            annotation = annotation_ref.get_object()
            action = annotation.get("/A")
            if action is None:
                continue
            uri = action.get("/URI")
            if uri is not None:
                targets.append(str(uri))
    return targets


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_version_is_defined() -> None:
    from oodocs import __version__

    assert __version__ == package_version("oodocs")
    assert __version__


def test_markup_creates_styled_fragments() -> None:
    fragments = markup("plain **bold** *italic* `code`")

    assert [fragment.value for fragment in fragments] == ["plain ", "bold", " ", "italic", " ", "code"]
    assert fragments[1].style.bold is True
    assert fragments[3].style.italic is True
    assert fragments[5].style.font_name == "Courier New"


def test_list_classes_create_block_instances() -> None:
    bullet = BulletList("first", Paragraph("second"))
    ordered = NumberedList("step one", "step two")

    assert isinstance(bullet, BulletList)
    assert bullet.ordered is False
    assert [item.plain_text() for item in bullet.items] == ["first", "second"]
    assert isinstance(ordered, NumberedList)
    assert ordered.ordered is True
    assert [item.plain_text() for item in ordered.items] == ["step one", "step two"]


def test_comment_and_math_helpers_create_renderable_fragments() -> None:
    inline_comment = comment("term", "Expanded note", author="pytest", initials="PT")
    inline_todo = todo("Verify units.", owner="QA", status="review")
    inline_margin_note = margin_note("Keep near the source paragraph.", side="left")
    inline_footnote = footnote("term", "Portable footnote note")
    inline_math = math(r"\alpha^2 + \beta^2")
    equation = Equation(r"\frac{1}{2}")

    assert isinstance(inline_comment, oodocs.Comment)
    assert inline_comment.plain_text() == "term[?]"
    assert inline_comment.author == "pytest"
    assert inline_comment.initials == "PT"
    assert isinstance(inline_todo, Todo)
    assert inline_todo.plain_text() == "TODO[?]"
    assert inline_todo.owner == "QA"
    assert inline_todo.status == "review"
    assert isinstance(inline_margin_note, MarginNote)
    assert inline_margin_note.side == "left"
    assert inline_margin_note.plain_text() == "[?]"
    assert isinstance(inline_footnote, Footnote)
    assert inline_footnote.plain_text() == "term[?]"
    assert isinstance(inline_math, Math)
    assert inline_math.plain_text() == "alpha2 + beta2"
    assert equation.plain_text() == "(1)/(2)"


def test_math_prescript_renders_to_all_outputs(tmp_path: Path) -> None:
    segments = parse_latex_segments(r"\prescript{14}{6}{C} + {}^{3}_{1}H")
    assert [(segment.text, segment.vertical_align) for segment in segments] == [
        ("14", SUPERSCRIPT),
        ("6", SUBSCRIPT),
        ("C + ", BASELINE),
        ("3", SUPERSCRIPT),
        ("1", SUBSCRIPT),
        ("H", BASELINE),
    ]

    document = Document(
        "Prescript Math",
        Paragraph("Inline isotope ", Math(r"\prescript{14}{6}{C} + {}^{3}_{1}H"), "."),
        Equation(r"\prescript{14}{6}{C} + {}^{3}_{1}H"),
    )

    docx_path = tmp_path / "prescript.docx"
    pdf_path = tmp_path / "prescript.pdf"
    html_path = tmp_path / "prescript.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    inline_paragraph = next(paragraph for paragraph in word_document.paragraphs if "Inline isotope" in paragraph.text)
    equation_paragraph = next(paragraph for paragraph in word_document.paragraphs if paragraph.text.startswith("146C + 31H"))
    assert any(run.text == "14" and run.font.superscript for run in inline_paragraph.runs)
    assert any(run.text == "6" and run.font.subscript for run in inline_paragraph.runs)
    assert any(run.text == "3" and run.font.superscript for run in equation_paragraph.runs)
    assert any(run.text == "1" and run.font.subscript for run in equation_paragraph.runs)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "146C + 31H" in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    assert "<sup>14</sup><sub>6</sub>C" in html_text
    assert "<sup>3</sup><sub>1</sub>H" in html_text


def test_amsmath_equation_blocks_number_reference_and_validate(tmp_path: Path) -> None:
    aligned = AlignedEquation(
        r"a &= b + c",
        r"  &= d",
        reference_label="Eq.",
    )
    cases = CasesEquation(
        ("0", "x < 0"),
        (r"x^2", r"x \geq 0"),
        left="f(x)",
    )
    unnumbered = Equation(r"\operatorname{loss}(x)", numbered=False)
    unsupported = Equation(r"\foo{x} + 1", numbered=False)
    document = Document(
        "AMSMath Blocks",
        Paragraph("See ", aligned.reference(), " and ", unnumbered.reference("the loss definition"), "."),
        aligned,
        cases,
        unnumbered,
        unsupported,
    )

    render_index = build_render_index(document)
    assert render_index.equation_number(aligned) == 1
    assert render_index.equation_number(cases) == 2
    assert render_index.equation_number(unnumbered) is None
    assert render_index.equation_number(unsupported) is None
    warning_codes = {issue.code for issue in document.validate().warnings}
    assert "unsupported-latex-command" in warning_codes

    unnumbered_target = Equation("x=1", numbered=False)
    invalid_reference = Document(
        "Invalid Equation Reference",
        Paragraph("See ", unnumbered_target.reference(), "."),
        unnumbered_target,
    )
    error_codes = {issue.code for issue in invalid_reference.validate().errors}
    assert "unnumbered-equation-reference" in error_codes

    docx_path = tmp_path / "amsmath.docx"
    pdf_path = tmp_path / "amsmath.pdf"
    html_path = tmp_path / "amsmath.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_pdf_text = " ".join(pdf_text.split())
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    assert "See Eq. 1 and the loss definition." in normalized_pdf_text
    assert "a = b + c" in docx_xml
    assert "= d" in docx_xml
    assert "f(x) = { 0, if x &lt; 0" in html_text
    assert "x<sup>2</sup>, if x &gt;= 0" in html_text
    assert "a = b + c" in normalized_html_text
    assert "= d" in normalized_html_text
    assert "Eq. 1" in normalized_html_text
    assert "(1)" in normalized_html_text
    assert "(2)" in normalized_html_text
    assert "(3)" not in normalized_html_text


def test_mhchem_formula_and_reaction_render_to_all_outputs(tmp_path: Path) -> None:
    from oodocs.chem import ce

    water = chemical_formula("H2O")
    sulfate = ChemicalFormula("SO4^2-")
    ammonium = ce("NH4+")
    reaction = ReactionEquation("2H2 + O2 -> 2H2O")
    document = Document(
        "Chemistry Blocks",
        Paragraph("Water is ", water, "; sulfate is ", sulfate, "; ammonium is ", ammonium, "."),
        Paragraph("See ", reaction.reference(), "."),
        reaction,
    )

    assert ChemicalFormula(r"\ce{CO2}").value == "CO_{2}"
    assert water.plain_text() == "H2O"
    assert sulfate.plain_text() == "SO42-"
    assert ammonium.value == "NH_{4}^{+}"
    assert build_render_index(document).equation_number(reaction) == 1
    assert reaction.reference_text(1) == "Reaction 1"

    docx_path = tmp_path / "mhchem.docx"
    pdf_path = tmp_path / "mhchem.pdf"
    html_path = tmp_path / "mhchem.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    formula_paragraph = next(paragraph for paragraph in word_document.paragraphs if "Water is" in paragraph.text)
    reaction_paragraph = next(paragraph for paragraph in word_document.paragraphs if "2H2 + O2" in paragraph.text)
    assert any(run.text == "2" and run.font.subscript for run in formula_paragraph.runs)
    assert any(run.text == "2-" and run.font.superscript for run in formula_paragraph.runs)
    assert any(run.text == "+" and run.font.superscript for run in formula_paragraph.runs)
    assert any(run.text == "2" and run.font.subscript for run in reaction_paragraph.runs)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_pdf_text = " ".join(pdf_text.split())
    assert "See Reaction 1." in normalized_pdf_text
    assert "2H2 + O2 -> 2H2O" in normalized_pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "H<sub>2</sub>O" in html_text
    assert "SO<sub>4</sub><sup>2-</sup>" in html_text
    assert "NH<sub>4</sub><sup>+</sup>" in html_text
    assert "2H<sub>2</sub> + O<sub>2</sub> -&gt; 2H<sub>2</sub>O" in html_text
    assert "Reaction 1" in normalized_html_text
    assert "(1)" in normalized_html_text


def test_method_style_inline_actions_create_renderable_fragments() -> None:
    source = CitationSource("Usage Guide", key="guide", year="2026")
    library = CitationLibrary([source])

    assert Text.bold("important").style.bold is True
    assert Text.italic("note").style.italic is True
    assert Text.inline_code("x = 1").style.font_name == "Courier New"
    assert Text.text_color("accent", "#0066AA").style.text_color == "0066AA"
    assert Text.highlight("marked", "#FFF2CC").style.highlight_color == "FFF2CC"
    assert Text.strikethrough("old").style.strikethrough is True
    assert Text.superscript("2").style.superscript is True
    assert Text.subscript("0").style.subscript is True
    assert [fragment.value for fragment in Text.from_markup("plain **bold**")] == [
        "plain ",
        "bold",
    ]
    assert bold("important").style.bold is True
    assert italic("note").style.italic is True
    assert inline_code("x = 1").style.font_name == "Courier New"
    assert text_color("accent", "#0066AA").style.text_color == "0066AA"
    assert highlight("marked", "#FFF2CC").style.highlight_color == "FFF2CC"
    assert strikethrough("old").style.strikethrough is True
    assert superscript("2").style.superscript is True
    assert subscript("0").style.subscript is True
    assert [fragment.plain_text() for fragment in prescript("14", "6", "C")] == ["14", "6", "C"]
    assert line_break().plain_text() == "\n"
    external_link = link("https://example.com", "Example")
    assert isinstance(external_link, inline_components.Hyperlink)
    assert external_link.target == "https://example.com"
    assert external_link.plain_text() == "Example"
    long_target = "https://example.com/releases/" + "stable-artifact-" * 8
    visible_url = url(long_target)
    assert isinstance(visible_url, inline_components.Hyperlink)
    assert visible_url.target == long_target
    assert "\u200b" in visible_url.plain_text()
    assert visible_url.plain_text().replace("\u200b", "") == long_target
    assert url(long_target, label="Project site").plain_text() == "Project site"
    assert source.cite().plain_text() == "[?]"
    assert library.cite("guide").plain_text() == "[?]"
    assert Comment.annotated("term", "Expanded note").plain_text() == "term[?]"
    assert Footnote.annotated("term", "Portable footnote note").plain_text() == "term[?]"
    assert Math.inline(r"\alpha^2").plain_text() == "alpha2"
    assert InlineChip("base").kind == "chip"
    assert tag("api").chip_style == "tag"
    assert tag("api", background_color="#EEF2FF").chip_style.background_color == "EEF2FF"
    assert badge(3).chip_style == "badge"
    assert status("ready", state="success").chip_style == "status.success"
    assert keyboard("Ctrl+Enter").chip_style == "keyboard"
    assert isinstance(tag("api"), InlineChip)
    assert isinstance(InlineChipStyle(text_color="#111827"), InlineChipStyle)

    try:
        status("ready", state="unknown")
    except ValueError as exc:
        assert "Unsupported status state" in str(exc)
    else:
        raise AssertionError("Expected invalid status state validation to fail")


def test_theme_validates_page_number_configuration() -> None:
    theme = Theme(
        page_numbers=PageNumberDefaults(
            show_page_numbers=True,
            page_number_template="Page {page}",
            page_number_alignment="right",
        )
    )

    assert theme.format_page_number(3) == "Page 3"
    assert theme.format_page_number(3, front_matter=True) == "Page iii"

    try:
        Theme(page_numbers=PageNumberDefaults(page_number_template="Page"))
    except ValueError as exc:
        assert "{page}" in str(exc)
    else:
        raise AssertionError("Expected page_number_template validation to fail")


def test_theme_accepts_grouped_defaults_objects() -> None:
    theme = Theme(
        typography=TypographyDefaults(body_font_name="Calibri", body_font_size=10.0),
        captions=CaptionDefaults(figure_label="Fig.", table_caption_position="below"),
        citations=CitationDefaults(citation_style="apa", reference_style="apa"),
        generated_content=GeneratedContentDefaults(table_of_contents_title="Outline"),
        page_numbers=PageNumberDefaults(show_page_numbers=True, page_number_template="p. {page}"),
        title_matter=TitleMatterDefaults(title_text_alignment="left"),
        blocks=BlockDefaults(
            paragraph_text_alignment="left",
            table_block_alignment="right",
            run_in_title_style=RunInTitleStyle(TextStyle(italic=True), separator=": "),
        ),
    )

    assert theme.typography.body_font_name == "Calibri"
    assert theme.typography.body_font_size == 10.0
    assert theme.captions.figure_label == "Fig."
    assert theme.captions.table_caption_position == "below"
    assert theme.citations.citation_style == "apa"
    assert theme.citations.reference_style == "apa"
    assert theme.citations.citation_style == "apa"
    assert theme.generated_content.table_of_contents_title == "Outline"
    assert theme.page_numbers.show_page_numbers is True
    assert theme.format_page_number(4) == "p. 4"
    assert theme.title_matter.title_text_alignment == "left"
    assert theme.blocks.paragraph_text_alignment == "left"
    assert theme.blocks.run_in_title_style.text_style.italic is True
    assert theme.blocks.run_in_title_style.separator == ": "
    assert theme.blocks.table_block_alignment == "right"

    title_style_theme = Theme(
        blocks=BlockDefaults(
            run_in_title_style=RunInTitleStyle(TextStyle(bold=True), separator=". ")
        ),
    )
    assert title_style_theme.blocks.run_in_title_style.text_style.bold is True
    assert title_style_theme.blocks.run_in_title_style.separator == ". "

    keyword_group = Theme(typography=TypographyDefaults(body_font_name="Aptos"))
    assert keyword_group.typography.body_font_name == "Aptos"
    generated_keyword_group = Theme(
        generated_content=GeneratedContentDefaults(reference_list_title="Bibliography")
    )
    assert generated_keyword_group.generated_content.reference_list_title == "Bibliography"
    try:
        Theme(typography=object())  # type: ignore[arg-type]
    except TypeError as exc:
        assert "Theme.typography" in str(exc)
    else:
        raise AssertionError("Expected invalid Theme defaults group to fail")


def test_common_block_styles_accept_direct_kwargs() -> None:
    paragraph = Paragraph("Right aligned", text_alignment="right", space_after=4)
    code_block = CodeBlock(
        "print('x')",
        language="python",
        caption="Minimal listing.",
        identifier="minimal-listing",
        line_numbers=True,
        highlight_lines={1},
        left_indent=0.25,
    )
    equation = Equation("x=1", space_after=2)
    bullet_list = BulletList("one", indent=0.4)
    numbered_list = NumberedList(
        "one",
        marker=CounterStyle(suffix=")"),
        start=3,
        item_spacing=5,
        block_spacing=9,
    )
    box = Box(Paragraph("inside"), background_color="#FFFFFF", padding=Padding.all(8), width=3.0)
    table = Table(
        headers=["A"],
        rows=[["B"]],
        header_background_color="#AABBCC",
        cell_text_alignment="center",
    )
    contents = TableOfContents(level_styles={1: {"bold": False, "space_after": 1}})

    assert paragraph.style.text_alignment == "right"
    assert paragraph.style.space_after == 4
    assert code_block.style.left_indent == 0.25
    assert code_block.show_language is True
    assert code_block.language_position == "top-right"
    assert code_block.caption is not None
    assert code_block.caption.plain_text() == "Minimal listing."
    assert code_block.identifier == "minimal-listing"
    assert code_block.line_numbers is True
    assert code_block.highlight_lines == frozenset({1})
    assert code_block.display_line(1, "print('x')") == "1 | print('x')"
    assert equation.style.space_after == 2
    assert bullet_list.style is not None
    assert bullet_list.style.marker.counter_format == "bullet"
    assert bullet_list.style.indent == 0.4
    assert numbered_list.style is not None
    assert numbered_list.start == 3
    assert numbered_list.style.marker.suffix == ")"
    assert numbered_list.style.item_spacing == 5
    assert numbered_list.style.block_spacing == 9
    assert box.style.background_color == "FFFFFF"
    assert box.style.padding == Padding.all(8)
    assert box.style.width == 3.0
    assert table.style.header_background_color == "AABBCC"
    assert table.style.cell_text_alignment == "center"
    assert contents.style_for_level(1).bold is False
    assert contents.style_for_level(1).space_after == 1

    hidden_label_block = CodeBlock("plain text", language="text", show_language=False, language_position="bottom-left")
    assert hidden_label_block.show_language is False
    assert hidden_label_block.language_position == "bottom-left"

    try:
        CodeBlock("x", language="python", language_position="middle")  # type: ignore[arg-type]
    except ValueError as exc:
        assert "language_position" in str(exc)
    else:
        raise AssertionError("Expected invalid CodeBlock language_position to fail")

    try:
        CodeBlock("x", highlight_lines={0})
    except ValueError as exc:
        assert "highlight_lines" in str(exc)
    else:
        raise AssertionError("Expected invalid CodeBlock highlight_lines to fail")


def test_paragraph_style_defaults_to_justify_alignment() -> None:
    paragraph = Paragraph("Default alignment paragraph.")

    assert paragraph.style.text_alignment is None
    assert Theme().resolve_paragraph_text_alignment(paragraph.style) == "justify"


def test_document_and_paragraph_text_alignment_options_render(tmp_path: Path) -> None:
    document = Document(
        "Alignment Test",
        Paragraph("Document-level centered paragraph."),
        Paragraph(
            "Paragraph-level right paragraph.",
            style=ParagraphStyle(text_alignment="right"),
        ),
        settings=DocumentSettings(
            theme=Theme(blocks=BlockDefaults(paragraph_text_alignment="center"))
        ),
    )

    docx_path = tmp_path / "alignment.docx"
    pdf_path = tmp_path / "alignment.pdf"
    html_path = tmp_path / "alignment.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    centered = next(
        paragraph
        for paragraph in word_document.paragraphs
        if paragraph.text == "Document-level centered paragraph."
    )
    right = next(
        paragraph
        for paragraph in word_document.paragraphs
        if paragraph.text == "Paragraph-level right paragraph."
    )
    assert centered.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert right.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    assert "Document-level centered paragraph." in pdf_text
    assert "Paragraph-level right paragraph." in pdf_text
    assert "text-align: center" in html_text
    assert "text-align: right" in html_text


def test_paragraph_titles_render_and_inherit_styles(tmp_path: Path) -> None:
    document = Document(
        "Paragraph Titles",
        Paragraph("Default body.", title="Default"),
        Section(
            "Scoped",
            Paragraph("Scoped body.", title="Scoped"),
            Paragraph(
                "Direct body.",
                title="Direct",
                title_style=RunInTitleStyle(
                    TextStyle(bold=True, text_color="991B1B"),
                    separator=" - ",
                ),
            ),
            run_in_title_style=RunInTitleStyle(
                TextStyle(bold=True, italic=True),
                separator=": ",
            ),
        ),
        Paragraph(
            "Individual body.",
            title="Individual",
            title_style=RunInTitleStyle(
                TextStyle(bold=True, text_color="166534"),
                separator=". ",
            ),
        ),
        settings=DocumentSettings(
            theme=Theme(
                blocks=BlockDefaults(
                    run_in_title_style=RunInTitleStyle(
                        TextStyle(bold=True),
                        separator=" ",
                    )
                )
            )
        ),
    )

    assert document.body.children[0].plain_text() == "Default Default body."

    docx_path = tmp_path / "paragraph-titles.docx"
    pdf_path = tmp_path / "paragraph-titles.pdf"
    html_path = tmp_path / "paragraph-titles.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    default = next(paragraph for paragraph in word_document.paragraphs if paragraph.text == "Default Default body.")
    scoped = next(paragraph for paragraph in word_document.paragraphs if paragraph.text == "Scoped: Scoped body.")
    direct = next(paragraph for paragraph in word_document.paragraphs if paragraph.text == "Direct - Direct body.")
    individual = next(paragraph for paragraph in word_document.paragraphs if paragraph.text == "Individual. Individual body.")
    assert default.runs[0].text == "Default"
    assert default.runs[0].bold is True
    assert scoped.runs[0].italic is True
    assert direct.runs[0].italic in (None, False)
    assert direct.runs[0].font.color.rgb == RGBColor(0x99, 0x1B, 0x1B)
    assert individual.runs[0].font.color.rgb == RGBColor(0x16, 0x65, 0x34)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html = html_path.read_text(encoding="utf-8")
    assert "Default Default body." in pdf_text
    assert "Scoped: Scoped body." in pdf_text
    assert "Direct - Direct body." in pdf_text
    assert "Individual. Individual body." in pdf_text
    assert "Default Default body." in unescape(re.sub(r"<[^>]+>", "", html))
    assert "Scoped: Scoped body." in unescape(re.sub(r"<[^>]+>", "", html))
    assert "Direct - Direct body." in unescape(re.sub(r"<[^>]+>", "", html))
    assert "#991B1B" in html
    assert "#166534" in html


def test_paragraph_style_supports_word_like_indents() -> None:
    style = ParagraphStyle(
        left_indent=1.27,
        right_indent=0.508,
        first_line_indent=0.635,
        unit="cm",
    )
    hanging = ParagraphStyle.hanging(left=1.524, by=0.635, unit="cm")

    assert style.unit == "cm"
    assert round(style.left_indent_in_inches("in"), 2) == 0.5
    assert round(style.right_indent_in_inches("in"), 2) == 0.2
    assert round(style.first_line_indent_in_inches("in"), 2) == 0.25
    assert hanging.unit == "cm"
    assert round(hanging.left_indent_in_inches("in"), 2) == 0.6
    assert round(hanging.first_line_indent_in_inches("in"), 2) == -0.25

    try:
        ParagraphStyle(left_indent=-0.1)
    except ValueError as exc:
        assert "left_indent" in str(exc)
    else:
        raise AssertionError("Expected negative left_indent validation to fail")

    try:
        ParagraphStyle.hanging(by=-0.1)
    except ValueError as exc:
        assert "hanging by" in str(exc)
    else:
        raise AssertionError("Expected negative hanging indent validation to fail")


def test_theme_defaults_center_media_objects_and_captions() -> None:
    theme = Theme()

    assert theme.blocks.page_background_color == "FFFFFF"
    assert theme.captions.caption_text_alignment == "center"
    assert theme.blocks.table_block_alignment == "center"
    assert theme.blocks.figure_block_alignment == "center"
    assert theme.blocks.box_block_alignment == "center"


def test_page_background_color_renders_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Color Test",
        Paragraph("Tinted page."),
        settings=DocumentSettings(
            theme=Theme(blocks=BlockDefaults(page_background_color="#F4F8FC"))
        ),
    )
    docx_path = tmp_path / "color.docx"
    pdf_path = tmp_path / "color.pdf"
    html_path = tmp_path / "color.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    assert 'w:background w:color="F4F8FC"' in _docx_document_xml(docx_path)
    assert b".956863 .972549 .988235 rg" in _pdf_content_bytes(pdf_path)
    assert "background: #F4F8FC;" in html_path.read_text(encoding="utf-8")


def test_document_save_infers_renderer_from_extension(tmp_path: Path) -> None:
    document = Document("Save Test", Paragraph("Saved by extension."))

    docx_path = document.save(tmp_path / "extension.docx")
    pdf_path = document.save(tmp_path / "extension.pdf")
    html_path = document.save(tmp_path / "extension.html")
    htm_path = document.save(tmp_path / "extension.htm")

    assert docx_path.exists()
    assert pdf_path.exists()
    assert html_path.exists()
    assert htm_path.exists()

    try:
        document.save(tmp_path / "extension.txt")
    except ValueError as exc:
        assert ".docx" in str(exc)
        assert ".pdf" in str(exc)
        assert ".html" in str(exc)
    else:
        raise AssertionError("Expected unsupported save extension to fail")


def test_document_save_all_renders_multiple_formats(tmp_path: Path) -> None:
    document = Document("Quarterly Review Draft", Paragraph("Saved as a bundle."))

    outputs = document.save_all(tmp_path)

    assert isinstance(outputs, OutputBundle)
    assert sorted(outputs.formats) == ["docx", "html", "pdf"]
    assert outputs["docx"] == tmp_path / "quarterly-review-draft.docx"
    assert outputs["pdf"] == tmp_path / "quarterly-review-draft.pdf"
    assert outputs["html"] == tmp_path / "quarterly-review-draft.html"
    assert all(path.exists() for path in outputs.values())
    assert sorted(output_format for output_format, _path in outputs) == ["docx", "html", "pdf"]

    selected_outputs = document.save_all(
        tmp_path / "selected",
        stem="review-pack",
        formats=(".docx", "htm"),
    )

    assert sorted(selected_outputs.keys()) == ["docx", "html"]
    assert selected_outputs["docx"] == tmp_path / "selected" / "review-pack.docx"
    assert selected_outputs["html"] == tmp_path / "selected" / "review-pack.html"
    assert all(path.exists() for path in selected_outputs.values())

    try:
        document.save_all(tmp_path, formats=("txt",))
    except ValueError as exc:
        assert "docx" in str(exc)
        assert "pdf" in str(exc)
        assert "html" in str(exc)
    else:
        raise AssertionError("Expected unsupported save_all format to fail")


def test_document_validate_returns_printable_result_with_format_scopes() -> None:
    document = Document(
        "Validation Report",
        TableOfContents(),
        Chapter("Findings", Paragraph("Body.")),
    )

    result = document.validate()

    assert isinstance(result, ValidationResult)
    assert result.ok
    assert [issue.code for issue in result.warnings] == ["html-toc-page-numbers"]
    assert result.warnings[0].formats == ("html",)
    assert result.warnings_for(("html",))
    assert result.warnings_for(("docx",)) == ()

    printable = str(result)
    assert "OODocs validation ok for All" in printable
    assert "Severity" in printable
    assert "Formats" in printable
    assert "HTML" in printable
    assert "html-toc-page-numbers" in printable


def test_named_styles_validate_and_render_across_formats(tmp_path: Path) -> None:
    styles = StyleSheet.default()
    styles.register("paragraph", "lead", ParagraphStyle(space_after=4, text_alignment="center"))
    styles.register("list", "steps", ListStyle(marker=CounterStyle()))
    styles.register("table_cell", "positive", TableCellStyle(text_color="166534", bold=True))

    document = Document(
        "Named Styles",
        Paragraph("Centered lead paragraph.", style="lead"),
        Box(
            Paragraph("The box uses a named style from the theme stylesheet."),
            title="Finding",
            style="info",
        ),
        NumberedList("Verify named list styles.", style="steps"),
        Table(
            headers=[TableCell("Metric", style="emphasis"), "Value"],
            rows=[["Accuracy", TableCell("98%", style="positive")]],
            caption="Styled metrics.",
            style="evidence",
            row_styles={0: "muted"},
            column_styles={1: "numeric"},
        ),
        Paragraph("Status: ", status("ready", chip_style="status.success")),
        settings=DocumentSettings(theme=Theme(stylesheet=styles)),
    )

    result = document.validate()
    assert result.errors == ()

    outputs = document.save_all(tmp_path, stem="named-styles")
    assert set(outputs.keys()) == {"docx", "html", "pdf"}
    assert all(path.exists() for path in outputs.values())

    html = outputs["html"].read_text(encoding="utf-8")
    assert "READY" in html
    assert "#ECFDF3" in html


def test_named_style_css_classes_render_in_html(tmp_path: Path) -> None:
    styles = StyleSheet.default()
    styles.register("paragraph", "lead.css", ParagraphStyle(css_class="lead dense"))
    styles.register("box", "panel.css", BoxStyle(css_class="panel highlight"))
    styles.register("table", "matrix.css", TableStyle(css_class="matrix"))
    styles.register("chip", "chip.css", InlineChipStyle(css_class="pill beta"))
    document = Document(
        "CSS Hooks",
        Paragraph("Lead paragraph.", style="lead.css"),
        Box(Paragraph("Panel body."), title="Panel", style="panel.css"),
        Table(["A"], [["B"]], caption="Matrix.", style="matrix.css"),
        Paragraph("Flag: ", tag("beta", chip_style="chip.css")),
        settings=DocumentSettings(theme=Theme(stylesheet=styles)),
    )
    html_path = tmp_path / "css-hooks.html"

    document.save_html(html_path)

    html = html_path.read_text(encoding="utf-8")
    assert 'class="oodocs-paragraph lead dense"' in html
    assert 'class="oodocs-box panel highlight"' in html
    assert 'class="oodocs-table matrix"' in html
    assert 'class="oodocs-inline-chip oodocs-inline-chip-tag pill beta"' in html


def test_document_validate_reports_unknown_and_wrong_named_styles() -> None:
    document = Document(
        "Named Style Errors",
        Paragraph("Unknown paragraph style.", style="missing"),
        BulletList("Unknown list style.", style="missing-list"),
        Table(
            headers=["A"],
            rows=[[TableCell("B", style="info")]],
            style="info",
        ),
    )

    result = document.validate()
    assert [issue.code for issue in result.errors] == [
        "unknown-style-name",
        "unknown-style-name",
        "wrong-style-category",
        "wrong-style-category",
    ]
    assert result.errors[0].path == "document.body.children[0].style"


def test_document_validate_reports_authoring_errors_and_blocks_render(
    tmp_path: Path,
) -> None:
    missing_image = tmp_path / "missing.png"
    document = Document(
        "Broken Document",
        Figure(missing_image, caption="Missing figure."),
    )

    result = document.validate()
    assert not result.ok
    assert [issue.code for issue in result.errors] == ["missing-image-file"]
    assert result.errors[0].formats == ("docx", "pdf", "html")
    assert "All" in str(result)

    pdf_path = tmp_path / "broken.pdf"
    try:
        document.save_pdf(pdf_path)
    except DocumentValidationError as exc:
        message = str(exc)
        assert "OODocs validation failed for PDF" in message
        assert "missing-image-file" in message
        assert "PDF" in message
    else:
        raise AssertionError("Expected invalid document rendering to stop before PDF build")
    assert not pdf_path.exists()


def test_image_data_sources_render_without_temp_files(tmp_path: Path) -> None:
    figure = Figure(ImageData(_build_sample_png()), caption="In-memory image.", width=1.0)
    document = Document("Image Data", figure)

    outputs = document.save_all(tmp_path, stem="image-data")

    assert set(outputs.keys()) == {"docx", "pdf", "html"}
    assert all(path.exists() for path in outputs.values())
    assert "data:image/png;base64," in outputs["html"].read_text(encoding="utf-8")


def test_figure_crop_rotation_and_alt_text_render_to_outputs(
    tmp_path: Path,
) -> None:
    image_path = tmp_path / "graphicx.png"
    image_path.write_bytes(_build_sample_png(width=120, height=80))
    figure = Figure(
        image_path,
        caption="Transformed figure.",
        width=1.0,
        crop=CropBox(left=10, right=20, top=5, bottom=15, unit="px"),
        rotation=90,
        alt_text="Cropped rotated sample",
    )
    document = Document("Graphicx Figure", figure)

    with pytest.raises(ValueError, match="CropBox"):
        CropBox(left=-1)
    with pytest.raises(ValueError, match="rotation"):
        Figure(image_path, rotation=float("nan"))

    docx_path = tmp_path / "graphicx.docx"
    pdf_path = tmp_path / "graphicx.pdf"
    html_path = tmp_path / "graphicx.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    html_text = html_path.read_text(encoding="utf-8")
    match = re.search(r'src="data:image/png;base64,([^"]+)" alt="Cropped rotated sample"', html_text)

    assert match is not None
    assert 'descr="Cropped rotated sample"' in docx_xml
    assert _pdf_image_draw_count(pdf_path) == 1

    from PIL import Image

    html_image = Image.open(BytesIO(b64decode(match.group(1))))
    assert html_image.size == (60, 90)
    with zipfile.ZipFile(docx_path) as archive:
        media_names = [
            name for name in archive.namelist()
            if name.startswith("word/media/")
        ]
        assert media_names
        docx_image = Image.open(BytesIO(archive.read(media_names[0])))
        assert docx_image.size == (60, 90)


def test_document_validate_catches_reference_mistakes() -> None:
    orphan = Paragraph("Outside the document.")
    uncaptioned_table = Table(headers=["Area"], rows=[["Validation"]])
    document = Document(
        "Reference Mistakes",
        Paragraph("See ", orphan.reference(), " and ", uncaptioned_table.reference(), "."),
        uncaptioned_table,
    )

    result = document.validate()
    assert [issue.code for issue in result.errors] == [
        "missing-reference-target",
        "uncaptioned-reference-target",
    ]

    try:
        document.validate(raise_on_error=True)
    except DocumentValidationError as exc:
        assert exc.errors == result.errors
        assert "missing-reference-target" in str(exc)
    else:
        raise AssertionError("Expected validate(raise_on_error=True) to fail")


def test_document_validate_reports_preflight_warnings(tmp_path: Path) -> None:
    image_path = tmp_path / "plot.png"
    _write_sample_image(image_path)
    missing_inline_image = tmp_path / "missing-inline.png"
    settings = DocumentSettings(
        unit="in",
        page_size=PageSize.letter(),
        page_margins=PageMargins.all(1.0, unit="in"),
    )
    wide_table = Table(
        headers=["A", "B"],
        rows=[["wide", "table"]],
        column_widths=[4.0, 4.0],
        unit="in",
    )
    captionless_figure = Figure(image_path)
    document = Document(
        "Preflight",
        wide_table,
        captionless_figure,
        Paragraph(
            "Inline asset ",
            ImageBox(
                missing_inline_image,
                placement="inline",
                width=0.2,
                height=0.2,
            ),
        ),
        settings=settings,
    )

    result = document.validate()
    warning_codes = {issue.code for issue in result.warnings}
    error_codes = {issue.code for issue in result.errors}

    assert warning_codes >= {
        "missing-table-caption",
        "missing-figure-caption",
        "wide-table",
    }
    assert "missing-image-file" in error_codes
    wide_table_warnings = [
        issue for issue in result.warnings
        if issue.code == "wide-table"
    ]
    assert any(issue.formats == ("docx", "pdf") for issue in wide_table_warnings)
    assert any("ColumnSpec(flex=...)" in issue.message for issue in wide_table_warnings)
    assert any("save_csv" in issue.message for issue in wide_table_warnings)


def test_document_validate_treats_subfigure_group_caption_as_reference_target(
    tmp_path: Path,
) -> None:
    image_path = tmp_path / "plot.png"
    _write_sample_image(image_path)
    grouped = SubFigureGroup(
        SubFigure(image_path, width=1.0),
        SubFigure(image_path, width=1.0),
        caption="Grouped result.",
    )
    captionless_group = SubFigureGroup(
        SubFigure(image_path, width=1.0),
    )

    valid_result = Document(
        "Grouped",
        Paragraph("See ", grouped.reference(), "."),
        grouped,
    ).validate()
    warning_codes = {issue.code for issue in valid_result.warnings}
    assert "missing-figure-caption" not in warning_codes
    assert valid_result.ok

    invalid_result = Document(
        "Captionless Group",
        Paragraph("See ", captionless_group.reference(), "."),
        captionless_group,
    ).validate()
    assert "missing-figure-caption" in {issue.code for issue in invalid_result.warnings}
    assert "uncaptioned-reference-target" in {issue.code for issue in invalid_result.errors}


def test_numbering_and_list_styles_are_customizable() -> None:
    heading_numbering = HeadingNumbering(
        level_styles=(
            CounterStyle(counter_format="upper-roman"),
            CounterStyle(counter_format="lower-alpha"),
        ),
        prefix="[",
        suffix="]",
    )
    ordered_style = ListStyle(
        marker=CounterStyle(counter_format="upper-roman", prefix="(", suffix=")")
    )
    bullet_style = ListStyle(
        marker=CounterStyle(counter_format="bullet", bullet="\u2022", suffix="")
    )

    assert heading_numbering.format_label([2, 3]) == "[II.c]"
    assert ordered_style.marker_for(0) == "(I)"
    assert ordered_style.marker_for(2) == "(III)"
    assert bullet_style.marker_for(1) == "\u2022"


def test_numbered_list_start_resume_and_spacing_render_to_outputs(tmp_path: Path) -> None:
    setup_steps = NumberedList(
        "Create the model",
        "Render the files",
        start=4,
        item_spacing=2,
        block_spacing=6,
    )
    follow_up = NumberedList("Publish the bundle", resume_from=setup_steps)
    nested = BulletList(
        "Inspect follow-up tasks",
        item_children=[[NumberedList("Audit links", start=2)]],
        item_spacing=7,
        block_spacing=11,
    )
    document = Document("Enumitem Lists", setup_steps, follow_up, nested)

    docx_path = tmp_path / "enumitem-lists.docx"
    pdf_path = tmp_path / "enumitem-lists.pdf"
    html_path = tmp_path / "enumitem-lists.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    word_text = "\n".join(
        [paragraph.text for paragraph in word_document.paragraphs]
        + [
            cell.text
            for table in word_document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_pdf_text = " ".join(pdf_text.split())
    html_text = _normalized_html_text(html_path)
    html_markup = html_path.read_text(encoding="utf-8")

    for expected in (
        "4. Create the model",
        "5. Render the files",
        "6. Publish the bundle",
        "2. Audit links",
    ):
        assert expected in word_text
        assert expected in normalized_pdf_text
        assert expected in html_text
    assert "margin: 0.0pt 0 7.0pt;" in html_markup
    assert "margin-bottom: 11.0pt;" in html_markup


def test_glossary_acronym_and_generated_list_render_to_outputs(tmp_path: Path) -> None:
    glossary = Glossary()
    glossary.term("SLO", "Service level objective")
    glossary.acronym("HVAC", "Heating, ventilation, and air conditioning")

    assert isinstance(glossary.get("HVAC"), Acronym)
    assert isinstance(glossary.get("SLO"), GlossaryTerm)

    document = Document(
        "Glossary Document",
        Paragraph(
            glossary.use("HVAC"),
            " controls support the ",
            glossary.use("SLO"),
            ". Later references use ",
            glossary.use("HVAC"),
            ".",
        ),
        GlossaryList(glossary, sort="key"),
    )

    result = document.validate()
    assert result.ok

    docx_path = tmp_path / "glossary.docx"
    pdf_path = tmp_path / "glossary.pdf"
    html_path = tmp_path / "glossary.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    word_text = "\n".join(
        [paragraph.text for paragraph in word_document.paragraphs]
        + [
            cell.text
            for table in word_document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_pdf_text = " ".join(pdf_text.split())
    html_text = _normalized_html_text(html_path)

    for expected in (
        "Heating, ventilation, and air conditioning (HVAC)",
        "Later references use HVAC.",
        "Service level objective",
        "Glossary",
    ):
        assert expected in word_text
        assert expected in normalized_pdf_text
        assert expected in html_text

    duplicate_glossary = Glossary()
    duplicate_glossary.term("API", "Application programming interface")
    duplicate_glossary.term("API", "Duplicate application term")
    duplicate_result = Document("Duplicate Glossary", GlossaryList(duplicate_glossary)).validate()
    assert "duplicate-glossary-key" in {issue.code for issue in duplicate_result.errors}

    empty_result = Document("Empty Glossary", GlossaryList(Glossary())).validate()
    assert "empty-glossary-list" in {issue.code for issue in empty_result.warnings}


def test_locale_theme_localizes_html_language_and_generated_labels(tmp_path: Path) -> None:
    glossary = Glossary()
    glossary.term("api", "Application programming interface", term="API")
    table = Table(["항목"], [["값"]], caption="측정값")
    source = CitationSource("Locale reference", key="locale2026", authors=("Kim",), year="2026")
    document = Document(
        "지역화 문서",
        Paragraph("본문 ", cite("locale2026"), "."),
        table,
        ListOfTables(),
        GlossaryList(glossary),
        ReferenceList(),
        settings=DocumentSettings(theme=Theme.from_locale("ko-KR")),
        citations=[source],
    )
    html_path = tmp_path / "locale.html"

    document.save_html(html_path)

    html_markup = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert '<html lang="ko-KR">' in html_markup
    assert "표 1. 측정값" in normalized_html_text
    assert "표 목록" in normalized_html_text
    assert "용어집" in normalized_html_text
    assert "용어 정의" in normalized_html_text
    assert "참고문헌" in normalized_html_text


def test_heading_style_renders_across_outputs(tmp_path: Path) -> None:
    heading_style = HeadingStyle(
        text_style=TextStyle(
            font_size=17,
            bold=False,
            italic=True,
            text_color="#1A2B3C",
        ),
        space_before=4,
        space_after=5,
        leading=19,
        text_alignment="right",
        numbering=CounterStyle(counter_format="upper-alpha", suffix=")"),
    )
    document = Document(
        "Heading Style",
        Section(
            "Styled",
            Paragraph("Body."),
            level=2,
            heading_style=heading_style,
        ),
    )

    docx_path = tmp_path / "heading-style.docx"
    html_path = tmp_path / "heading-style.html"
    pdf_path = tmp_path / "heading-style.pdf"
    document.save_docx(docx_path)
    document.save_html(html_path)
    document.save_pdf(pdf_path)

    word_document = WordDocument(docx_path)
    heading = next(
        paragraph
        for paragraph in word_document.paragraphs
        if paragraph.text == "1.A) Styled"
    )
    assert heading.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert heading.paragraph_format.space_before.pt == 4
    assert heading.paragraph_format.space_after.pt == 5
    assert heading.paragraph_format.line_spacing.pt == 19
    assert heading.runs[0].font.size.pt == 17
    assert heading.runs[0].font.bold is False
    assert heading.runs[0].font.italic is True
    assert heading.runs[0].font.color.rgb == RGBColor.from_string("1A2B3C")

    html = html_path.read_text(encoding="utf-8")
    assert "1.A) Styled" in html
    assert "font-size: 17.0pt" in html
    assert "text-align: right" in html
    assert "margin: 4.0pt 0 5.0pt" in html
    assert "line-height: 19.0pt" in html
    assert "font-weight: 400" in html
    assert "font-style: italic" in html
    assert "color: #1A2B3C" in html

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf_path).pages)
    assert "1.A) Styled" in pdf_text


def test_table_accepts_dataframe_like_inputs_and_spans() -> None:
    dataframe = FakeDataFrame(
        columns=[("Metrics", "Latency"), ("Metrics", "Quality"), ("Summary", "")],
        rows=[["14 ms", "stable", "ready"]],
    )
    table = Table(
        dataframe,
        caption="Span test.",
        column_widths=[1.5, 1.5, 1.5],
        style=TableStyle(alternate_row_background_color="#F4F8FC"),
    )
    merged_header = Table(
        headers=[
            [TableCell("Metrics", colspan=2), TableCell("Summary", rowspan=2)],
            ["Latency", "Quality"],
        ],
        rows=[["14 ms", "stable", "ready"]],
        column_widths=[1.5, 1.5, 1.5],
    )

    assert len(table.header_rows) == 2
    assert table.header_rows[0][0].colspan == 2
    assert table.header_rows[0][1].rowspan == 2
    assert table._layout().column_count == 3
    assert merged_header._layout().row_count == 3


def test_table_grouped_headers_helper_renders_spans(
    tmp_path: Path,
) -> None:
    table = Table.grouped_headers(
        groups=[("Geometry", 2), ("Performance", 2)],
        columns=["Page", "Orientation", "Latency", "Status"],
        rows=[
            [TableCell("Letter", rowspan=2), "portrait", "14 ms", "ok"],
            ["landscape", "12 ms", "ok"],
        ],
        caption="Grouped renderer matrix.",
    )
    document = Document("Grouped Headers", table)

    assert table.header_rows[0][0].colspan == 2
    assert table.header_rows[0][1].colspan == 2
    assert table.rows[0][0].rowspan == 2

    with pytest.raises(ValueError, match="group spans"):
        Table.grouped_headers(
            groups=[("Only one", 1)],
            columns=["A", "B"],
            rows=[["a", "b"]],
        )

    docx_path = tmp_path / "grouped-headers.docx"
    pdf_path = tmp_path / "grouped-headers.pdf"
    html_path = tmp_path / "grouped-headers.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf_path).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert "Geometry" in docx_xml
    assert "Grouped renderer matrix." in pdf_text
    assert "landscape" in pdf_text
    assert 'colspan="2"' in html_text
    assert 'rowspan="2"' in html_text


def test_table_cell_alignment_options_are_validated() -> None:
    cell = TableCell(
        "42",
        text_alignment="right",
        vertical_alignment="center",
    )
    style = TableStyle(
        cell_text_alignment="center",
        cell_vertical_alignment="bottom",
        header_text_alignment="right",
        header_vertical_alignment="middle",
    )

    assert cell.text_alignment == "right"
    assert cell.vertical_alignment == "middle"
    assert style.cell_text_alignment == "center"
    assert style.cell_vertical_alignment == "bottom"
    assert style.header_text_alignment == "right"
    assert style.header_vertical_alignment == "middle"

    try:
        TableCell("bad", text_alignment="diagonal")
    except ValueError as exc:
        assert "alignment" in str(exc)
    else:
        raise AssertionError("Expected invalid horizontal alignment to fail")

    try:
        TableStyle(cell_vertical_alignment="baseline")
    except ValueError as exc:
        assert "vertical" in str(exc)
    else:
        raise AssertionError("Expected invalid vertical alignment to fail")


def test_table_column_specs_render_flex_widths_and_wrapping(
    tmp_path: Path,
) -> None:
    table = Table(
        headers=["Case", "Status", "Notes"],
        rows=[["case-001", "pass", "Fixed-page renderers share remaining text width."]],
        caption="Column spec table.",
        columns=[
            ColumnSpec(width=1.0, unit="in", text_alignment="right", wrap=False),
            ColumnSpec(flex=1, text_alignment="center"),
            ColumnSpec(flex=2, text_alignment="left"),
        ],
    )
    settings = DocumentSettings(
        unit="in",
        page_size=PageSize.letter(),
        page_margins=PageMargins.all(0.75, unit="in"),
    )
    document = Document("Column Specs", table, settings=settings)

    assert table._column_widths_in_inches("in", available_width=7.0) == [
        1.0,
        2.0,
        4.0,
    ]

    with pytest.raises(ValueError, match="mutually exclusive"):
        Table(
            headers=["A"],
            rows=[["B"]],
            column_widths=[1.0],
            columns=[ColumnSpec(flex=1)],
        )

    docx_path = tmp_path / "column-specs.docx"
    pdf_path = tmp_path / "column-specs.pdf"
    html_path = tmp_path / "column-specs.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf_path).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert 'w:w="1440"' in docx_xml
    assert 'w:w="2880"' in docx_xml
    assert 'w:w="5760"' in docx_xml
    assert '<w:jc w:val="right"' in docx_xml
    assert "Column spec table." in pdf_text
    assert "case-001" in pdf_text
    assert 'style="width: 1.00in;"' in html_text
    assert 'style="width: 2.00in;"' in html_text
    assert 'style="width: 4.00in;"' in html_text
    assert "white-space: nowrap" in html_text
    assert "text-align: right" in html_text
    assert "text-align: center" in html_text


def test_table_split_and_media_placement_options_render(tmp_path: Path) -> None:
    image_path = tmp_path / "placement.png"
    _write_sample_image(image_path)
    long_rows = [[f"Item {index}", f"Value {index}"] for index in range(34)]
    long_table = Table(
        headers=["Item", "Value"],
        rows=long_rows,
        caption="Long table with repeated headers.",
        column_widths=[2.0, 2.0],
        split=False,
    )
    here_table = Table(
        headers=["Mode", "Behavior"],
        rows=[["split=True", "render here and allow page breaks"]],
        caption="Explicit split table.",
        split=True,
        continuation_label="continued",
        continued_caption_template="{caption} -- {continuation_label}",
    )
    invalid_template_error = None
    try:
        Table(
            headers=["Mode"],
            rows=[["bad"]],
            continued_caption_template="{missing}",
        )
    except ValueError as exc:
        invalid_template_error = exc
    top_figure = Figure(
        image_path,
        caption=Paragraph("Figure with top placement."),
        width=1.0,
        placement="top",
    )
    document = Document(
        "Placement Test",
        Paragraph("Before media."),
        long_table,
        here_table,
        top_figure,
        settings=DocumentSettings(
            page_size=PageSize.letter(),
            page_margins=PageMargins.all(0.5, unit="in"),
        ),
    )

    assert long_table._resolve_split() is True
    assert long_table._resolve_placement() == "here"
    assert here_table._resolve_split() is True
    assert here_table._resolve_placement() == "here"
    assert here_table.continued_caption_text() == "Explicit split table. -- continued"
    assert invalid_template_error is not None
    assert "continued_caption_template" in str(invalid_template_error)
    assert top_figure.resolved_placement() == "top"

    docx_path = tmp_path / "placement.docx"
    pdf_path = tmp_path / "placement.pdf"
    html_path = tmp_path / "placement.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert long_table.total_row_count == 35
    assert 'w:tblHeader' in docx_xml
    assert docx_xml.count('<w:tblHeader') >= 2
    assert '<w:br w:type="page"/>' in docx_xml
    assert "Long table with repeated headers." in pdf_text
    assert pdf_text.count("Item") >= 2
    assert len(pdf_reader.pages) >= 2
    assert 'oodocs-table-split' in html_text
    assert 'data-continuation-label="continued"' in html_text
    assert 'data-continued-caption="Explicit split table. -- continued"' in html_text
    assert '.oodocs-table-split thead' in html_text
    assert 'display: table-header-group' in html_text
    assert 'oodocs-placement-here' in html_text
    assert 'oodocs-placement-top' in html_text
    assert 'break-before: page' in html_text


def test_pdf_float_tables_can_move_after_following_prose(tmp_path: Path) -> None:
    floating_table = Table(
        headers=["Metric", "Value"],
        rows=[["Latency", "14 ms"], ["Quality", "stable"]],
        caption="Deferred table.",
        column_widths=[1.4, 1.4],
    )
    here_table = Table(
        headers=["Metric", "Value"],
        rows=[["Throughput", "ready"]],
        caption="Here table.",
        placement="here",
        column_widths=[1.4, 1.4],
    )
    document = Document(
        "Float Test",
        Chapter(
            "Main",
            Section(
                "Flow",
                Paragraph("Before floating table."),
                floating_table,
                Paragraph("Following prose fills the available page before the float."),
                Paragraph("Before here table."),
                here_table,
                Paragraph("After here table."),
            ),
        ),
    )

    pdf_path = tmp_path / "float.pdf"
    document.save_pdf(pdf_path)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert pdf_text.index("Following prose fills the available page before the float.") < pdf_text.index("Table 1. Deferred table.")
    assert pdf_text.index("Table 2. Here table.") < pdf_text.index("After here table.")


def test_heading_hierarchy_uses_latex_like_levels() -> None:
    chapter = Chapter(
        "Part I",
        Section(
            "Overview",
            Subsection(
                "Details",
                SubSubsection("Examples"),
            ),
        ),
    )

    assert chapter.level == 1
    assert chapter.children[0].level == 2
    assert chapter.children[0].children[0].level == 3
    assert chapter.children[0].children[0].children[0].level == 4


def test_sections_can_be_rendered_without_numbering() -> None:
    unnumbered = Section("Abstract", Paragraph("Summary."), level=2, numbered=False)
    numbered = Section("Introduction", Paragraph("Body."), level=1)
    document = Document("Article", unnumbered, numbered)

    render_index = build_render_index(document)

    assert render_index.heading_number(unnumbered) is None
    assert render_index.heading_number(numbered) == "1"


def test_parts_use_dedicated_pages_without_resetting_chapters(tmp_path: Path) -> None:
    first_part = Part(
        "Foundations",
        Chapter("Start", Section("Scope", Paragraph("Body one."))),
    )
    second_part = Part(
        "Reference",
        Chapter("Continue", Section("API", Paragraph("Body two."))),
    )
    document = Document("Part Test", TableOfContents(), first_part, second_part)

    render_index = build_render_index(document)

    assert render_index.heading_number(first_part) == "Part I"
    assert render_index.heading_number(first_part.children[0]) == "1"
    assert render_index.heading_number(second_part) == "Part II"
    assert render_index.heading_number(second_part.children[0]) == "2"
    assert [entry.level for entry in render_index.headings[:4]] == [0, 1, 2, 0]

    docx_path = tmp_path / "parts.docx"
    pdf_path = tmp_path / "parts.pdf"
    html_path = tmp_path / "parts.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    assert "Part I" in paragraph_texts
    assert "Foundations" in paragraph_texts
    assert "1 Start" in paragraph_texts
    assert "Part II" in paragraph_texts
    assert "2 Continue" in paragraph_texts
    assert _docx_document_xml(docx_path).count('<w:br w:type="page"/>') >= 2

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Part I" in pdf_text
    assert "1 Start" in pdf_text
    assert "Part II" in pdf_text
    assert "2 Continue" in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "Part I Foundations" in normalized_html_text
    assert "2 Continue" in normalized_html_text
    assert 'class="oodocs-part-page oodocs-page-break-before oodocs-page-break-after"' in html_text
    assert 'oodocs-toc-entry-level-0' in html_text


def test_appendix_switches_child_heading_numbers_to_letters(tmp_path: Path) -> None:
    appendix = Appendix(
        Chapter(
            "Input Data Schema",
            Section("Fields", Paragraph("Field definitions.")),
        ),
        Chapter("Validation Cases", Paragraph("Reference checks.")),
    )
    document = Document(
        "Appendix Test",
        TableOfContents(show_page_numbers=False),
        Chapter("Main", Paragraph("Body.")),
        Paragraph("See ", reference(appendix.children[0]), "."),
        appendix,
    )

    render_index = build_render_index(document)

    assert render_index.heading_number(document.body.children[1]) == "1"
    assert render_index.heading_number(appendix) is None
    assert render_index.heading_number(appendix.children[0]) == "A"
    assert render_index.heading_number(appendix.children[0].children[0]) == "A.1"
    assert render_index.heading_number(appendix.children[1]) == "B"
    assert [
        ("".join(fragment.plain_text() for fragment in entry.title), entry.number)
        for entry in render_index.headings
    ] == [
        ("Main", "1"),
        ("Appendices", None),
        ("Input Data Schema", "A"),
        ("Fields", "A.1"),
        ("Validation Cases", "B"),
    ]

    docx_path = tmp_path / "appendix.docx"
    pdf_path = tmp_path / "appendix.pdf"
    html_path = tmp_path / "appendix.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    paragraph_texts = [paragraph.text for paragraph in WordDocument(docx_path).paragraphs]
    assert "Appendices" in paragraph_texts
    assert "See Chapter A." in paragraph_texts
    assert "A Input Data Schema" in paragraph_texts
    assert "A.1 Fields" in paragraph_texts
    assert "B Validation Cases" in paragraph_texts

    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages
    )
    assert "Appendices" in pdf_text
    assert "See Chapter A." in pdf_text
    assert "A Input Data Schema" in pdf_text
    assert "A.1 Fields" in pdf_text
    assert "B Validation Cases" in pdf_text

    normalized_html_text = _normalized_html_text(html_path)
    assert "Appendices" in normalized_html_text
    assert "See Chapter A" in normalized_html_text
    assert "A Input Data Schema" in normalized_html_text
    assert "A.1 Fields" in normalized_html_text
    assert "B Validation Cases" in normalized_html_text


def test_appendix_validation_warns_when_nested() -> None:
    document = Document("Nested Appendix", Chapter("Main", Appendix(Chapter("Data"))))

    warning_codes = {issue.code for issue in document.validate().warnings}

    assert "nested-appendix" in warning_codes


def test_public_api_prefers_classes_for_structural_nodes() -> None:
    assert hasattr(oodocs, "Document")
    assert hasattr(oodocs, "DocumentSettings")
    assert hasattr(oodocs, "Chapter")
    assert hasattr(oodocs, "AuthorLayout")
    assert hasattr(oodocs, "Section")
    assert hasattr(oodocs, "Shape")
    assert hasattr(oodocs, "Paragraph")
    assert hasattr(oodocs, "Part")
    assert hasattr(oodocs, "Appendix")
    assert hasattr(oodocs, "BulletList")
    assert hasattr(oodocs, "ColumnSpan")
    assert hasattr(oodocs, "CountableBlock")
    assert hasattr(oodocs, "Definition")
    assert hasattr(oodocs, "Lemma")
    assert hasattr(oodocs, "Proposition")
    assert hasattr(oodocs, "Theorem")
    assert hasattr(oodocs, "Corollary")
    assert hasattr(oodocs, "Proof")
    assert hasattr(oodocs, "Example")
    assert hasattr(oodocs, "Remark")
    assert hasattr(oodocs, "Assumption")
    assert hasattr(oodocs, "Axiom")
    assert hasattr(oodocs, "Claim")
    assert hasattr(oodocs, "Conjecture")
    assert hasattr(oodocs, "create_countable_block_type")
    assert hasattr(oodocs, "MultiColumn")
    assert hasattr(oodocs, "NumberedList")
    assert hasattr(oodocs, "PageBreak")
    assert hasattr(oodocs, "PageMargins")
    assert hasattr(oodocs, "PageSize")
    assert hasattr(oodocs, "ListOfTables")
    assert hasattr(oodocs, "ListOfFigures")
    assert hasattr(oodocs, "cite")
    assert hasattr(oodocs, "Box")
    assert hasattr(oodocs, "BoxStyle")
    assert hasattr(oodocs, "CitationDefaults")
    assert hasattr(oodocs, "HeadingStyle")
    assert hasattr(oodocs, "HeadingNumbering")
    assert hasattr(oodocs, "ImageBox")
    assert hasattr(oodocs, "ImageData")
    assert hasattr(oodocs, "ListStyle")
    assert hasattr(oodocs, "Table")
    assert hasattr(oodocs, "TableCell")
    assert hasattr(oodocs, "TableStyle")
    assert hasattr(oodocs, "Figure")
    assert hasattr(oodocs, "TableOfContents")
    assert hasattr(oodocs, "TocLevelStyle")
    assert hasattr(oodocs, "Comment")
    assert hasattr(oodocs, "CommentList")
    assert not hasattr(oodocs, "Todo")
    assert not hasattr(oodocs, "MarginNote")
    assert hasattr(review_components, "Todo")
    assert hasattr(review_components, "MarginNote")
    assert hasattr(oodocs, "Footnote")
    assert hasattr(oodocs, "FootnoteList")
    assert hasattr(oodocs, "Equation")
    assert hasattr(oodocs, "ChemicalFormula")
    assert hasattr(oodocs, "ReactionEquation")
    assert hasattr(oodocs, "Math")
    assert hasattr(oodocs, "InlineChip")
    assert hasattr(oodocs, "InlineChipStyle")
    assert hasattr(oodocs, "TextStyle")
    assert hasattr(oodocs, "TextBox")
    assert hasattr(oodocs, "LineBreak")
    assert hasattr(oodocs, "VerticalSpace")
    assert hasattr(oodocs, "Divider")
    assert not hasattr(oodocs, "Sheet")
    assert not hasattr(oodocs, "Body")
    assert not hasattr(oodocs, "Bold")
    assert not hasattr(oodocs, "Hyperlink")
    assert not hasattr(oodocs, "Italic")
    assert not hasattr(oodocs, "InlineCode")
    assert hasattr(inline_components, "Bold")
    assert hasattr(inline_components, "Hyperlink")
    assert hasattr(inline_components, "Italic")
    assert hasattr(inline_components, "InlineCode")
    assert hasattr(generated_components, "FootnoteList")
    assert not hasattr(oodocs, "ListBlock")
    assert not hasattr(oodocs, "Citation")
    assert not hasattr(oodocs, "TableReference")
    assert not hasattr(oodocs, "FigureReference")
    assert not hasattr(oodocs, "Strong")
    assert not hasattr(oodocs, "Emphasis")
    assert not hasattr(oodocs, "Code")
    assert not hasattr(oodocs, "PageBreaker")
    assert not hasattr(oodocs, "VSpace")
    assert not hasattr(oodocs, "HorizontalRule")
    assert hasattr(oodocs, "comment")
    assert not hasattr(oodocs, "todo")
    assert not hasattr(oodocs, "margin_note")
    assert hasattr(review_components, "todo")
    assert hasattr(review_components, "margin_note")
    assert hasattr(oodocs, "footnote")
    assert hasattr(oodocs, "from_notebook")
    assert hasattr(oodocs, "from_markdown")
    assert hasattr(oodocs, "from_markdown_file")
    assert not hasattr(oodocs, "from_ipynb")
    assert hasattr(oodocs, "math")
    assert hasattr(oodocs, "chemical_formula")
    assert hasattr(oodocs, "prescript")
    assert hasattr(oodocs, "reference")
    assert hasattr(oodocs, "bold")
    assert hasattr(oodocs, "italic")
    assert hasattr(oodocs, "inline_code")
    assert hasattr(oodocs, "text_color")
    assert not hasattr(oodocs, "code")
    assert not hasattr(oodocs, "color")
    assert hasattr(oodocs, "highlight")
    assert hasattr(oodocs, "link")
    assert hasattr(oodocs, "line_break")
    assert not hasattr(oodocs, "vspace")
    assert not hasattr(oodocs, "hrule")
    assert not hasattr(oodocs, "strike")
    assert hasattr(oodocs, "strikethrough")
    assert hasattr(oodocs, "tag")
    assert hasattr(oodocs, "badge")
    assert hasattr(oodocs, "status")
    assert hasattr(oodocs, "keyboard")
    assert hasattr(oodocs, "subscript")
    assert hasattr(oodocs, "superscript")
    assert hasattr(oodocs, "parse_notebook")
    assert hasattr(oodocs, "parse_markdown")
    assert hasattr(oodocs, "parse_markdown_file")
    assert not hasattr(oodocs, "parse_ipynb")
    assert hasattr(inline_components, "InlineChip")
    assert hasattr(inline_components, "InlineChipStyle")
    assert not hasattr(inline_components, "Strong")
    assert not hasattr(inline_components, "Emphasis")
    assert not hasattr(inline_components, "Code")
    assert hasattr(inline_components, "tag")
    assert hasattr(inline_components, "badge")
    assert hasattr(inline_components, "status")
    assert hasattr(inline_components, "keyboard")
    assert hasattr(inline_components, "reference")
    assert ImageData(_build_sample_png()).data

    for removed_name in (
        "document",
        "body",
        "chapter",
        "section",
        "subsection",
        "subsubsection",
        "paragraph",
        "code_block",
        "bullet_list",
        "numbered_list",
        "table",
        "figure",
    ):
        assert not hasattr(oodocs, removed_name)


def test_multicolumn_layout_renders_across_outputs(tmp_path: Path) -> None:
    image_path = tmp_path / "multicolumn.png"
    _write_sample_image(image_path)

    wide_table = Table(
        headers=["Metric", "Value"],
        rows=[
            ["Traceability checks", "passed"],
            ["Late revision cost", "reduced"],
        ],
        caption="Wide result table.",
        column_widths=[2.4, 2.4],
        placement="here",
    )
    wide_figure = Figure(
        image_path,
        caption="Wide result figure.",
        width=5.6,
        placement="here",
    )
    document = Document(
        "Multicolumn Report",
        Section(
            "Findings",
            MultiColumn(
                Paragraph("Column paragraph alpha."),
                Paragraph("Column paragraph beta."),
                wide_table,
                Paragraph("Column paragraph gamma."),
                ColumnSpan(wide_figure),
                Paragraph("Column paragraph delta."),
                columns=2,
                column_gap=0.18,
            ),
        ),
    )

    docx_path = tmp_path / "multicolumn.docx"
    pdf_path = tmp_path / "multicolumn.pdf"
    html_path = tmp_path / "multicolumn.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    word_text = "\n".join(
        [paragraph.text for paragraph in word_document.paragraphs]
        + [
            cell.text
            for table in word_document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "Column paragraph alpha." in word_text
    assert "Column paragraph delta." in word_text
    assert "Table 1. Wide result table." in word_text
    assert "Figure 1. Wide result figure." in word_text
    docx_xml = _docx_document_xml(docx_path)
    assert "<w:cols" in docx_xml
    assert 'w:num="2"' in docx_xml

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Column paragraph alpha." in pdf_text
    assert "Column paragraph delta." in pdf_text
    assert "Table 1. Wide result table." in pdf_text
    assert "Figure 1. Wide result figure." in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert 'class="oodocs-multi-column-layout"' in html_text
    assert "column-count: 2" in html_text
    assert html_text.count('class="oodocs-column-span"') >= 2
    assert "Column paragraph alpha." in normalized_html_text
    assert "Column paragraph delta." in normalized_html_text
    assert "Table 1. Wide result table." in normalized_html_text
    assert "Figure 1. Wide result figure." in normalized_html_text


def test_page_items_render_without_affecting_document_flow(tmp_path: Path) -> None:
    image_path = tmp_path / "page-logo.png"
    _write_sample_image(image_path)
    page_items = [
        Shape.rect(
            name="frame",
            x=0.25,
            y=0.25,
            width=7.0,
            height=10.0,
            stroke=StrokeStyle.solid("#476172", width=1.4),
        ),
        Shape.ellipse(
            anchor="frame",
            x=5.7,
            y=0.25,
            width=0.7,
            height=0.7,
            stroke=StrokeStyle.solid("#B2783D"),
            fill_color="#FFF1D8",
        ),
        ImageBox(
            image_path,
            anchor="frame",
            x=0.4,
            y=0.35,
            width=0.7,
            height=0.42,
            z_index=1,
        ),
        TextBox(
            "Positioned Page Overlay",
            anchor="margin",
            x=0.0,
            y=0.0,
            width=3.0,
            height=0.45,
            font_size=12,
        ),
        TextBox(
            "Anchored to the named frame shape.",
            anchor="frame",
            x=0.7,
            y=0.95,
            width=5.4,
            height=0.5,
            text_alignment="center",
            vertical_alignment="middle",
            font_size=11,
            z_index=2,
        ),
    ]
    document = Document(
        "Page Item Test",
        Paragraph("Body text keeps its normal position."),
        settings=DocumentSettings(
            page_size=PageSize.letter(),
            page_items=page_items,
            theme=Theme(page_numbers=PageNumberDefaults(show_page_numbers=True)),
        ),
    )

    docx_path = tmp_path / "page-items.docx"
    pdf_path = tmp_path / "page-items.pdf"
    html_path = tmp_path / "page-items.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    word_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    html_text = html_path.read_text(encoding="utf-8")
    word_xml = _docx_word_xml(docx_path)

    assert "Body text keeps its normal position." in word_text
    assert "Positioned Page Overlay" in word_xml
    assert "Anchored to the named frame shape." in word_xml
    assert "<v:rect" in word_xml
    assert "<v:oval" in word_xml
    assert "<v:imagedata" in word_xml
    assert "Positioned Page Overlay" in pdf_text
    assert "Anchored to the named frame shape." in pdf_text
    assert _pdf_image_draw_count(pdf_path) == 1
    assert 'class="oodocs-page-items"' in html_text
    assert 'class="oodocs-page-item oodocs-imagebox"' in html_text
    assert "Positioned Page Overlay" in html_text
    assert "Anchored to the named frame shape." in html_text
    assert html_text.count("data:image/png;base64,") == 1
    assert "oodocs-sheet" not in html_text


def test_page_item_scopes_filter_pdf_and_warn_for_static_outputs(tmp_path: Path) -> None:
    document = Document(
        "Scoped Overlay Test",
        Chapter("Main Matter", Paragraph("Main body text.")),
        settings=DocumentSettings(
            page_size=PageSize.letter(),
            cover_page=True,
            page_items=[
                TextBox("ALL SCOPE", x=0.3, y=0.2, width=1.8, height=0.25, font_size=8),
                TextBox(
                    "COVER SCOPE",
                    x=0.3,
                    y=0.5,
                    width=1.8,
                    height=0.25,
                    font_size=8,
                    scope="cover",
                ),
                TextBox(
                    "MAIN SCOPE",
                    x=0.3,
                    y=0.8,
                    width=1.8,
                    height=0.25,
                    font_size=8,
                    scope=PageItemScope.main(),
                ),
                TextBox(
                    "PAGE TWO SCOPE",
                    x=0.3,
                    y=1.1,
                    width=2.1,
                    height=0.25,
                    font_size=8,
                    scope=PageItemScope.pages(2),
                ),
                TextBox(
                    "FRONT SCOPE",
                    x=0.3,
                    y=1.4,
                    width=1.8,
                    height=0.25,
                    font_size=8,
                    scope="front",
                ),
            ],
        ),
    )

    validation = document.validate()
    assert "page-item-scope-static-output" in {issue.code for issue in validation.warnings}
    assert validation.warnings_for(("pdf",)) == ()

    docx_path = tmp_path / "scoped-page-items.docx"
    pdf_path = tmp_path / "scoped-page-items.pdf"
    html_path = tmp_path / "scoped-page-items.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    page_texts = [
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages
    ]
    assert len(page_texts) >= 2
    assert "ALL SCOPE" in page_texts[0]
    assert "COVER SCOPE" in page_texts[0]
    assert "MAIN SCOPE" not in page_texts[0]
    assert "PAGE TWO SCOPE" not in page_texts[0]
    assert "FRONT SCOPE" not in "\n".join(page_texts)
    assert "ALL SCOPE" in page_texts[1]
    assert "MAIN SCOPE" in page_texts[1]
    assert "PAGE TWO SCOPE" in page_texts[1]
    assert "COVER SCOPE" not in page_texts[1]

    word_xml = _docx_word_xml(docx_path)
    assert "COVER SCOPE" in word_xml
    assert "MAIN SCOPE" in word_xml

    html_text = html_path.read_text(encoding="utf-8")
    assert "ALL SCOPE" in html_text
    assert "COVER SCOPE" in html_text
    assert "MAIN SCOPE" not in html_text


def test_header_footer_templates_render_across_outputs(tmp_path: Path) -> None:
    document = Document(
        "Header Footer",
        Chapter(
            "Overview",
            Section("Scope", Paragraph("Body text.")),
        ),
        settings=DocumentSettings(
            theme=Theme(
                header_footer=HeaderFooterDefaults(
                    header_left="{chapter}",
                    header_right="{page}",
                    footer_center="{title}",
                    first_header_center="First {title}",
                    first_footer_center="Cover {page}",
                    even_header_left="Even {section}",
                    different_first_page=True,
                    different_odd_even_pages=True,
                )
            )
        ),
    )

    docx_path = tmp_path / "header-footer.docx"
    pdf_path = tmp_path / "header-footer.pdf"
    html_path = tmp_path / "header-footer.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    settings_xml = word_document.settings.element.xml
    header_xml = word_document.sections[0].header.paragraphs[0]._p.xml
    first_header_xml = word_document.sections[0].first_page_header.paragraphs[0]._p.xml
    first_footer_xml = word_document.sections[0].first_page_footer.paragraphs[0]._p.xml
    even_header_xml = word_document.sections[0].even_page_header.paragraphs[0]._p.xml

    assert "evenAndOddHeaders" in settings_xml
    assert 'STYLEREF "Heading 1"' in header_xml
    assert " PAGE " in header_xml
    assert "First " in first_header_xml
    assert "Header Footer" in first_header_xml
    assert "Cover " in first_footer_xml
    assert " PAGE " in first_footer_xml
    assert 'STYLEREF "Heading 2"' in even_header_xml

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "First Header Footer" in pdf_text
    assert "Cover 1" in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    assert 'class="oodocs-header-footer"' in html_text
    assert "First Header Footer" in html_text
    assert "Cover 1" in html_text
    assert "@media print" in html_text


def test_positioned_items_can_render_inline_like_text(tmp_path: Path) -> None:
    image_path = tmp_path / "inline-logo.png"
    _write_sample_image(image_path)
    document = Document(
        "Inline Drawing Test",
        Paragraph("Before inline drawing."),
        Shape.rect(
            width=1.2,
            height=0.35,
            placement="inline",
            stroke=StrokeStyle.solid("#476172"),
            fill_color="#EEF6FF",
        ),
        ImageBox(
            image_path,
            width=0.7,
            height=0.42,
            placement="inline",
        ),
        TextBox(
            "Inline textbox",
            width=1.8,
            height=0.35,
            placement="inline",
            text_alignment="center",
        ),
        Paragraph(
            "Image can sit ",
            ImageBox(
                image_path,
                width=0.45,
                height=0.27,
                placement="inline",
            ),
            " inside text.",
        ),
        Paragraph("After inline drawing."),
        settings=DocumentSettings(page_size=PageSize.letter()),
    )

    docx_path = tmp_path / "inline-drawing.docx"
    pdf_path = tmp_path / "inline-drawing.pdf"
    html_path = tmp_path / "inline-drawing.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    word_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
    word_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert "Before inline drawing." in word_text
    assert "After inline drawing." in word_text
    assert "<v:rect" in word_xml
    assert "Inline textbox" in word_xml
    assert len(word_document.inline_shapes) == 2
    assert "Inline textbox" in pdf_text
    assert _pdf_image_draw_count(pdf_path) == 2
    assert "display: inline-block" in html_text
    assert "Inline textbox" in html_text
    assert html_text.count("data:image/png;base64,") == 2


def test_box_style_supports_tcolorbox_like_layout_controls(tmp_path: Path) -> None:
    image_path = tmp_path / "panel-logo.png"
    _write_sample_image(image_path)
    document = Document(
        "Box Layout Test",
        Box(
            Paragraph("Editable content inside a styled report panel."),
            Table(
                headers=["Surface", "Behavior"],
                rows=[["Box", "compact nested table"]],
                column_widths=[1.2, 2.0],
            ),
            Figure(image_path, width=0.7),
            title="Panel",
            style=BoxStyle(
                border=BorderStyle.solid("#1058A3", width=0.5),
                background_color="#FFFFFF",
                title_background_color="#1058A3",
                title_text_color="#FFFFFF",
                padding=Padding(2, 5, 3, 7),
                width=10,
                unit="cm",
                block_alignment="left",
            ),
        ),
        CalloutBox(
            Paragraph("Stop the release until the reviewer resolves this item."),
            variant="danger",
            icon="!",
            title_position="side",
            shadow=True,
        ),
    )
    validation = document.validate()
    assert "box-shadow-html-only" in {issue.code for issue in validation.warnings}

    docx_path = tmp_path / "box-layout.docx"
    pdf_path = tmp_path / "box-layout.pdf"
    html_path = tmp_path / "box-layout.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    html_text = html_path.read_text(encoding="utf-8")
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_pdf_text = " ".join(pdf_text.split())

    assert "Editable content inside a styled report panel." in docx_xml
    assert 'w:fill="1058A3"' in docx_xml
    assert 'w:color="FFFFFF"' in docx_xml
    assert '<w:tblW w:w="5669" w:type="dxa"/>' in docx_xml
    assert '<w:top w:w="40" w:type="dxa"/>' in docx_xml
    assert '<w:left w:w="140" w:type="dxa"/>' in docx_xml
    assert "width: 3.9370in" in html_text
    assert "padding: 2.0pt 5.0pt 3.0pt 7.0pt" in html_text
    assert "padding: 0;" in html_text
    assert "box-shadow: none" in html_text
    assert "color: #FFFFFF" in html_text
    assert "grid-template-columns: max-content minmax(0, 1fr)" in html_text
    assert "box-shadow: 0 10pt 22pt rgba(15, 23, 42, 0.14)" in html_text
    assert "Danger" in docx_xml
    assert "Stop the release until the reviewer resolves this item." in docx_xml
    assert "Editable content inside a styled report panel." in pdf_text
    assert "compact nested table" in pdf_text
    assert "Danger" in pdf_text
    assert "Stop the release until the reviewer resolves this item." in normalized_pdf_text


def test_explicit_page_break_renders_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Break Test",
        Paragraph("Before break."),
        PageBreak(),
        Paragraph("After break."),
    )

    docx_path = tmp_path / "break.docx"
    pdf_path = tmp_path / "break.pdf"
    html_path = tmp_path / "break.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    assert '<w:br w:type="page"/>' in _docx_document_xml(docx_path)
    assert len(PdfReader(BytesIO(pdf_path.read_bytes())).pages) >= 2
    assert 'class="oodocs-page-break"' in html_path.read_text(encoding="utf-8")


def test_inline_highlight_strike_and_line_break_render_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Inline Word Features",
        Paragraph(
            "Keep ",
            highlight("review focus", "#FFF2CC"),
            ", remove ",
            strikethrough("old value"),
            line_break(),
            "Continue same paragraph.",
        ),
    )

    docx_path = tmp_path / "inline-word-features.docx"
    pdf_path = tmp_path / "inline-word-features.pdf"
    html_path = tmp_path / "inline-word-features.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert 'w:fill="FFF2CC"' in docx_xml
    assert "<w:strike" in docx_xml
    assert "<w:br" in docx_xml
    assert "review focus" in pdf_text
    assert "old value" in pdf_text
    assert "Continue same paragraph." in pdf_text
    assert "background-color: #FFF2CC" in html_text
    assert "line-through" in html_text
    assert "<br/>" in html_text


def test_inline_caps_and_vertical_align_render_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Inline Detail Features",
        Paragraph(
            Text.styled("small caps", small_caps=True),
            " ",
            Text.styled("upper", uppercase=True),
            " H",
            Text.styled("2", subscript=True),
            " x",
            Text.styled("2", superscript=True),
            " isotope ",
            prescript("14", "6", "C"),
            " index ",
            subscript("i"),
            superscript("j"),
        ),
    )

    docx_path = tmp_path / "inline-detail-features.docx"
    pdf_path = tmp_path / "inline-detail-features.pdf"
    html_path = tmp_path / "inline-detail-features.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert "<w:smallCaps" in docx_xml
    assert "<w:caps" in docx_xml
    assert 'w:val="subscript"' in docx_xml
    assert 'w:val="superscript"' in docx_xml
    word_paragraph = next(paragraph for paragraph in WordDocument(docx_path).paragraphs if "isotope 146C" in paragraph.text)
    assert any(run.text == "14" and run.font.superscript for run in word_paragraph.runs)
    assert any(run.text == "6" and run.font.subscript for run in word_paragraph.runs)
    assert any(run.text == "i" and run.font.subscript for run in word_paragraph.runs)
    assert any(run.text == "j" and run.font.superscript for run in word_paragraph.runs)
    assert "UPPER" in pdf_text
    assert "146C" in pdf_text
    assert "font-variant: small-caps" in html_text
    assert "text-transform: uppercase" in html_text
    assert "vertical-align: sub" in html_text
    assert "vertical-align: super" in html_text


def test_inline_chips_render_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Inline Chips",
        Paragraph(
            "Route ",
            tag("api"),
            " with ",
            status("ready", state="success"),
            ", show ",
            badge(3),
            ", and press ",
            keyboard("Ctrl+Enter"),
            ".",
        ),
    )

    docx_path = tmp_path / "inline-chips.docx"
    pdf_path = tmp_path / "inline-chips.pdf"
    html_path = tmp_path / "inline-chips.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html = _normalized_html_text(html_path)

    assert len(word_document.inline_shapes) == 4
    assert "Route  with , show , and press ." in [paragraph.text for paragraph in word_document.paragraphs]
    assert "api" in pdf_text
    assert "READY" in pdf_text
    assert "3" in pdf_text
    assert "Ctrl+Enter" in pdf_text
    assert "oodocs-inline-chip-tag" in html_text
    assert "oodocs-inline-chip-badge" in html_text
    assert "oodocs-inline-chip-status" in html_text
    assert "oodocs-inline-chip-keyboard" in html_text
    assert "api" in normalized_html
    assert "READY" in normalized_html
    assert "Ctrl+Enter" in normalized_html


def test_paragraph_indents_render_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Indent Test",
        Paragraph(
            "Indented paragraph.",
            style=ParagraphStyle(
                left_indent=1.27,
                right_indent=0.508,
                first_line_indent=0.635,
            ),
        ),
        Paragraph(
            "Hanging indent paragraph.",
            style=ParagraphStyle.hanging(left=1.524, by=0.635),
        ),
        settings=DocumentSettings(unit="cm"),
    )

    docx_path = tmp_path / "indent.docx"
    pdf_path = tmp_path / "indent.pdf"
    html_path = tmp_path / "indent.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert 'w:left="720"' in docx_xml
    assert 'w:right="288"' in docx_xml
    assert 'w:firstLine="360"' in docx_xml
    assert 'w:hanging="360"' in docx_xml
    assert "Indented paragraph." in pdf_text
    assert "Hanging indent paragraph." in pdf_text
    assert "margin-left: 0.50in" in html_text
    assert "margin-right: 0.20in" in html_text
    assert "text-indent: 0.25in" in html_text
    assert "text-indent: -0.25in" in html_text


def test_paragraph_spacing_and_pagination_options_render_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Paragraph Detail Options",
        Paragraph(
            "Keep this with the next block.",
            space_before=6,
            space_after=3,
            keep_together=True,
            keep_with_next=True,
            page_break_before=True,
            widow_control=False,
        ),
        Paragraph("Next block."),
    )

    docx_path = tmp_path / "paragraph-details.docx"
    pdf_path = tmp_path / "paragraph-details.pdf"
    html_path = tmp_path / "paragraph-details.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert 'w:before="120"' in docx_xml
    assert 'w:after="60"' in docx_xml
    assert "<w:keepLines" in docx_xml
    assert "<w:keepNext" in docx_xml
    assert "<w:pageBreakBefore" in docx_xml
    assert 'w:widowControl w:val="0"' in docx_xml
    assert "Keep this with the next block." in pdf_text
    assert "margin: 6.0pt 0 3.0pt" in html_text
    assert "break-inside: avoid" in html_text
    assert "page-break-before: always" in html_text
    assert "widows: 1" in html_text


def test_vertical_space_and_divider_render_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Spacing Blocks",
        Paragraph("Before spacer."),
        VerticalSpace(18),
        Divider(
            color="C8CDD6",
            thickness=1.5,
            space_before=4,
            space_after=5,
            width=2.0,
            alignment="center",
            unit="in",
        ),
        Paragraph("After divider."),
    )

    docx_path = tmp_path / "spacing-blocks.docx"
    pdf_path = tmp_path / "spacing-blocks.pdf"
    html_path = tmp_path / "spacing-blocks.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert 'w:after="360"' in docx_xml
    assert 'w:before="80"' in docx_xml
    assert 'w:after="100"' in docx_xml
    assert '<w:bottom w:val="single" w:sz="12" w:space="1" w:color="C8CDD6"/>' in docx_xml
    assert "Before spacer." in pdf_text
    assert "After divider." in pdf_text
    assert 'class="oodocs-vertical-space"' in html_text
    assert "height: 18.0pt" in html_text
    assert 'class="oodocs-divider"' in html_text
    assert "border-top: 1.50pt solid #C8CDD6" in html_text
    assert "width: 2.0000in" in html_text


def test_table_cell_alignment_renders_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Table Alignment Test",
        Table(
            headers=[
                [
                    TableCell(
                        "Metric",
                        text_alignment="center",
                        vertical_alignment="middle",
                    ),
                    "Value",
                ]
            ],
            rows=[
                [
                    "Latency",
                    TableCell(
                        "14 ms",
                        text_alignment="right",
                        vertical_alignment="bottom",
                    ),
                ],
            ],
            style=TableStyle(
                header_text_alignment="center",
                header_vertical_alignment="middle",
                cell_text_alignment="left",
                cell_vertical_alignment="top",
            ),
        ),
    )

    docx_path = tmp_path / "table-alignment.docx"
    pdf_path = tmp_path / "table-alignment.pdf"
    html_path = tmp_path / "table-alignment.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert '<w:jc w:val="center"' in docx_xml
    assert '<w:jc w:val="right"' in docx_xml
    assert '<w:vAlign w:val="center"' in docx_xml
    assert '<w:vAlign w:val="bottom"' in docx_xml
    assert "Metric" in pdf_text
    assert "14 ms" in pdf_text
    assert "text-align: center" in html_text
    assert "text-align: right" in html_text


def test_default_table_cell_alignment_is_left_in_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Default Table Alignment Test",
        Table(
            headers=["Area", "Status"],
            rows=[["Release notes", "Compatibility notes wrap cleanly in narrow cells."]],
        ),
        settings=DocumentSettings(
            theme=Theme(blocks=BlockDefaults(paragraph_text_alignment="justify"))
        ),
    )

    docx_path = tmp_path / "default-table-alignment.docx"
    pdf_path = tmp_path / "default-table-alignment.pdf"
    html_path = tmp_path / "default-table-alignment.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert '<w:jc w:val="left"' in docx_xml
    assert "Compatibility notes wrap cleanly" in pdf_text
    assert "text-align: left" in html_text


def test_table_cell_row_and_column_styles_render_to_all_outputs(tmp_path: Path) -> None:
    table = Table(
        headers=["Metric", "Value"],
        rows=[
            [
                "Latency",
                TableCell(
                    "14 ms",
                    style=TableCellStyle(
                        background_color="#FFE699",
                        text_color="#7F1D1D",
                        bold=True,
                    ),
                ),
            ],
            ["Quality", "Stable"],
        ],
        caption="Styled table.",
        row_styles={
            1: TableCellStyle(background_color="#E2F0D9", italic=True),
        },
        column_styles={
            0: TableCellStyle(text_color="#1F4E79", bold=True),
        },
        header_row_styles={
            0: TableCellStyle(background_color="#1F4E79", text_color="#FFFFFF"),
        },
    )
    document = Document("Table Style Test", table)

    docx_path = tmp_path / "table-style.docx"
    pdf_path = tmp_path / "table-style.pdf"
    html_path = tmp_path / "table-style.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert "1F4E79" in docx_xml
    assert "FFE699" in docx_xml
    assert "E2F0D9" in docx_xml
    assert '<w:b' in docx_xml
    assert '<w:i' in docx_xml
    assert "Styled table." in pdf_text
    assert "Latency" in pdf_text
    assert "Stable" in pdf_text
    assert "background: #1F4E79" in html_text
    assert "background: #FFE699" in html_text
    assert "background: #E2F0D9" in html_text
    assert "color: #7F1D1D" in html_text
    assert "font-weight: 700" in html_text
    assert "font-style: italic" in html_text


def test_table_detail_style_options_render_to_all_outputs(tmp_path: Path) -> None:
    table = Table(
        headers=["Metric", "Value"],
        rows=[["Latency", "14 ms"], ["Quality", "Stable"]],
        caption="Detailed table.",
        border=BorderStyle.solid("#334155", width=1.25),
        cell_padding=Padding(2, 3, 4, 5),
        repeat_header_rows=True,
    )
    document = Document("Table Detail Test", table)

    docx_path = tmp_path / "table-detail-style.docx"
    pdf_path = tmp_path / "table-detail-style.pdf"
    html_path = tmp_path / "table-detail-style.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert 'w:sz="10"' in docx_xml
    assert 'w:color="334155"' in docx_xml
    assert '<w:top w:w="40" w:type="dxa"/>' in docx_xml
    assert '<w:right w:w="60" w:type="dxa"/>' in docx_xml
    assert '<w:bottom w:w="80" w:type="dxa"/>' in docx_xml
    assert '<w:left w:w="100" w:type="dxa"/>' in docx_xml
    assert "<w:tblHeader" in docx_xml
    assert "Detailed table." in pdf_text
    assert "border: 1.25pt solid #334155" in html_text
    assert "padding: 2.0pt 3.0pt 4.0pt 5.0pt" in html_text


def test_booktabs_table_style_renders_horizontal_rules_to_all_outputs(tmp_path: Path) -> None:
    table = Table(
        headers=["Metric", "Value"],
        rows=[["Accuracy", "0.91"], ["F1", "0.88"]],
        caption="Booktabs table.",
        style="booktabs",
    )
    document = Document("Booktabs Table Test", table)

    docx_path = tmp_path / "booktabs-table.docx"
    pdf_path = tmp_path / "booktabs-table.pdf"
    html_path = tmp_path / "booktabs-table.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages
    )
    html_text = html_path.read_text(encoding="utf-8")
    pdf_border_commands = PdfRenderer()._table_border_style_commands(
        TableStyle.booktabs(),
        build_table_layout(table.header_rows, table.rows),
    )

    assert "Booktabs table." in pdf_text
    assert "Accuracy" in pdf_text
    assert 'w:color="111827"' in docx_xml
    assert 'w:sz="8"' in docx_xml
    assert 'w:sz="5"' in docx_xml
    assert '<w:left w:val="nil"/>' in docx_xml
    assert '<w:right w:val="nil"/>' in docx_xml
    assert "border: none" in html_text
    assert "border-top: 1.00pt solid #111827" in html_text
    assert "border-bottom: 0.60pt solid #111827" in html_text
    assert "border-bottom: 1.00pt solid #111827" in html_text
    assert [command[0] for command in pdf_border_commands] == [
        "LINEABOVE",
        "LINEBELOW",
        "LINEBELOW",
    ]


def test_component_and_template_presets_build_renderable_documents(tmp_path: Path) -> None:
    callout = CalloutBox(
        Paragraph("Check terminology before review."),
        title="Review focus",
        style="warning",
    )
    metadata = KeyValueTable(
        {"Preset namespace": "oodocs.presets.components", "Output": "DOCX/PDF/HTML"},
        caption="Preset metadata.",
    )
    nomenclature = Nomenclature(
        [
            ("A", "area", "m2"),
            ("E", "energy", "kWh"),
            ("q", "heat flux", "W/m2"),
            ("T", "temperature", "degC"),
        ],
        double_column=True,
        title="Nomenclature",
    )
    document = Document("Preset Components", callout, metadata, nomenclature)

    docx_path = tmp_path / "preset-components.docx"
    html_path = tmp_path / "preset-components.html"
    document.save_docx(docx_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    html_text = html_path.read_text(encoding="utf-8")

    assert "Review focus" in docx_xml
    assert "Preset namespace" in docx_xml
    assert "Nomenclature" in docx_xml
    assert "heat flux" in docx_xml
    assert "FFFBEB" in docx_xml
    assert "Review focus" in html_text
    assert "Preset metadata." in _normalized_html_text(html_path)

    article_document = JournalArticleTemplate(include_contents=True, include_references=False).build(
        "Readable manuscript generation",
        authors=[Author("Research Lead", affiliations=["Example Lab"], corresponding=True)],
        abstract="A concise abstract paragraph.",
        keywords=["document generation", "python"],
        sections=[
            ManuscriptSection("Introduction", [Paragraph("Problem and contribution.")]),
            ("Methods", [Paragraph("Data, model, and validation.")]),
        ],
        acknowledgements="The authors thank the review team.",
        data_availability=None,
    )

    assert isinstance(article_document, Document)
    assert any(isinstance(child, TableOfContents) for child in article_document.body.children)
    article_headings = [
        child.plain_title()
        for child in article_document.body.children
        if isinstance(child, Section)
    ]
    assert "Acknowledgements" in article_headings
    assert "Data Availability" not in article_headings

    report_document = TechnicalReportTemplate(include_references=False).build(
        "Validation Report",
        executive_summary="All release checks passed.",
        sections=[("Findings", [Paragraph("The evidence package is complete.")])],
        appendices=[("Evidence Tables", [Paragraph("Detailed evidence.")])],
        back_matter=Paragraph("Distribution list."),
    )
    assert report_document.validate().ok
    assert report_document.settings.cover_page is True
    report_front_matter, report_main_matter = report_document.split_top_level_children()
    assert any(isinstance(child, TableOfContents) for child in report_front_matter)
    assert any(
        isinstance(child, Section) and child.plain_title() == "Executive Summary"
        for child in report_front_matter
    )
    assert report_main_matter[0].plain_title() == "Findings"
    assert any(isinstance(child, Appendix) for child in report_main_matter)
    assert not any(isinstance(child, ReferenceList) for child in report_document.body.children)

    manual_document = SoftwareManualTemplate().build(
        "Command Manual",
        overview="This manual explains the release command workflow.",
        sections=[ManuscriptSection("Install", [Paragraph("Install the package.")])],
        appendices=[("Command Reference", [Paragraph("Command flags.")])],
    )
    assert manual_document.validate().ok
    manual_titles = [
        child.plain_title()
        for child in manual_document.body.children
        if isinstance(child, Section)
    ]
    assert "Overview" in manual_titles
    assert "Install" in manual_titles
    assert not any(isinstance(child, ReferenceList) for child in manual_document.body.children)

    book_document = BookTemplate(include_references=True).build(
        "Engineering Handbook",
        front_matter=Section(
            "Preface",
            Paragraph("Why this handbook exists."),
            level=1,
            numbered=False,
        ),
        parts=[("Operations", [("Getting Started", [Paragraph("Start here.")])])],
        chapters=[("Standalone Chapter", [Paragraph("Standalone content.")])],
        appendices=[("Appendix Data", [Paragraph("Reference tables.")])],
        back_matter=Paragraph("Index placeholder."),
    )
    assert book_document.validate().ok
    book_front_matter, book_main_matter = book_document.split_top_level_children()
    assert any(
        isinstance(child, Section) and child.plain_title() == "Preface"
        for child in book_front_matter
    )
    assert any(
        isinstance(child, Part) and child.plain_title() == "Operations"
        for child in book_main_matter
    )
    assert any(
        isinstance(child, Chapter) and child.plain_title() == "Standalone Chapter"
        for child in book_main_matter
    )
    assert any(isinstance(child, Appendix) for child in book_main_matter)
    assert any(isinstance(child, ReferenceList) for child in book_document.body.children)

    cover_preset = CoverPagePreset.eplus_simple(footer_label="Internal Review")
    cover_settings = cover_preset.settings(
        metadata=DocumentMetadata(subject="Cover preset test"),
        subtitle="Release gate evidence",
        authors=[Author("QA Lead", affiliations=["Example Lab"])],
    )
    assert cover_settings.cover_page is True
    assert cover_settings.author_layout.mode == "stacked"
    assert len(cover_settings.page_items) == 3
    assert all(item.scope.kind == "cover" for item in cover_settings.page_items)
    cover_document = Document(
        "Cover Preset",
        Paragraph("Body content."),
        settings=cover_settings,
    )
    assert cover_document.validate().ok
    cover_html_path = tmp_path / "cover-preset.html"
    cover_document.save_html(cover_html_path)
    cover_html = cover_html_path.read_text(encoding="utf-8")
    assert "Internal Review" in cover_html
    assert "#2563EB" in cover_html

    unitless = Nomenclature([("x", "value"), ("y", "other value")], double_column=True)
    assert unitless.children[0].header_rows[0][0].content.content[0].value == "Symbol"
    assert len(unitless.children[0].header_rows[0]) == 4
    assert unitless.children[0].style == "nomenclature.inner"
    assert note_box("Note body").style == "note"
    assert info_box("Info body").style == "info"
    assert warning_box("Warning body").style == "warning"
    assert success_box("Success body").style == "success"
    danger_callout = CalloutBox("Danger body", variant="danger", icon="!")
    assert danger_callout.style == "danger"
    assert danger_callout.title_fragments() is not None

    try:
        Nomenclature([("x", "value", "-", "extra")])  # type: ignore[list-item]
    except ValueError as exc:
        assert "entries" in str(exc)
    else:
        raise AssertionError("Expected invalid Nomenclature entry shape to fail")


def test_subfigure_group_renders_labels_and_references(tmp_path: Path) -> None:
    image_path = tmp_path / "subfigure.png"
    _write_sample_image(image_path)
    baseline = SubFigure(image_path, caption="Baseline output.", width=1.0)
    treatment = SubFigure(image_path, caption="Treatment output.", width=1.0)
    group = SubFigureGroup(
        baseline,
        treatment,
        caption="Paired outputs.",
        columns=2,
        column_gap=0.2,
    )
    document = Document(
        "Subfigure Test",
        Paragraph("Compare ", baseline.reference(), " with ", treatment.reference(), " in ", group.reference(), "."),
        group,
    )

    docx_path = tmp_path / "subfigure.docx"
    pdf_path = tmp_path / "subfigure.pdf"
    html_path = tmp_path / "subfigure.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
    table_text = "\n".join(
        paragraph.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    assert "Compare Figure 1(a) with Figure 1(b) in Figure 1." in paragraph_text
    assert "(a) Baseline output." in table_text
    assert "(b) Treatment output." in table_text
    assert "Figure 1. Paired outputs." in paragraph_text
    assert len(word_document.inline_shapes) == 2

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Compare Figure 1(a) with Figure 1(b) in Figure 1." in pdf_text
    assert "(a) Baseline output." in pdf_text
    assert "(b) Treatment output." in pdf_text
    assert "Figure 1. Paired outputs." in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "Compare Figure 1(a) with Figure 1(b) in Figure 1" in normalized_html_text
    assert "(a) Baseline output." in normalized_html_text
    assert "(b) Treatment output." in normalized_html_text
    assert "Figure 1. Paired outputs." in normalized_html_text
    assert 'href="#figure_1_a"' in html_text
    assert 'href="#figure_1_b"' in html_text
    assert 'id="figure_1_a"' in html_text
    assert 'id="figure_1_b"' in html_text
    assert html_text.count("data:image/png;base64,") == 2


def test_subtable_group_renders_labels_and_references(tmp_path: Path) -> None:
    baseline = SubTable(
        Table(
            ["Metric", "Value"],
            [["AUC", "0.91"], ["F1", "0.84"]],
            caption="Baseline sensitivity.",
            column_widths=[0.7, 0.7],
        )
    )
    tuned = SubTable(
        Table(
            ["Metric", "Value"],
            [["AUC", "0.94"], ["F1", "0.88"]],
            column_widths=[0.7, 0.7],
        ),
        caption="Tuned sensitivity.",
    )
    group = SubTableGroup(
        baseline,
        tuned,
        caption="Sensitivity table variants.",
        columns=2,
        column_gap=0.15,
    )
    document = Document(
        "Subtable Test",
        Paragraph("Compare ", baseline.reference(), " with ", tuned.reference(), " in ", group.reference(), "."),
        group,
    )

    docx_path = tmp_path / "subtable.docx"
    pdf_path = tmp_path / "subtable.pdf"
    html_path = tmp_path / "subtable.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)

    def collect_table_text(tables: list[object]) -> list[str]:
        parts: list[str] = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(paragraph.text for paragraph in cell.paragraphs)
                    parts.extend(collect_table_text(list(cell.tables)))
        return parts

    table_text = "\n".join(collect_table_text(list(word_document.tables)))
    assert "Compare Table 1(a) with Table 1(b) in Table 1." in paragraph_text
    assert "(a) Baseline sensitivity." in table_text
    assert "(b) Tuned sensitivity." in table_text
    assert "Table 1. Sensitivity table variants." in paragraph_text
    assert "AUC" in table_text
    assert "0.94" in table_text

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Compare Table 1(a) with Table 1(b) in Table 1." in pdf_text
    assert "(a) Baseline sensitivity." in pdf_text
    assert "(b) Tuned sensitivity." in pdf_text
    assert "Table 1. Sensitivity table variants." in pdf_text
    assert "Table 2." not in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "Compare Table 1(a) with Table 1(b) in Table 1" in normalized_html_text
    assert "(a) Baseline sensitivity." in normalized_html_text
    assert "(b) Tuned sensitivity." in normalized_html_text
    assert "Table 1. Sensitivity table variants." in normalized_html_text
    assert "Table 2." not in normalized_html_text
    assert 'href="#table_1_a"' in html_text
    assert 'href="#table_1_b"' in html_text
    assert 'id="table_1_a"' in html_text
    assert 'id="table_1_b"' in html_text


def test_pdf_pages_inserts_external_pdf_pages_with_fallbacks(tmp_path: Path) -> None:
    external_path = tmp_path / "external.pdf"
    Document(
        "External Appendix",
        Paragraph("External page one."),
        PageBreak(),
        Paragraph("External page two."),
    ).save_pdf(external_path)

    imported_pages = PdfPages(external_path, pages=[2], title="External appendix")
    document = Document(
        "PDF Pages Test",
        Paragraph("Before external."),
        imported_pages,
        Paragraph("After external."),
    )
    validation = document.validate()
    assert validation.ok_for(("pdf",))
    assert any(issue.code == "pdf-pages-fallback" for issue in validation.warnings_for(("docx", "html")))

    pdf_path = tmp_path / "merged.pdf"
    docx_path = tmp_path / "merged.docx"
    html_path = tmp_path / "merged.html"
    document.save_pdf(pdf_path)
    document.save_docx(docx_path)
    document.save_html(html_path)

    page_texts = [page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages]
    before_index = next(index for index, text in enumerate(page_texts) if "Before external." in text)
    external_index = next(index for index, text in enumerate(page_texts) if "External page two." in text)
    after_index = next(index for index, text in enumerate(page_texts) if "After external." in text)
    assert before_index < external_index < after_index
    assert "External page one." not in "\n".join(page_texts)
    assert "oodocs-pdfpages-" not in "\n".join(page_texts)

    word_document = WordDocument(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
    assert "PDF pages: External appendix (page 2)" in paragraph_text

    html_text = _normalized_html_text(html_path)
    assert "PDF pages: External appendix" in html_text
    assert "(page 2)" in html_text


def test_caption_and_reference_labels_can_differ_by_theme(tmp_path: Path) -> None:
    image_path = tmp_path / "labels.png"
    _write_sample_image(image_path)
    table = Table(
        headers=["Metric", "Value"],
        rows=[["Latency", "14 ms"]],
        caption="Localized table caption.",
    )
    figure = Figure(image_path, caption="Localized figure caption.", width=1.0)
    document = Document(
        "Caption Labels Test",
        Paragraph("See ", table.reference(), " and ", figure.reference(), "."),
        table,
        figure,
        ListOfTables(),
        ListOfFigures(),
        settings=DocumentSettings(
            theme=Theme(
                captions=CaptionDefaults(
                    table_caption_label="Table",
                    table_reference_label="Tbl.",
                    figure_caption_label="Figure",
                    figure_reference_label="Fig.",
                )
            )
        ),
    )

    docx_path = tmp_path / "caption-labels.docx"
    pdf_path = tmp_path / "caption-labels.pdf"
    html_path = tmp_path / "caption-labels.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    assert "See Tbl. 1 and Fig. 1." in word_text
    assert "Table 1. Localized table caption." in word_text
    assert "Figure 1. Localized figure caption." in word_text

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "See Tbl. 1 and Fig. 1." in pdf_text
    assert "Table 1. Localized table caption." in pdf_text
    assert "Figure 1. Localized figure caption." in pdf_text

    normalized_html_text = _normalized_html_text(html_path)
    assert "See Tbl. 1 and Fig. 1" in normalized_html_text
    assert "Table 1. Localized table caption." in normalized_html_text
    assert "Figure 1. Localized figure caption." in normalized_html_text


def test_cleveref_style_reference_helpers_render_lists_ranges_and_wrappers(tmp_path: Path) -> None:
    first = Table(["Metric"], [["A"]], caption="First table.")
    second = Table(["Metric"], [["B"]], caption="Second table.")
    theorem = Theorem("Every generated reference should keep its label.")
    section = Section("Targets", first, second, theorem)
    document = Document(
        "Clever References",
        section,
        Paragraph(
            "See ",
            Ref(first, style=TextStyle(text_color="C00000")),
            ", ",
            refs([first, second], plural_label="Tables", style=TextStyle(bold=True)),
            ", ",
            ref_range(
                first,
                second,
                plural_label="Tables",
                range_separator="--",
                style=TextStyle(underline=True),
            ),
            ", ",
            paren_ref(theorem),
            ", and ",
            page_ref(section),
            ".",
        ),
    )

    warning_codes = {issue.code for issue in document.validate().warnings}
    assert "page-aware-reference-degrades" in warning_codes

    docx_path = tmp_path / "clever-references.docx"
    pdf_path = tmp_path / "clever-references.pdf"
    html_path = tmp_path / "clever-references.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    with zipfile.ZipFile(docx_path) as archive:
        word_xml = archive.read("word/document.xml").decode("utf-8")
    html_markup = html_path.read_text(encoding="utf-8")
    html_text = _normalized_html_text(html_path)

    expected = "See Table 1, Tables 1 and 2, Tables 1--2, (Theorem 1), and Section 1.1."
    assert expected in word_text
    assert expected in pdf_text
    assert 'w:val="C00000"' in word_xml
    assert "<w:b" in word_xml
    assert 'w:val="single"' in word_xml
    assert "color: #C00000" in html_markup
    assert "font-weight: 700" in html_markup
    assert "text-decoration: underline" in html_markup
    assert "Table 1" in html_text
    assert "Tables 1 and 2" in html_text
    assert "Tables 1 -- 2" in html_text
    assert "Theorem 1" in html_text
    assert "Section 1.1" in html_text


def test_reference_format_overrides_labels_for_helpers(tmp_path: Path) -> None:
    image_path = tmp_path / "reference-format.png"
    _write_sample_image(image_path)
    first = Figure(image_path, caption="First figure.", width=0.5)
    second = Figure(image_path, caption="Second figure.", width=0.5)
    document = Document(
        "Reference Format",
        Paragraph(
            "Compare ",
            refs(
                [first, second],
                reference_format=ReferenceFormat(label="fig.", plural_label="figs."),
                last_separator=" + ",
            ),
            ".",
        ),
        first,
        second,
    )

    html_path = tmp_path / "reference-format.html"
    document.save_html(html_path)

    assert "Compare figs. 1 + 2" in _normalized_html_text(html_path)


def test_explicit_reference_api_covers_numbered_blocks(tmp_path: Path) -> None:
    intro = Paragraph("Reusable paragraph target.")
    equation = Equation(r"\alpha^2 + \beta^2 = \gamma^2")
    snippet = CodeBlock("const ready = true;\nconsole.log(ready);", language="javascript")
    detail_box = Box(Paragraph("Box body."), title="Referenceable Box")
    section = Section("Reference Targets", intro)
    section.children.append(Paragraph(
        "Targets: ",
        intro.reference(),
        ", ",
        reference(equation),
        ", ",
        snippet.reference(),
        ", ",
        detail_box.reference("the boxed note"),
        ", and ",
        section.reference(),
        ".",
    ))
    document = Document(
        "Explicit References",
        section,
        equation,
        snippet,
        detail_box,
    )

    docx_path = tmp_path / "explicit-references.docx"
    pdf_path = tmp_path / "explicit-references.pdf"
    html_path = tmp_path / "explicit-references.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    assert "Targets: Paragraph 1, Equation 1, Code block 1, the boxed note, and Section 1.1." in word_text
    assert "alpha2 + beta2 = gamma2 (1)" in word_text
    assert "const ready = true;" in word_text
    assert "Targets: Paragraph 1, Equation 1, Code block 1, the boxed note, and Section 1.1." in pdf_text
    assert "alpha2 + beta2 = gamma2 (1)" in pdf_text
    assert "Targets:" in normalized_html_text
    assert "Paragraph 1" in normalized_html_text
    assert "Equation 1" in normalized_html_text
    assert "Code block 1" in normalized_html_text
    assert "the boxed note" in normalized_html_text
    assert "Section 1.1" in normalized_html_text
    assert 'href="#paragraph_' in html_text
    assert 'href="#equation_' in html_text
    assert 'href="#code_' in html_text
    assert 'href="#box_' in html_text
    assert 'href="#heading_' in html_text
    assert "color: #" in html_text

    try:
        Paragraph("Ambiguous ", equation, ".")
    except TypeError as exc:
        assert "reference(obj)" in str(exc)
    else:
        raise AssertionError("Expected raw document object references in Paragraph to fail")


def test_countable_blocks_share_document_counter_and_render_references(tmp_path: Path) -> None:
    CustomClaim = create_countable_block_type("Claim", counter="theorem")
    definition = Definition("A countable block participates in the document-wide theorem counter.")
    lemma = Lemma("A later theorem-like block advances the same counter.")
    theorem = Theorem("The shared counter is stable across output formats.", title="Main result")
    proof = Proof("The proof is intentionally unnumbered.")
    example = Example("Examples keep counting after the proof.")
    remark = Remark("Remarks share the same theorem-like sequence.")
    assumption = Assumption("Assumptions can be referenced and numbered.")
    claim = CustomClaim("Custom countable kinds can join the same counter.")
    document = Document(
        "Countable Blocks",
        Paragraph("See ", theorem.reference(), ", ", proof.reference("the proof"), ", and ", claim.reference(), "."),
        definition,
        lemma,
        theorem,
        proof,
        example,
        remark,
        assumption,
        claim,
    )

    render_index = build_render_index(document)
    assert isinstance(claim, CountableBlock)
    assert render_index.countable_number(definition) == 1
    assert render_index.countable_number(lemma) == 2
    assert render_index.countable_number(theorem) == 3
    assert render_index.countable_number(proof) is None
    assert render_index.countable_number(example) == 4
    assert render_index.countable_number(remark) == 5
    assert render_index.countable_number(assumption) == 6
    assert render_index.countable_number(claim) == 7

    docx_path = tmp_path / "countable-blocks.docx"
    pdf_path = tmp_path / "countable-blocks.pdf"
    html_path = tmp_path / "countable-blocks.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    expected_texts = (
        "See Theorem 3, the proof, and Claim 7.",
        "Definition 1.",
        "Lemma 2.",
        "Theorem 3. Main result",
        "Proof.",
        "Example 4.",
        "Remark 5.",
        "Assumption 6.",
        "Claim 7.",
    )
    for text in expected_texts:
        assert text in word_text
        assert text in pdf_text
    for text in expected_texts[1:]:
        assert text in normalized_html_text
    assert "See Theorem 3 , the proof , and Claim 7 ." in normalized_html_text
    assert 'class="oodocs-countable-block' in html_text
    assert 'href="#countable_' in html_text


def test_countable_blocks_can_render_as_theorem_boxes(tmp_path: Path) -> None:
    definition = Definition(
        "A boxed theorem-like block keeps its counter and reference target.",
        title="Boxed term",
        box_style=BoxStyle(
            border=BorderStyle.solid("2563EB", width=0.75),
            background_color="EFF6FF",
            title_background_color="DBEAFE",
            title_text_color="1E3A8A",
            title_position="side",
        ),
    )
    document = Document(
        "Boxed Countable Blocks",
        Paragraph("See ", definition.reference(), "."),
        definition,
    )

    docx_path = tmp_path / "boxed-countable.docx"
    pdf_path = tmp_path / "boxed-countable.pdf"
    html_path = tmp_path / "boxed-countable.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    docx_xml = _docx_document_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_pdf_text = " ".join(pdf_text.split())
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    assert "See Definition 1." in word_text
    assert "Definition 1." in docx_xml
    assert "Boxed term" in docx_xml
    assert "A boxed theorem-like block keeps its counter and reference target." in docx_xml
    assert "See Definition 1." in pdf_text
    assert "Definition 1. Boxed term" in pdf_text
    assert "A boxed theorem-like block keeps its counter and reference target." in normalized_pdf_text
    assert "See Definition 1 ." in normalized_html_text
    assert "Definition 1. Boxed term" in normalized_html_text
    assert "A boxed theorem-like block keeps its counter and reference target." in normalized_html_text
    assert 'class="oodocs-box"' in html_text
    assert 'href="#countable_' in html_text
    assert 'id="countable_' in html_text


def test_algorithm_blocks_render_clauses_steps_code_and_references(tmp_path: Path) -> None:
    prose_algorithm = Algorithm(
        "Coverage aggregation",
        inputs=["test results", "coverage map"],
        outputs=["coverage summary"],
        steps=[
            "Load evidence records.",
            "Group records by feature.",
            "Compute pass/fail counts.",
        ],
        caption="Coverage aggregation algorithm.",
    )
    code_algorithm = Algorithm(
        "Euclidean algorithm",
        code="while b:\n    a, b = b, a % b\nreturn a",
        language="python",
        caption="Euclidean algorithm.",
        body_style="code",
        line_numbers=True,
    )
    document = Document(
        "Algorithm Blocks",
        ListOfAlgorithms(),
        Paragraph("See ", prose_algorithm.reference(), " and ", code_algorithm.reference(), "."),
        prose_algorithm,
        code_algorithm,
    )

    render_index = build_render_index(document)
    assert isinstance(prose_algorithm, CountableBlock)
    assert render_index.countable_number(prose_algorithm) == 1
    assert render_index.countable_number(code_algorithm) == 2
    assert [entry.number for entry in render_index.scoped_algorithms(document.body.children[0])] == [1, 2]
    assert any(issue.code == "html-algorithm-list-page-numbers" for issue in document.validate().warnings)

    docx_path = tmp_path / "algorithms.docx"
    pdf_path = tmp_path / "algorithms.pdf"
    html_path = tmp_path / "algorithms.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    expected_texts = (
        "See Algorithm 1 and Algorithm 2.",
        "List of Algorithms",
        "Algorithm 1. Coverage aggregation algorithm.",
        "Input: test results, coverage map",
        "Output: coverage summary",
        "Load evidence records.",
        "Algorithm 2. Euclidean algorithm.",
        "1 | while b:",
        "2 |     a, b = b, a % b",
        "3 | return a",
    )
    for text in expected_texts:
        assert text in word_text
        assert text in pdf_text
    assert "See Algorithm 1 and Algorithm 2 ." in normalized_html_text
    assert "List of Algorithms" in normalized_html_text
    assert "Algorithm 1. Coverage aggregation algorithm." in normalized_html_text
    assert "Input: test results, coverage map" in normalized_html_text
    assert "Output: coverage summary" in normalized_html_text
    assert "Load evidence records." in normalized_html_text
    assert "Algorithm 2. Euclidean algorithm." in normalized_html_text
    assert "1 | while b:" in normalized_html_text
    assert "2 | a, b = b, a % b" in normalized_html_text
    assert "3 | return a" in normalized_html_text
    assert 'class="oodocs-countable-block oodocs-countable-algorithm"' in html_text
    assert 'href="#countable_' in html_text


def test_algorithm_rejects_ambiguous_body_inputs() -> None:
    try:
        Algorithm("Broken", steps=["one"], code="one")
    except ValueError as exc:
        assert "mutually exclusive" in str(exc)
    else:
        raise AssertionError("Expected Algorithm steps and code conflict to fail")

    try:
        Algorithm("Broken", body_style="table")  # type: ignore[arg-type]
    except ValueError as exc:
        assert "body_style" in str(exc)
    else:
        raise AssertionError("Expected invalid Algorithm body_style to fail")


def test_unnumbered_countable_reference_requires_custom_label() -> None:
    proof = Proof("No counter is assigned to proofs by default.")
    document = Document("Proof Reference", Paragraph("See ", proof.reference(), "."), proof)

    result = document.validate()

    assert [issue.code for issue in result.errors] == ["unnumbered-countable-reference"]


def test_code_block_language_label_can_move_or_hide(tmp_path: Path) -> None:
    visible = CodeBlock("print('visible')", language="python", language_position="bottom-left")
    hidden = CodeBlock("README.md\nsrc/oodocs", language="text", show_language=False)
    document = Document("Code Labels", visible, hidden)

    docx_path = tmp_path / "code-labels.docx"
    pdf_path = tmp_path / "code-labels.pdf"
    html_path = tmp_path / "code-labels.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert "PYTHON" in word_text
    assert "PYTHON" in pdf_text
    assert "oodocs-code-language-bottom-left" in html_text
    assert "oodocs-code-has-label-bottom" in html_text
    assert "TEXT" not in word_text
    assert "TEXT" not in pdf_text
    assert ">TEXT<" not in html_text
    assert "README.md" in word_text
    assert "README.md" in pdf_text
    assert "README.md" in html_text


def test_code_block_caption_line_numbers_highlights_and_from_file(tmp_path: Path) -> None:
    source_path = tmp_path / "example.py"
    source_path.write_text("def double(value):\n    return value * 2\n", encoding="utf-8")
    block = CodeBlock.from_file(
        source_path,
        caption="Reusable implementation.",
        identifier="double-implementation",
        line_numbers=True,
        highlight_lines={2},
        show_language=False,
    )
    document = Document(
        "Code Listing Options",
        Paragraph("See ", block.reference(), "."),
        block,
    )

    docx_path = tmp_path / "code-listing.docx"
    pdf_path = tmp_path / "code-listing.pdf"
    html_path = tmp_path / "code-listing.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    assert block.language == "py"
    assert "See Code block 1." in word_text
    assert "Code block 1. Reusable implementation." in word_text
    assert "1 | def double(value):" in word_text
    assert "2 |     return value * 2" in word_text
    assert "See Code block 1." in pdf_text
    assert "Code block 1. Reusable implementation." in pdf_text
    assert "1 | def double(value):" in pdf_text
    assert "2 |     return value * 2" in pdf_text
    assert "See Code block 1 ." in normalized_html_text
    assert "Code block 1. Reusable implementation." in normalized_html_text
    assert "1 | def double" in normalized_html_text
    assert "(value):" in normalized_html_text
    assert "2 | return value * 2" in normalized_html_text
    assert 'data-oodocs-identifier="double-implementation"' in html_text
    assert "oodocs-code-line-highlight" in html_text
    assert "background-color: #FFF3B0" in html_text
    assert "FFF3B0" in _docx_document_xml(docx_path)


def test_code_block_validation_reports_unrenderable_options() -> None:
    document = Document(
        "Code Listing Validation",
        CodeBlock("print('x')", identifier=" ", highlight_lines={2}),
    )

    result = document.validate()

    assert [issue.code for issue in result.errors] == ["blank-code-block-identifier"]
    assert [issue.code for issue in result.warnings] == ["missing-code-highlight-line"]


def test_pdf_code_block_flowable_wraps_long_unbroken_lines() -> None:
    from oodocs.renderers.pdf import CodeBlockFlowable
    from oodocs.renderers.syntax import SyntaxToken

    flowable = CodeBlockFlowable(
        [SyntaxToken("x" * 120)],
        font_names={
            (False, False): "Courier",
            (True, False): "Courier-Bold",
            (False, True): "Courier-Oblique",
            (True, True): "Courier-BoldOblique",
        },
        font_size=9,
        leading=12,
    )

    width, height = flowable.wrap(90, 400)

    assert width == 90
    assert height > 12
    assert len(flowable._lines) > 1


def test_table_of_contents_uses_page_numbers_and_leaders_by_default(tmp_path: Path) -> None:
    document = Document(
        "TOC Test",
        TableOfContents(),
        Chapter("One", Section("Two", Paragraph("Body"))),
    )

    docx_path = tmp_path / "toc.docx"
    pdf_path = tmp_path / "toc.pdf"
    html_path = tmp_path / "toc.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    docx_settings_xml = _docx_settings_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")

    assert " TOC \\f \\l \"1-9\" \\h \\z " in docx_xml
    assert ' TC "1 One" \\l 1 ' in docx_xml
    assert ' TC "1.1 Two" \\l 2 ' in docx_xml
    assert "PAGEREF heading_" not in docx_xml
    assert 'w:updateFields w:val="true"' in docx_settings_xml
    assert 'w:dirty="true"' in docx_xml
    assert "<w:fldSimple" not in docx_xml
    assert ".  .  ." in pdf_text
    assert "1 One" in pdf_text
    assert 'class="oodocs-toc-page-number"' not in html_text
    assert "oodocs-toc-leader" not in html_text
    assert "target-counter(attr(data-target), page)" not in html_text
    assert "oodocs-toc-entry-no-page" in html_text


def test_table_of_contents_options_can_hide_pages_and_limit_depth(tmp_path: Path) -> None:
    document = Document(
        "TOC Options",
        TableOfContents(
            show_page_numbers=False,
            max_level=1,
            level_styles={1: TocLevelStyle(bold=False, space_after=1)},
        ),
        Chapter("One", Section("Two", Paragraph("Body"))),
    )

    docx_path = tmp_path / "toc-options.docx"
    html_path = tmp_path / "toc-options.html"
    document.save_docx(docx_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    html_text = html_path.read_text(encoding="utf-8")

    assert "PAGEREF heading_" not in docx_xml
    assert 'class="oodocs-toc-page-number"' not in html_text
    assert 'class="oodocs-toc-entry oodocs-toc-entry-no-page oodocs-toc-entry-level-1"' in html_text
    assert 'class="oodocs-toc-entry oodocs-toc-entry-level-2"' not in html_text


def test_caption_lists_use_page_numbers_and_leaders_by_default(tmp_path: Path) -> None:
    document = Document(
        "Caption Lists",
        ListOfTables(),
        ListOfFigures(),
        Table(["Metric"], [["Latency"]], caption="Metric summary."),
        Figure(ImageData(_build_sample_png()), caption="Result plot.", width=1.0),
    )

    docx_path = tmp_path / "caption-lists.docx"
    pdf_path = tmp_path / "caption-lists.pdf"
    html_path = tmp_path / "caption-lists.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    docx_settings_xml = _docx_settings_xml(docx_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    html_text = html_path.read_text(encoding="utf-8")
    warning_codes = {issue.code for issue in document.validate().warnings}

    assert "PAGEREF table_1 \\h" in docx_xml
    assert "PAGEREF figure_1 \\h" in docx_xml
    assert 'w:updateFields w:val="true"' in docx_settings_xml
    assert "Table 1. Metric summary." in pdf_text
    assert "Figure 1. Result plot." in pdf_text
    assert ".  .  ." in pdf_text
    assert 'class="oodocs-caption-list-entry"' in html_text
    assert "oodocs-toc-page-number" not in html_text
    assert "html-table-list-page-numbers" in warning_codes
    assert "html-figure-list-page-numbers" in warning_codes


def test_caption_lists_can_hide_page_numbers(tmp_path: Path) -> None:
    document = Document(
        "Caption List Options",
        ListOfTables(show_page_numbers=False),
        ListOfFigures(show_page_numbers=False),
        Table(["Metric"], [["Latency"]], caption="Metric summary."),
        Figure(ImageData(_build_sample_png()), caption="Result plot.", width=1.0),
    )

    docx_path = tmp_path / "caption-list-options.docx"
    document.save_docx(docx_path)
    warning_codes = {issue.code for issue in document.validate().warnings}

    docx_xml = _docx_document_xml(docx_path)
    assert "PAGEREF table_1 \\h" not in docx_xml
    assert "PAGEREF figure_1 \\h" not in docx_xml
    assert "html-table-list-page-numbers" not in warning_codes
    assert "html-figure-list-page-numbers" not in warning_codes


def test_generated_lists_can_filter_to_part_chapter_or_section(tmp_path: Path) -> None:
    part_figures = ListOfFigures("Part figures", scope="part")
    chapter_toc = TableOfContents("Chapter contents", scope="chapter")
    section_tables = ListOfTables("Section tables", scope="section")
    alpha_table = Table(["Metric"], [["A"]], caption="Alpha table.")
    alpha_nested_table = Table(["Metric"], [["Nested"]], caption="Alpha nested table.")
    beta_table = Table(["Metric"], [["B"]], caption="Beta table.")
    alpha_figure = Figure(ImageData(_build_sample_png()), caption="Alpha figure.", width=1.0)
    beta_figure = Figure(ImageData(_build_sample_png()), caption="Beta figure.", width=1.0)
    outside_figure = Figure(ImageData(_build_sample_png()), caption="Outside figure.", width=1.0)
    document = Document(
        "Scoped Generated Lists",
        Part(
            "Part One",
            part_figures,
            Chapter(
                "Alpha",
                chapter_toc,
                Section(
                    "Alpha Data",
                    section_tables,
                    alpha_table,
                    alpha_figure,
                    Subsection("Alpha Detail", alpha_nested_table),
                ),
                Section("Alpha Notes", Paragraph("Scoped chapter content.")),
            ),
            Chapter("Beta", Section("Beta Data", beta_table, beta_figure)),
        ),
        Part("Part Two", Chapter("Gamma", Section("Gamma Data", outside_figure))),
    )

    render_index = build_render_index(document)

    scoped_heading_titles = [
        "".join(fragment.plain_text() for fragment in entry.title)
        for entry in render_index.scoped_headings(chapter_toc)
    ]
    scoped_table_captions = [
        entry.block.caption.plain_text()
        for entry in render_index.scoped_tables(section_tables)
    ]
    scoped_figure_captions = [
        entry.block.caption.plain_text()
        for entry in render_index.scoped_figures(part_figures)
    ]

    assert scoped_heading_titles == ["Alpha Data", "Alpha Detail", "Alpha Notes"]
    assert scoped_table_captions == ["Alpha table.", "Alpha nested table."]
    assert scoped_figure_captions == ["Alpha figure.", "Beta figure."]

    docx_path = tmp_path / "scoped-generated.docx"
    pdf_path = tmp_path / "scoped-generated.pdf"
    html_path = tmp_path / "scoped-generated.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    html_text = html_path.read_text(encoding="utf-8")
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf_path).pages)

    assert "PAGEREF heading_" in docx_xml
    assert "PAGEREF figure_1 \\h" in docx_xml
    assert "PAGEREF figure_2 \\h" in docx_xml
    assert "PAGEREF figure_3 \\h" not in docx_xml
    part_figures_html = html_text.split("Part figures", 1)[1].split("Chapter contents", 1)[0]
    assert "Alpha figure." in part_figures_html
    assert "Beta figure." in part_figures_html
    assert "Outside figure." not in part_figures_html
    assert "Part figures" in pdf_text
    assert "Chapter contents" in pdf_text
    assert "Section tables" in pdf_text


def test_generated_list_scope_rejects_unknown_values() -> None:
    with pytest.raises(ValueError, match="generated list scope"):
        TableOfContents(scope="page")

    with pytest.raises(ValueError, match="generated list scope"):
        ListOfTables(scope="appendix")


def test_table_of_contents_default_styles_emphasize_only_top_level(tmp_path: Path) -> None:
    document = Document(
        "TOC Style",
        TableOfContents(),
        Chapter("Top", Section("Second", Subsection("Third", Paragraph("Body")))),
    )
    docx_path = tmp_path / "toc-style.docx"
    html_path = tmp_path / "toc-style.html"
    document.save_docx(docx_path)
    document.save_html(html_path)

    docx_xml = _docx_document_xml(docx_path)
    assert ' TC "1 Top" \\l 1 ' in docx_xml
    assert ' TC "1.1 Second" \\l 2 ' in docx_xml
    assert ' TC "1.1.1 Third" \\l 3 ' in docx_xml
    assert " TOC \\f \\l \"1-9\" \\h \\z " in docx_xml
    html_text = html_path.read_text(encoding="utf-8")
    assert 'oodocs-toc-entry-level-1" style="margin-left: 0.00in; margin-top: 12.0pt; margin-bottom: 7.0pt' in html_text
    assert 'oodocs-toc-entry-level-2" style="margin-left: 0.24in; margin-top: 3.0pt; margin-bottom: 3.0pt' in html_text
    assert 'oodocs-toc-entry-level-2" style=' in html_text and "font-weight: 400" in html_text


def test_bibtex_string_creates_citation_library() -> None:
    document = Document(
        "Bibliography Test",
        Paragraph("Example"),
        citations="""@misc{oodocs-repository,
  title = {oodocs},
  organization = {Gonie-Gonie},
  year = {2026},
  url = {https://github.com/Gonie-Gonie/oo-docs},
  note = {GitHub repository}
}""",
    )

    entry = document.citations.resolve("oodocs-repository")
    assert entry.title == "oodocs"
    assert entry.organization == "Gonie-Gonie"
    assert entry.url == "https://github.com/Gonie-Gonie/oo-docs"
    assert "GitHub repository" in entry.format_reference()


def test_bibtex_file_parser_handles_quotes_and_duplicate_keys(tmp_path: Path) -> None:
    bib_path = tmp_path / "references.bib"
    bib_path.write_text(
        """@string{ignored = "Journal"}
@article{doe2024,
  title = "Reliable {APIs}, Revisited",
  author = "Doe, Jane and Smith, John",
  journal = {Journal of Docs},
  year = {2024},
  url = {https://example.com/reliable-apis}
}
""",
        encoding="utf-8",
    )

    library = CitationLibrary.from_bibtex_file(bib_path)
    entry = library.resolve("doe2024")

    assert entry.title == "Reliable APIs, Revisited"
    assert entry.authors == ("Doe, Jane", "Smith, John")
    assert entry.publisher == "Journal of Docs"
    assert entry.url == "https://example.com/reliable-apis"

    with pytest.raises(OODocsError, match="Duplicate citation key"):
        CitationLibrary.from_bibtex(
            """@misc{same, title={First}}
@misc{same, title={Second}}
"""
        )


def test_citation_and_reference_styles_can_be_configured(tmp_path: Path) -> None:
    source = CitationSource(
        "Literate Programming",
        key="knuth",
        authors=("Donald E. Knuth",),
        publisher="The Computer Journal",
        year="1984",
        url="https://doi.org/10.1093/comjnl/27.2.97",
    )
    document = Document(
        "Citation Style Test",
        Paragraph("Prior work ", cite("knuth"), " remains relevant."),
        ReferenceList(),
        settings=DocumentSettings(
            theme=Theme(citations=CitationDefaults(citation_style="apa", reference_style="apa")),
        ),
        citations=[source],
    )

    assert source.format_reference("apa").startswith(
        "Knuth, D. E. (1984). Literate Programming."
    )

    docx_path = tmp_path / "citation-style.docx"
    pdf_path = tmp_path / "citation-style.pdf"
    html_path = tmp_path / "citation-style.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    paragraph_texts = [paragraph.text for paragraph in WordDocument(docx_path).paragraphs]
    assert any("Prior work (Knuth, 1984) remains relevant." in text for text in paragraph_texts)
    assert any(
        text.startswith("Knuth, D. E. (1984). Literate Programming.")
        for text in paragraph_texts
    )
    assert not any(text.startswith("[1] Knuth") for text in paragraph_texts)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Prior work (Knuth, 1984) remains relevant." in pdf_text
    assert "Knuth, D. E. (1984). Literate Programming." in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "Prior work (Knuth, 1984) remains relevant." in normalized_html_text
    assert "Knuth, D. E. (1984). Literate Programming." in normalized_html_text
    assert 'href="#citation_1"' in html_text
    assert 'id="citation_1"' in html_text
    assert '<span class="oodocs-generated-marker">' not in html_text


def test_reference_list_can_include_uncited_entries_and_sort(tmp_path: Path) -> None:
    cited = CitationSource(
        "Beta Work",
        key="beta",
        authors=("Zed, Zelda",),
        publisher="Journal B",
        year="2020",
    )
    uncited = CitationSource(
        "Alpha Work",
        key="alpha",
        authors=("Able, Alice",),
        publisher="Journal A",
        year="2021",
    )
    document = Document(
        "Reference Policy",
        Paragraph("Cited work ", cite("beta"), "."),
        ReferenceList(include_uncited=True),
        settings=DocumentSettings(
            theme=Theme(
                citations=CitationDefaults(
                    citation_style="apa",
                    reference_style="apa",
                    reference_sort="author",
                )
            ),
        ),
        citations=[cited, uncited],
    )

    html_path = tmp_path / "references.html"
    document.save_html(html_path)

    html_text = _normalized_html_text(html_path)
    assert "Cited work (Zed, 2020)" in html_text
    assert "Able, A. (2021). Alpha Work." in html_text
    assert "Zed, Z. (2020). Beta Work." in html_text
    assert html_text.index("Able, A. (2021). Alpha Work.") < html_text.index("Zed, Z. (2020). Beta Work.")
    assert "empty-references-page" not in {issue.code for issue in document.validate().warnings}


def test_document_accepts_document_settings() -> None:
    settings = DocumentSettings(
        metadata=DocumentMetadata(keywords="settings, metadata"),
        metadata_author="OODocs",
        summary="Settings test",
        subtitle="Grouped metadata",
        authors=[
            Author(
                "Example Author",
                affiliations=[Affiliation(organization="Example Lab")],
            )
        ],
        author_layout=AuthorLayout(mode="stacked"),
        cover_page=True,
        unit="cm",
        theme=Theme(page_numbers=PageNumberDefaults(show_page_numbers=True)),
    )

    document = Document("Configured", Paragraph("Body"), settings=settings)

    assert document.settings.resolved_author() == "OODocs"
    assert document.settings.resolved_metadata_title(document.title) == "Configured"
    assert document.settings.resolved_metadata_subject() == "Settings test"
    assert document.settings.resolved_metadata_keywords() == ("settings", "metadata")
    assert document.settings.summary == "Settings test"
    assert document.settings.subtitle is not None
    assert document.settings.subtitle[0].plain_text() == "Grouped metadata"
    assert document.settings.authors[0].name == "Example Author"
    assert document.settings.authors[0].affiliations[0].formatted() == "Example Lab"
    assert document.settings.author_layout.mode == "stacked"
    assert document.settings.cover_page is True
    assert document.settings.unit == "cm"
    assert round(document.settings.get_text_width(), 2) == 15.92
    assert document.settings.theme.page_numbers.show_page_numbers is True


def test_document_metadata_maps_to_renderer_outputs(tmp_path: Path) -> None:
    settings = DocumentSettings(
        metadata=DocumentMetadata(
            title="Metadata Title",
            author="Metadata Author",
            subject="Metadata Subject",
            keywords=["alpha", "beta"],
            description="Metadata description.",
        ),
        summary="Legacy summary fallback",
    )
    document = Document("Visible Title", Paragraph("Body"), settings=settings)

    docx_path = tmp_path / "metadata.docx"
    pdf_path = tmp_path / "metadata.pdf"
    html_path = tmp_path / "metadata.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    properties = WordDocument(docx_path).core_properties
    assert properties.title == "Metadata Title"
    assert properties.author == "Metadata Author"
    assert properties.subject == "Metadata Subject"
    assert properties.keywords == "alpha, beta"
    assert properties.comments == "Metadata description."

    pdf_metadata = PdfReader(str(pdf_path)).metadata
    assert pdf_metadata.title == "Metadata Title"
    assert pdf_metadata.author == "Metadata Author"
    assert pdf_metadata.subject == "Metadata Subject"
    assert pdf_metadata["/Keywords"] == "alpha, beta"

    html_text = html_path.read_text(encoding="utf-8")
    assert "<title>Metadata Title</title>" in html_text
    assert 'name="description" content="Metadata description."' in html_text
    assert 'name="author" content="Metadata Author"' in html_text
    assert 'name="subject" content="Metadata Subject"' in html_text
    assert 'name="keywords" content="alpha, beta"' in html_text


def test_theme_link_style_controls_html_link_defaults(tmp_path: Path) -> None:
    document = Document(
        "Styled Links",
        Paragraph("Open ", link("https://example.com", "Example")),
        settings=DocumentSettings(
            theme=Theme(
                links=LinkDefaults(
                    TextStyle(text_color="C00000", underline=False),
                )
            )
        ),
    )

    html_path = tmp_path / "links.html"
    document.save_html(html_path)
    html_text = html_path.read_text(encoding="utf-8")

    assert "color: #C00000;" in html_text
    assert "text-decoration: none;" in html_text


def test_url_helper_preserves_targets_and_warns_for_raw_long_urls(tmp_path: Path) -> None:
    long_target = (
        "https://example.com/releases/download/v1.2.3/"
        + "portable-artifact-segment-" * 5
        + "?checksum="
        + "abcdef1234567890" * 4
    )
    document = Document(
        "Breakable URLs",
        Paragraph("Download ", url(long_target), " or open ", url(long_target, label="release page"), "."),
    )

    result = document.validate()
    assert "overly-long-url" not in {issue.code for issue in result.warnings}

    docx_path = tmp_path / "breakable-url.docx"
    pdf_path = tmp_path / "breakable-url.pdf"
    html_path = tmp_path / "breakable-url.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert long_target.replace("&", "&amp;") in html_text
    assert "\u200b" in html_text
    assert long_target in _docx_relationships_xml(docx_path)
    assert long_target in _pdf_uri_targets(pdf_path)

    raw_document = Document("Raw URL", Paragraph("Download ", link(long_target), "."))
    labeled_document = Document("Labeled URL", Paragraph("Download ", link(long_target, "release page"), "."))

    assert "overly-long-url" in {issue.code for issue in raw_document.validate().warnings}
    assert "overly-long-url" not in {issue.code for issue in labeled_document.validate().warnings}


def test_validation_reports_broken_internal_hyperlinks() -> None:
    document = Document(
        "Broken Links",
        Paragraph(
            "See ",
            inline_components.Hyperlink.internal_anchor("missing-anchor", "missing"),
        ),
    )

    result = document.validate()

    assert "broken-internal-link" in {issue.code for issue in result.errors}


def test_validation_accepts_internal_hyperlinks_to_generated_anchors() -> None:
    target = Section("Target", Paragraph("Details"), anchor="target")
    document = Document(
        "Valid Links",
        target,
        Paragraph(
            "See ",
            inline_components.Hyperlink.internal_anchor("target", "target"),
        ),
    )

    result = document.validate()

    assert "broken-internal-link" not in {issue.code for issue in result.errors}


def test_print_units_include_common_document_units() -> None:
    assert round(length_to_inches(25.4, "mm"), 8) == 1.0
    assert length_to_inches(6, "pc") == 1.0
    assert length_to_inches(1440, "twip") == 1.0
    assert length_to_inches(72, "pt") == 1.0


def test_document_unit_applies_to_media_dimensions(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    _write_sample_image(image_path)
    table = Table(
        headers=["A", "B"],
        rows=[["one", "two"]],
        column_widths=[2.54, 5.08],
    )
    figure = Figure(image_path, width=2.54)
    document = Document(
        "Metric Dimensions",
        table,
        figure,
        settings=DocumentSettings(unit="cm"),
    )

    docx_path = tmp_path / "metric.docx"
    pdf_path = tmp_path / "metric.pdf"
    html_path = tmp_path / "metric.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_document = WordDocument(docx_path)
    assert abs(int(word_document.inline_shapes[0].width) - 914400) <= 1
    assert len(PdfReader(BytesIO(pdf_path.read_bytes())).pages) == 1
    html_text = html_path.read_text(encoding="utf-8")
    assert 'style="width: 1.00in;"' in html_text
    assert 'style="width: 2.00in;"' in html_text
    assert 'width: 1.00in; max-width: 100%; height: auto' in html_text


def test_page_size_and_margins_render_to_all_outputs(tmp_path: Path) -> None:
    settings = DocumentSettings(
        unit="cm",
        page_size=PageSize(20, 10, unit="cm"),
        page_margins=PageMargins.symmetric(vertical=1.5, horizontal=2.0, unit="cm"),
    )
    document = Document("Margins", Paragraph("Body"), settings=settings)

    docx_path = tmp_path / "margins.docx"
    pdf_path = tmp_path / "margins.pdf"
    html_path = tmp_path / "margins.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_section = WordDocument(docx_path).sections[0]
    assert abs(int(word_section.page_width) - int(20 / 2.54 * 914400)) <= 300
    assert abs(int(word_section.left_margin) - int(2 / 2.54 * 914400)) <= 300
    pdf_page = PdfReader(BytesIO(pdf_path.read_bytes())).pages[0]
    assert round(float(pdf_page.mediabox.width), 1) == round(20 / 2.54 * 72, 1)
    html_text = html_path.read_text(encoding="utf-8")
    assert "size: 7.87in 3.94in;" in html_text
    assert "margin: 0.59in 0.79in 0.59in 0.79in;" in html_text
    assert "max-width: 6.30in;" in html_text


def test_page_layout_landscape_renders_to_all_outputs(tmp_path: Path) -> None:
    layout = PageLayout.landscape(
        PageSize.a4(),
        PageMargins.all(1.0, unit="cm"),
    )
    settings = DocumentSettings(unit="cm", page_layout=layout)
    document = Document("Landscape Layout", Paragraph("Body"), settings=settings)

    assert settings.page_layout.orientation == "landscape"
    assert settings.page_width_in_inches() > settings.page_height_in_inches()
    assert round(settings.get_text_width(unit="cm"), 1) == 27.7
    assert round(PageSize.a4().landscape().width, 1) == 29.7
    assert round(PageSize(11.0, 8.5, unit="in").portrait().height, 1) == 11.0

    docx_path = tmp_path / "landscape.docx"
    pdf_path = tmp_path / "landscape.pdf"
    html_path = tmp_path / "landscape.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_section = WordDocument(docx_path).sections[0]
    assert int(word_section.page_width) > int(word_section.page_height)
    assert abs(int(word_section.page_width) - int(29.7 / 2.54 * 914400)) <= 300
    pdf_page = PdfReader(BytesIO(pdf_path.read_bytes())).pages[0]
    assert float(pdf_page.mediabox.width) > float(pdf_page.mediabox.height)
    html_text = html_path.read_text(encoding="utf-8")
    assert "size: 11.69in 8.27in;" in html_text
    assert "margin: 0.39in 0.39in 0.39in 0.39in;" in html_text
    assert "max-width: 10.91in;" in html_text


def test_page_layout_rejects_ambiguous_settings_geometry() -> None:
    with pytest.raises(ValueError, match="page_layout cannot be combined"):
        DocumentSettings(
            page_layout=PageLayout.landscape(),
            page_size=PageSize.letter(),
        )


def test_section_page_layout_switches_docx_pdf_and_html(tmp_path: Path) -> None:
    default_layout = PageLayout.portrait(
        PageSize.a4(),
        PageMargins.all(1.0, unit="cm"),
    )
    landscape_layout = PageLayout.landscape(
        PageSize.a4(),
        PageMargins.all(1.0, unit="cm"),
    )
    document = Document(
        "Section Layout",
        Section("Default Section", Paragraph("Portrait body."), level=1),
        Section(
            "Wide Section",
            Paragraph("Landscape body."),
            level=1,
            page_layout=landscape_layout,
        ),
        Paragraph("Back to the default page layout."),
        settings=DocumentSettings(unit="cm", page_layout=default_layout),
    )

    warning_codes = {issue.code for issue in document.validate().warnings}
    assert "section-page-layout-html-degrade" in warning_codes

    docx_path = tmp_path / "section-layout.docx"
    pdf_path = tmp_path / "section-layout.pdf"
    html_path = tmp_path / "section-layout.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_sections = list(WordDocument(docx_path).sections)
    assert len(word_sections) >= 3
    assert int(word_sections[1].page_width) > int(word_sections[1].page_height)
    assert int(word_sections[-1].page_width) < int(word_sections[-1].page_height)

    pdf_pages = PdfReader(BytesIO(pdf_path.read_bytes())).pages
    page_orientations = [
        float(page.mediabox.width) > float(page.mediabox.height)
        for page in pdf_pages
    ]
    assert True in page_orientations
    assert page_orientations[-1] is False

    html_text = html_path.read_text(encoding="utf-8")
    assert "oodocs-section-page-layout" in html_text
    assert "--oodocs-page-width: 11.69in;" in html_text
    assert "--oodocs-page-height: 8.27in;" in html_text


def test_object_unit_overrides_document_unit(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    _write_sample_image(image_path)
    figure = Figure(image_path, width=1.0, unit="in")
    table = Table(
        headers=["A"],
        rows=[["one"]],
        column_widths=[1.0],
        unit="in",
    )
    document = Document(
        "Local Dimensions",
        table,
        figure,
        settings=DocumentSettings(unit="cm"),
    )

    html_path = tmp_path / "local.html"
    document.save_html(html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert 'style="width: 1.00in;"' in html_text
    assert 'width: 1.00in; max-width: 100%; height: auto' in html_text


def test_figure_height_and_text_width_helpers_render(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    _write_sample_image(image_path)
    settings = DocumentSettings(
        unit="cm",
        page_size=PageSize.a4(),
        page_margins=PageMargins.all(2.0, unit="cm"),
    )
    figure = Figure(
        image_path,
        width=settings.get_text_width(0.5),
        height=3.0,
    )
    document = Document("Figure Size", figure, settings=settings)

    docx_path = tmp_path / "figure-size.docx"
    html_path = tmp_path / "figure-size.html"
    document.save_docx(docx_path)
    document.save_html(html_path)

    shape = WordDocument(docx_path).inline_shapes[0]
    assert abs(int(shape.width) - int((settings.text_width_in_inches() * 0.5) * 914400)) <= 1
    assert abs(int(shape.height) - int((3.0 / 2.54) * 914400)) <= 1
    html_text = html_path.read_text(encoding="utf-8")
    assert "width: 3.35in" in html_text
    assert "height: 1.18in" in html_text


def test_auto_footnotes_page_can_be_disabled(tmp_path: Path) -> None:
    document = Document(
        "Inline Notes",
        Paragraph(
            "Portable ",
            footnote("term", "Suppressed footnote note."),
            " stays inline.",
        ),
        settings=DocumentSettings(
            theme=Theme(
                blocks=BlockDefaults(
                    auto_footnotes_page=False,
                    footnote_placement="document",
                )
            )
        ),
    )

    docx_path = tmp_path / "inline-notes.docx"
    pdf_path = tmp_path / "inline-notes.pdf"
    html_path = tmp_path / "inline-notes.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    paragraph_texts = [paragraph.text for paragraph in WordDocument(docx_path).paragraphs]
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_html_text = _normalized_html_text(html_path)

    assert "Footnotes" not in paragraph_texts
    assert all("Suppressed footnote note." not in text for text in paragraph_texts)
    assert "Footnotes" not in pdf_text
    assert "Suppressed footnote note." not in pdf_text
    assert "Footnotes" not in normalized_html_text
    assert "Suppressed footnote note." not in normalized_html_text


def test_docx_native_page_footnotes_are_default(tmp_path: Path) -> None:
    document = Document(
        "Inline Notes",
        Paragraph(
            "Portable ",
            footnote("term", "Default native footnote note."),
            " stays inline.",
        ),
    )

    docx_path = tmp_path / "inline-notes.docx"
    document.save_docx(docx_path)

    paragraph_texts = [paragraph.text for paragraph in WordDocument(docx_path).paragraphs]
    assert "Footnotes" not in paragraph_texts
    assert all("Default native footnote note." not in text for text in paragraph_texts)

    with zipfile.ZipFile(docx_path) as archive:
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
    assert "Default native footnote note." in footnotes_xml
    assert "w:footnoteReference" in _docx_document_xml(docx_path)


def test_custom_footnote_streams_use_generated_markers_across_outputs(tmp_path: Path) -> None:
    document = Document(
        "Custom Footnote Streams",
        Paragraph(
            "Symbol ",
            footnote("alpha", "Symbol stream note.", stream="symbols"),
            " review ",
            footnote("beta", "Review stream note.", stream="review"),
            " default ",
            footnote("gamma", "Default stream note."),
            ".",
        ),
        settings=DocumentSettings(
            theme=Theme(
                footnotes=FootnoteDefaults(
                    stream_styles={
                        "symbols": FootnoteStyle.symbol(("*", "#")),
                        "review": FootnoteStyle(CounterStyle(prefix="R")),
                    }
                )
            )
        ),
    )

    validation = document.validate()
    assert "docx-footnote-stream-generated-list" in {
        issue.code for issue in validation.warnings_for(("docx",))
    }
    assert validation.warnings_for(("pdf", "html")) == ()

    docx_path = tmp_path / "custom-footnote-streams.docx"
    pdf_path = tmp_path / "custom-footnote-streams.pdf"
    html_path = tmp_path / "custom-footnote-streams.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    word_text = "\n".join(paragraph.text for paragraph in WordDocument(docx_path).paragraphs)
    word_xml = _docx_document_xml(docx_path)
    with zipfile.ZipFile(docx_path) as archive:
        assert "word/footnotes.xml" not in archive.namelist()
    assert "Footnotes" in word_text
    assert "[*] Symbol stream note." in word_text
    assert "[R1] Review stream note." in word_text
    assert "[1] Default stream note." in word_text
    assert "w:footnoteReference" not in word_xml

    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages
    )
    assert "[*] Symbol stream note." in pdf_text
    assert "[R1] Review stream note." in pdf_text
    assert "[1] Default stream note." in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert 'id="footnote_1"' in html_text
    assert 'id="footnote_2"' in html_text
    assert "[*]" in normalized_html_text
    assert "[R1]" in normalized_html_text
    assert "[1]" in normalized_html_text
    assert "Symbol stream note." in normalized_html_text


def test_todo_and_margin_notes_render_with_comment_fallbacks(tmp_path: Path) -> None:
    document = Document(
        "Review Annotations",
        Paragraph(
            "Assumption ",
            MarginNote("Check this assumption beside the source paragraph.", side="left"),
            " needs review ",
            Todo("Verify units before release.", owner="QA"),
            ".",
        ),
    )

    validation = document.validate()
    assert "margin-note-renderer-fallback" in {
        issue.code for issue in validation.warnings_for(("docx", "pdf"))
    }
    assert validation.warnings_for(("html",)) == ()

    docx_path = tmp_path / "review-annotations.docx"
    pdf_path = tmp_path / "review-annotations.pdf"
    html_path = tmp_path / "review-annotations.html"
    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    with zipfile.ZipFile(docx_path) as archive:
        comments_xml = archive.read("word/comments.xml").decode("utf-8")
    word_xml = _docx_document_xml(docx_path)
    assert "Check this assumption beside the source paragraph." in comments_xml
    assert "Verify units before release." in comments_xml
    assert "QA" in comments_xml
    assert "TODO" in word_xml
    assert "w:commentReference" in word_xml

    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages
    )
    assert "TODO" in pdf_text
    assert "Comments" in pdf_text
    assert "Check this assumption beside the source paragraph." in pdf_text
    assert "Verify units before release." in pdf_text

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert 'class="oodocs-margin-note oodocs-margin-note-left"' in html_text
    assert "TODO" in normalized_html_text
    assert "Check this assumption beside the source paragraph." in normalized_html_text
    assert "Verify units before release." in normalized_html_text


def test_document_renders_to_docx_and_pdf(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    _write_sample_image(image_path)
    repository_source = CitationSource(
        "oodocs",
        organization="Gonie-Gonie",
        publisher="GitHub repository",
        year="2026",
        url="https://github.com/Gonie-Gonie/oo-docs",
    )
    registered_source = CitationSource(
        "Release Notes",
        key="release-notes",
        organization="Gonie-Gonie",
        publisher="Documentation index",
        year="2026",
        url="https://github.com/Gonie-Gonie/oo-docs/releases",
    )
    unused_source = CitationSource(
        "Internal Draft",
        key="internal-draft",
        organization="Gonie-Gonie",
        year="2026",
        url="https://example.invalid/internal-draft",
    )
    artifacts_table = Table(
        headers=["Type", "Status"],
        rows=[
            ["DOCX", Paragraph("generated ", footnote("state", "Table cell footnote note."))],
            ["PDF", "generated"],
        ],
        caption="Generated artifacts.",
        column_widths=[2.5, 2.5],
        style=TableStyle(
            header_background_color="#DCE8F4",
            alternate_row_background_color="#F7FAFD",
        ),
    )
    workflow_frame = FakeDataFrame(
        columns=[("Workflow", "Step"), ("Workflow", "Target"), ("Result", "")],
        rows=[
            ["Draft review", "DOCX", "ready"],
            ["Release", "PDF", "published"],
        ],
    )
    workflow_table = Table(
        workflow_frame,
        caption="Output workflow.",
        column_widths=[2.2, 1.6, 1.6],
        style=TableStyle(alternate_row_background_color="#EEF4FA"),
    )
    merged_header_table = Table(
        headers=[
            [TableCell("Metrics", colspan=2), TableCell("Summary", rowspan=2)],
            ["Latency", "Quality"],
        ],
        rows=[
            [
                TableCell("14 ms"),
                TableCell("stable", background_color="#EEF6FF"),
                TableCell("ready"),
            ]
        ],
        caption="Merged header table.",
        column_widths=[1.6, 1.6, 1.6],
        style=TableStyle(header_background_color="#D9E6F2"),
    )
    preview_figure = Figure(
        image_path,
        caption=Paragraph("Tiny sample image."),
        width=1.0,
    )
    figure_object = FakeFigure(_build_sample_png(width=320, height=180))
    preview_figure_second = Figure(
        figure_object,
        caption=Paragraph("Second tiny sample image."),
        width=1.2,
    )
    boxed_detail = Box(
        Paragraph("A boxed paragraph can live alongside nested objects."),
        Table(
            headers=["Scope", "State"],
            rows=[["Box", "stable"]],
            column_widths=[1.4, 1.4],
        ),
        Figure(
            image_path,
            width=0.7,
        ),
        title="Review Box",
        style=BoxStyle(
            border=BorderStyle.solid("#7A8CA5", width=0.75),
            background_color="#F4F8FC",
            title_background_color="#DDE8F4",
        ),
    )

    document = Document(
        "Pipeline Report",
        TableOfContents(),
        Chapter(
            "Summary",
            Section(
                "Highlights",
                Paragraph(
                    "The review ",
                    comment("note", "Check the generated outputs before release.", author="pytest", initials="PT"),
                    " appears inline and is also exported to the comments page.",
                ),
                HighlightedParagraph(
                    "The ",
                    bold("oodocs"),
                    " pipeline supports ",
                    italic("styled"),
                    " text, ",
                    inline_code("code"),
                    ", and ",
                    text_color("custom color", "#0066AA", style=TextStyle(bold=True)),
                    ".",
                    style=ParagraphStyle(space_after=14),
                ),
                Paragraph(
                    markup("Inline helpers also support **bold** and *italic* markup."),
                    " Inline math such as ",
                    math(r"\alpha^2 + \beta^2 = \gamma^2"),
                    " is supported as well.",
                ),
                Paragraph(
                    "See ",
                    reference(artifacts_table),
                    " and ",
                    preview_figure.reference(),
                    " for the generated outputs.",
                ),
                Paragraph(
                    "Repository status is tracked in ",
                    cite(repository_source),
                    ".",
                ),
                Paragraph(
                    "Registered bibliography entries can still be cited as ",
                    cite("release-notes"),
                    ".",
                ),
                Paragraph(
                    "Portable footnotes such as ",
                    footnote("term", "Paragraph footnote note."),
                    " are collected automatically on the footnotes page.",
                ),
                boxed_detail,
                Subsection(
                    "Artifacts",
                    BulletList(
                        "Lists render into both DOCX and PDF.",
                        Paragraph("Paragraph instances can also be list items."),
                    ),
                    SubSubsection(
                        "Export Steps",
                        CodeBlock(
                            "from oodocs import Document\n\ndocument.save_docx('report.docx')\ndocument.save_pdf('report.pdf')",
                            language="python",
                        ),
                    ),
                    Equation(r"\int_0^1 \alpha x^2 \, dx = \frac{\alpha}{3}"),
                    NumberedList("Create the model", "Render the files"),
                    artifacts_table,
                    workflow_table,
                    merged_header_table,
                    preview_figure,
                    preview_figure_second,
                ),
            ),
        ),
        ListOfTables(),
        ListOfFigures(),
        CommentList(),
        ReferenceList(),
        settings=DocumentSettings(
            metadata_author="pytest",
            summary="Renderer integration test",
            theme=Theme(
                page_numbers=PageNumberDefaults(
                    show_page_numbers=True,
                    page_number_template="Page {page}",
                    page_number_alignment="center",
                ),
                blocks=BlockDefaults(
                    footnote_placement="document",
                    heading_numbering=HeadingNumbering(),
                    bullet_list_style=ListStyle(
                        marker=CounterStyle(
                            counter_format="bullet",
                            bullet="\u2022",
                            suffix="",
                        ),
                    ),
                    numbered_list_style=ListStyle(
                        marker=CounterStyle(suffix="."),
                    ),
                ),
            ),
        ),
        citations=[registered_source, unused_source],
    )

    docx_path = tmp_path / "report.docx"
    pdf_path = tmp_path / "report.pdf"
    html_path = tmp_path / "report.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    assert docx_path.exists()
    assert pdf_path.exists()
    assert html_path.exists()
    assert docx_path.stat().st_size > 0
    assert pdf_path.stat().st_size > 0
    assert html_path.stat().st_size > 0

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    assert "Pipeline Report" in paragraph_texts
    assert "1 Summary" in paragraph_texts
    assert "1.1 Highlights" in paragraph_texts
    assert "1.1.1 Artifacts" in paragraph_texts
    assert "1.1.1.1 Export Steps" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert "Comments" in paragraph_texts
    assert "Footnotes" in paragraph_texts
    assert "List of Tables" in paragraph_texts
    assert "List of Figures" in paragraph_texts
    assert "References" in paragraph_texts
    assert any("The review note[1] appears inline" in text for text in paragraph_texts)
    assert any("oodocs" in text for text in paragraph_texts)
    assert any(text == "1 Summary" for text in paragraph_texts)
    assert any(text == "1.1 Highlights" for text in paragraph_texts)
    assert any("See Table 1 and Figure 1 for the generated outputs." in text for text in paragraph_texts)
    assert any("Repository status is tracked in [1]." in text for text in paragraph_texts)
    assert any("Registered bibliography entries can still be cited as [2]." in text for text in paragraph_texts)
    assert any("Portable footnotes such as term" in text and "collected automatically on the footnotes page." in text for text in paragraph_texts)
    assert any("Inline math such as" in text and "2 + " in text and " = " in text for text in paragraph_texts)
    assert any("dx = (" in text and ")/(3)" in text for text in paragraph_texts)
    assert any("Table cell footnote note." in text for text in paragraph_texts)
    assert any("Paragraph footnote note." in text for text in paragraph_texts)
    assert any("[1] Check the generated outputs before release." in text for text in paragraph_texts)
    assert any(text == "\u2022 Lists render into both DOCX and PDF." for text in paragraph_texts)
    assert any(text == "1. Create the model" for text in paragraph_texts)
    assert sum("Table 1. Generated artifacts." in text for text in paragraph_texts) >= 2
    assert sum("Table 2. Output workflow." in text for text in paragraph_texts) >= 2
    assert sum("Table 3. Merged header table." in text for text in paragraph_texts) >= 2
    assert sum("Figure 1. Tiny sample image." in text for text in paragraph_texts) >= 2
    assert sum("Figure 2. Second tiny sample image." in text for text in paragraph_texts) >= 2
    assert any("https://github.com/Gonie-Gonie/oo-docs" in text for text in paragraph_texts)
    assert any("https://github.com/Gonie-Gonie/oo-docs/releases" in text for text in paragraph_texts)
    assert all("internal-draft" not in text.lower() for text in paragraph_texts)
    assert any("from oodocs import Document" in text for text in paragraph_texts)
    assert len(word_document.inline_shapes) == 3
    summary_paragraph = next(
        paragraph
        for paragraph in word_document.paragraphs
        if "The review note[1] appears inline" in paragraph.text
    )
    assert summary_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    assert len(word_document.tables) == 4
    assert "Review Box" in word_document.tables[0].cell(0, 0).text
    assert word_document.tables[1].cell(1, 0).text == "DOCX"
    assert word_document.tables[1].cell(1, 1).text.startswith("generated")
    assert word_document.tables[1].cell(2, 1).text == "generated"
    assert word_document.tables[2].cell(2, 0).text == "Draft review"
    assert word_document.tables[2].cell(3, 1).text == "PDF"
    assert word_document.tables[3].cell(2, 0).text == "14 ms"
    assert word_document.tables[3].cell(2, 2).text == "ready"
    assert word_document.styles["Normal"].font.name == "Times New Roman"
    assert word_document.styles["Title"].font.name == "Times New Roman"
    assert word_document.styles["Heading 1"].font.name == "Times New Roman"
    assert word_document.styles["Heading 2"].font.name == "Times New Roman"
    assert word_document.styles["Heading 3"].font.name == "Times New Roman"
    assert word_document.styles["Heading 4"].font.name == "Times New Roman"
    assert word_document.styles["Title"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 1"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 2"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 3"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 4"].font.color.rgb == RGBColor(0, 0, 0)
    heading_styles = {paragraph.text: paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text in {"Comments", "List of Tables", "List of Figures", "References"}}
    assert heading_styles["Comments"] == "Heading 2"
    assert heading_styles["List of Tables"] == "Heading 2"
    assert heading_styles["List of Figures"] == "Heading 2"
    assert heading_styles["References"] == "Heading 2"
    assert next(paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text == "Footnotes") == "Heading 2"
    assert len(word_document.comments) == 1
    assert "Check the generated outputs before release." in "\n".join(
        paragraph.text
        for comment_item in word_document.comments
        for paragraph in comment_item.paragraphs
    )
    footer_xml = word_document.sections[0].footer.paragraphs[0]._p.xml
    assert "<w:instrText" in footer_xml
    assert " PAGE " in footer_xml
    assert word_document.sections[0].footer.paragraphs[0].text.startswith("Page ")
    inline_math_paragraph = next(paragraph for paragraph in word_document.paragraphs if "Inline math such as" in paragraph.text)
    assert any(run.text == "2" and run.font.superscript for run in inline_math_paragraph.runs)
    equation_paragraph = next(paragraph for paragraph in word_document.paragraphs if "dx = (" in paragraph.text and ")/(3)" in paragraph.text)
    assert any(run.text == "2" and run.font.superscript for run in equation_paragraph.runs)
    assert any(run.text == "0" and run.font.subscript for run in equation_paragraph.runs)
    assert any(run.text == "1" and run.font.superscript for run in equation_paragraph.runs)
    docx_xml = _docx_document_xml(docx_path)
    assert "Review Box" in docx_xml
    assert "A boxed paragraph can live alongside nested objects." in docx_xml
    assert "stable" in docx_xml
    assert "D9E6F2" in docx_xml
    assert "DCE8F4" in docx_xml
    assert "PAGEREF table_1 \\h" in docx_xml
    assert "PAGEREF figure_1 \\h" in docx_xml
    assert len(figure_object.calls) >= 2
    assert all(call.get("format") == "png" for call in figure_object.calls)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Pipeline Report" in pdf_text
    assert "1 Summary" in pdf_text
    assert "1.1 Highlights" in pdf_text
    assert "1.1.1 Artifacts" in pdf_text
    assert "1.1.1.1 Export Steps" in pdf_text
    assert "Contents" in pdf_text
    assert "Comments" in pdf_text
    assert "Footnotes" in pdf_text
    assert "The review note[1] appears inline and is also exported to the comments page." in pdf_text
    assert "See Table 1 and Figure 1 for the generated outputs." in pdf_text
    assert "Repository status is tracked in [1]." in pdf_text
    assert "Registered bibliography entries can still be cited as [2]." in pdf_text
    assert "Portable footnotes such as term" in pdf_text and "collected automatically on the footnotes page." in pdf_text
    assert "Inline math such as" in pdf_text
    assert "dx = (" in pdf_text
    assert "Table cell footnote note." in pdf_text
    assert "Paragraph footnote note." in pdf_text
    assert "[1] Check the generated outputs before release." in pdf_text
    assert "Review Box" in pdf_text
    assert "A boxed paragraph can live alongside nested objects." in pdf_text
    assert "stable" in pdf_text
    assert pdf_text.count("Table 1. Generated artifacts.") >= 2
    assert pdf_text.count("Table 2. Output workflow.") >= 2
    assert pdf_text.count("Table 3. Merged header table.") >= 2
    assert pdf_text.count("Figure 1. Tiny sample image.") >= 2
    assert pdf_text.count("Figure 2. Second tiny sample image.") >= 2
    assert "List of Tables" in pdf_text
    assert "List of Figures" in pdf_text
    assert "References" in pdf_text
    assert "https://github.com/Gonie-Gonie/oo-docs" in pdf_text
    assert "https://github.com/Gonie-Gonie/oo-docs/releases" in pdf_text
    assert "Internal Draft" not in pdf_text
    assert "Lists render into both DOCX and PDF." in pdf_text
    assert "from oodocs import Document" in pdf_text
    assert "1\nLists render into both DOCX and PDF." not in pdf_text
    assert "1.\nCreate the model" in pdf_text
    assert _pdf_image_draw_count(pdf_path) == 3
    pdf_fonts = _pdf_font_names(pdf_path)
    assert any(font == "/Times-Roman" or "TimesNewRomanPSMT" in font for font in pdf_fonts)
    assert any(font == "/Times-Bold" or "TimesNewRomanPS-Bold" in font for font in pdf_fonts)
    assert any(
        font in {"/Times-BoldItalic", "/Times-Italic"}
        or "TimesNewRomanPS-BoldItalic" in font
        or "TimesNewRomanPS-Italic" in font
        for font in pdf_fonts
    )
    assert any(font.startswith("/Courier") or "CourierNewPS" in font for font in pdf_fonts)
    pdf_content = _pdf_content_bytes(pdf_path)
    assert b"Page 1" in pdf_content
    assert b"18 Tf" in pdf_content
    assert b"15 Tf" in pdf_content
    assert b"13 Tf" in pdf_content
    assert b"11.5 Tf" in pdf_content
    assert b"Comments" in pdf_content
    assert b"15 Tf" in _pdf_text_context(pdf_path, "List of Tables")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "List of Figures")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "Comments")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "References")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "Contents")

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "Pipeline Report" in normalized_html_text
    assert "1 Summary" in normalized_html_text
    assert "1.1 Highlights" in normalized_html_text
    assert "1.1.1 Artifacts" in normalized_html_text
    assert "1.1.1.1 Export Steps" in normalized_html_text
    assert "Contents" in normalized_html_text
    assert "Comments" in normalized_html_text
    assert "Footnotes" in normalized_html_text
    assert "List of Tables" in normalized_html_text
    assert "List of Figures" in normalized_html_text
    assert "References" in normalized_html_text
    assert "The review note [1] appears inline and is also exported to the comments page." in normalized_html_text
    assert "See Table 1 and Figure 1 for the generated outputs." in normalized_html_text
    assert "Repository status is tracked in [1]" in normalized_html_text
    assert "Registered bibliography entries can still be cited as [2]" in normalized_html_text
    assert "Portable footnotes such as term" in normalized_html_text
    assert "Table cell footnote note." in normalized_html_text
    assert "Paragraph footnote note." in normalized_html_text
    assert "[1] Check the generated outputs before release." in normalized_html_text
    assert "Review Box" in normalized_html_text
    assert "A boxed paragraph can live alongside nested objects." in normalized_html_text
    assert "stable" in normalized_html_text
    assert normalized_html_text.count("Table 1. Generated artifacts.") >= 2
    assert normalized_html_text.count("Table 2. Output workflow.") >= 2
    assert normalized_html_text.count("Table 3. Merged header table.") >= 2
    assert normalized_html_text.count("Figure 1. Tiny sample image.") >= 2
    assert normalized_html_text.count("Figure 2. Second tiny sample image.") >= 2
    assert "https://github.com/Gonie-Gonie/oo-docs" in normalized_html_text
    assert "https://github.com/Gonie-Gonie/oo-docs/releases" in normalized_html_text
    assert "Internal Draft" not in normalized_html_text
    assert "Lists render into both DOCX and PDF." in normalized_html_text
    assert "from oodocs import Document" in normalized_html_text
    assert html_text.count("data:image/png;base64,") == 3
    assert 'href="#table_1"' in html_text
    assert 'href="#figure_1"' in html_text
    assert 'id="table_1"' in html_text
    assert 'id="figure_1"' in html_text
    assert 'id="citation_1"' in html_text
    assert 'id="comment_1"' in html_text
    assert 'id="footnote_1"' in html_text
