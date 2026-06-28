from __future__ import annotations

import importlib.util
import json
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module():
    module_path = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "project_metadata_report"
        / "main.py"
    )
    spec = importlib.util.spec_from_file_location("project_metadata_report_main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_project_metadata_report_example_builds_outputs(tmp_path: Path) -> None:
    example = _load_example_module()
    document = example.build_document()
    outputs = example.build(tmp_path)

    assert document.validate().ok
    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])
    assert outputs.metadata_json.exists()

    sidecar = json.loads(outputs.metadata_json.read_text(encoding="utf-8"))
    assert sidecar["pyproject"] == "pyproject.toml"
    assert sidecar["workflow"] == ".github/workflows/release.yml"
    assert sidecar["project"]["name"] == "oodocs"
    assert sidecar["workflow_jobs"]

    word_document = WordDocument(outputs["docx"])
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages)
    html_text = _normalized_html_text(outputs["html"])

    assert "Project Metadata Report" in paragraph_texts
    assert "1 Configuration Sources" in paragraph_texts
    assert "2 Package Metadata" in paragraph_texts
    assert "3 Release Workflow" in paragraph_texts
    assert "pyproject.toml" in table_text
    assert "release.yml" in table_text
    assert any("Package metadata from pyproject.toml." in text for text in paragraph_texts)
    assert any("GitHub Actions jobs in the release workflow." in text for text in paragraph_texts)
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=(
            "Project Metadata Report",
            "1 Configuration Sources",
            "2 Package Metadata",
            "3 Release Workflow",
        ),
        min_tables=4,
    )
    assert "Project Metadata Report" in pdf_text
    assert "Package Metadata" in pdf_text
    assert "GitHub Actions workflow" in pdf_text
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=("Project Metadata Report", "pyproject.toml", "release.yml"),
        min_pages=1,
    )
    assert "Project Metadata Report" in html_text
    assert "ProjectMetadata.from_pyproject(...)" in html_text
    assert_html_internal_links_resolve(outputs["html"])


def test_project_metadata_report_example_supports_common_cli(
    tmp_path: Path,
    capsys,
) -> None:
    example = _load_example_module()
    output_dir = tmp_path / "cli"

    outputs = example.build(output_dir / "programmatic", output_formats=("html",))
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()
    assert outputs.metadata_json.exists()

    example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "project-metadata-report.html").exists()
    assert (output_dir / "project-metadata.json").exists()
