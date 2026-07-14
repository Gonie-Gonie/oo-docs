from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader
import pytest

from oodocs import (
    BulletList,
    CodeBlock,
    Document,
    Equation,
    Paragraph,
    Text,
    footnote,
    inline_math,
)
from oodocs.components.descriptions import DescriptionItem, DescriptionList
from oodocs.styles import DescriptionListStyle, ParagraphStyle, StyleSheet, TextStyle


pytestmark = pytest.mark.render


def test_description_list_model_coerces_blocks_and_returns_plain_records() -> None:
    item = DescriptionItem(
        [Text("alpha"), " / ", inline_math(r"x^2")],
        "A string definition.",
        BulletList("First", "Second"),
        CodeBlock("value = 42", language="python"),
        Equation(r"x^2", numbered=False),
    )
    descriptions = DescriptionList([item]).add("beta", "Another definition.")

    assert isinstance(item.children[0], Paragraph)
    assert descriptions.as_records() == [
        {
            "term": "alpha / x2",
            "definition": "A string definition.\nFirst\nSecond\nvalue = 42\nx2",
        },
        {"term": "beta", "definition": "Another definition."},
    ]
    assert descriptions.add("gamma", Paragraph("Third.")) is descriptions
    assert DescriptionItem("delta").add("Fourth.").plain_definition() == "Fourth."

    with pytest.raises(TypeError, match="DescriptionList items"):
        DescriptionList(["not an item"])  # type: ignore[list-item]


def test_description_list_styles_validate_register_and_round_trip() -> None:
    styles = StyleSheet.default()

    assert set(styles.description_list) >= {
        "description.default",
        "description.compact",
        "description.symbols",
    }
    symbols = styles.resolve("description_list", "description.symbols")
    assert symbols.term_text_style.font_name == "Cambria Math"
    assert symbols.term_width_in_inches() == pytest.approx(0.9)

    restored = StyleSheet.from_dict(styles.to_dict())
    restored_compact = restored.resolve("description_list", "description.compact")
    assert restored_compact.item_spacing == pytest.approx(3.0)
    assert isinstance(restored_compact.term_text_style, TextStyle)
    assert isinstance(restored_compact.definition_style, ParagraphStyle)

    with pytest.raises(ValueError, match="layout"):
        DescriptionListStyle(layout="columns")  # type: ignore[arg-type]
    with pytest.raises(ValueError, match="term_width"):
        DescriptionListStyle(term_width=0)


def _rich_description_document(style: DescriptionListStyle | str | None = None) -> Document:
    target = Paragraph("Linked target.")
    descriptions = DescriptionList(style=style)
    descriptions.add(
        "--output FORMAT",
        Paragraph(
            "Select an output. See ",
            target.ref("the linked target"),
            " and note",
            footnote("note marker", "Portable footnote inside a definition."),
            ".",
        ),
        BulletList("HTML output", "PDF output"),
        CodeBlock("tool --output html", language="text"),
        Equation(r"E = mc^2", numbered=False),
    )
    descriptions.add(
        "a-very-long-configuration-key-that-needs-to-wrap-cleanly",
        Paragraph("First line.\nSecond line with a longer explanation."),
    )
    return Document("Description List", target, descriptions)


def test_description_list_renders_semantically_to_html_docx_and_pdf(tmp_path: Path) -> None:
    document = _rich_description_document("description.default")
    result = document.validate()

    assert result.ok, result.format_text()
    outputs = document.save_all(tmp_path, stem="descriptions")
    html = outputs["html"].read_text(encoding="utf-8")
    word = WordDocument(outputs["docx"])
    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages
    )

    assert '<dl class="oodocs-description-list oodocs-description-hanging"' in html
    assert html.count('<dt class="oodocs-description-term"') == 2
    assert html.count('<dd class="oodocs-description-definition"') == 2
    assert 'href="#paragraph_1"' in html
    assert "Portable footnote inside a definition." in html
    assert "tool --output html" in html

    assert word.tables
    table_xml = word.tables[0]._tbl.xml
    assert 'w:val="nil"' in table_xml
    table_text = "\n".join(
        cell.text
        for table in word.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "--output FORMAT" in table_text
    assert "HTML output" in table_text
    assert "tool --output html" in table_text

    assert "--output FORMAT" in pdf_text
    assert "Select an output" in pdf_text
    assert "E = mc" in pdf_text


@pytest.mark.parametrize("layout", ["stacked", "run-in"])
def test_description_list_alternate_layouts_render(layout: str, tmp_path: Path) -> None:
    style = DescriptionListStyle(layout=layout)  # type: ignore[arg-type]
    document = _rich_description_document(style)
    outputs = document.save_all(tmp_path / layout, stem=layout)

    html = outputs["html"].read_text(encoding="utf-8")
    assert f"oodocs-description-{layout}" in html
    assert outputs["docx"].exists()
    assert outputs["pdf"].exists()


def test_long_description_list_splits_across_pdf_pages(tmp_path: Path) -> None:
    descriptions = DescriptionList(style="description.compact")
    for index in range(90):
        descriptions.add(
            f"very-long-term-{index:03d}-with-wrapping-content",
            Paragraph(
                "First definition line with enough content to wrap.\n"
                "Second definition line remains attached to the same semantic item."
            ),
        )
    document = Document("Long Description List", descriptions)

    pdf_path = tmp_path / "long-descriptions.pdf"
    document.save_pdf(pdf_path)
    reader = PdfReader(BytesIO(pdf_path.read_bytes()))

    assert len(reader.pages) >= 3
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "very-long-term-000" in text
    assert "very-long-term-089" in text
