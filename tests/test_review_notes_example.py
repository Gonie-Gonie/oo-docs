from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import zipfile

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module():
    module_path = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "review_notes_example"
        / "main.py"
    )
    spec = importlib.util.spec_from_file_location("review_notes_example_main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_review_notes_example_builds_outputs(tmp_path: Path) -> None:
    example = _load_example_module()
    document = example.build_document()
    outputs = example.build(tmp_path)

    assert document.validate().ok
    assert example.review_queue_rows()[0] == [
        "copy",
        "open",
        "Confirm release date and product wording.",
    ]
    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])

    word_document = WordDocument(outputs["docx"])
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    comment_text = "\n".join(
        paragraph.text
        for word_comment in word_document.comments
        for paragraph in word_comment.paragraphs
    )
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages)
    html_text = _normalized_html_text(outputs["html"])
    with zipfile.ZipFile(outputs["docx"]) as archive:
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")

    assert "Review Notes Example" in paragraph_texts
    assert "1 Review Workflow" in paragraph_texts
    assert "2 Generated Review Pages" in paragraph_texts
    assert "Collected Review Notes" in paragraph_texts
    assert "Collected Footnotes" in paragraph_texts
    assert any("Core comments stay attached to source prose" in text for text in paragraph_texts)
    assert "Verify benchmark fixture names." in comment_text
    assert any("Verify benchmark fixture names." in text for text in paragraph_texts)
    assert "Service-level agreement used for the review scenario." in footnotes_xml
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=(
            "Review Notes Example",
            "1 Review Workflow",
            "2 Generated Review Pages",
        ),
        min_tables=1,
    )
    assert "Review Notes Example" in pdf_text
    assert "oodocs.review" in pdf_text
    assert "Collected Review Notes" in pdf_text
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=("Review Notes Example", "Collected Review Notes", "Collected Footnotes"),
        min_pages=1,
    )
    assert "Review Notes Example" in html_text
    assert "risk note" in html_text
    assert_html_internal_links_resolve(outputs["html"])


def test_review_notes_example_supports_common_cli(tmp_path: Path, capsys) -> None:
    example = _load_example_module()
    output_dir = tmp_path / "cli"

    outputs = example.build(output_dir / "programmatic", output_formats=("html",))
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()

    example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "review-notes.html").exists()
