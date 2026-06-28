from __future__ import annotations

import importlib.util
import json
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module():
    module_path = (
        Path(__file__).resolve().parents[1]
        / "examples"
        / "validation_gate_report"
        / "main.py"
    )
    spec = importlib.util.spec_from_file_location("validation_gate_report_main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_validation_gate_report_example_builds_outputs(tmp_path: Path) -> None:
    example = _load_example_module()
    candidate_result = example.build_candidate_document().validate()
    outputs = example.build(tmp_path)

    assert {issue.code for issue in candidate_result.warnings} >= {
        "html-toc-page-numbers",
        "wide-table",
    }
    passed, denied = example.evaluate_gate(candidate_result)
    assert passed
    assert denied == ()
    assert_rendered_bundle(outputs["docx"], outputs["pdf"], outputs["html"])
    assert outputs.validation_json.exists()

    sidecar = json.loads(outputs.validation_json.read_text(encoding="utf-8"))
    codes = {issue["code"] for issue in sidecar["issues"]}
    assert {"html-toc-page-numbers", "wide-table"} <= codes

    word_document = WordDocument(outputs["docx"])
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(outputs["pdf"].read_bytes())).pages)
    html_text = _normalized_html_text(outputs["html"])

    assert "Validation Gate Report" in paragraph_texts
    assert "1 Validation Policy" in paragraph_texts
    assert "2 Validation Result Table" in paragraph_texts
    assert "3 Release Gate Pattern" in paragraph_texts
    assert "html-toc-page-numbers" in table_text
    assert "wide-table" in table_text
    assert any("Temporary write pattern for avoiding stale release artifacts." in text for text in paragraph_texts)
    assert_docx_structure(
        outputs["docx"],
        required_paragraphs=(
            "Validation Gate Report",
            "1 Validation Policy",
            "2 Validation Result Table",
        ),
        min_tables=3,
    )
    assert "Validation Gate Report" in pdf_text
    assert "Document.validate()" in pdf_text
    assert "validation-result.json" in pdf_text
    assert_pdf_text_and_pages(
        outputs["pdf"],
        required_text=("Validation Gate Report", "wide-table", "validation-result.json"),
        min_pages=1,
    )
    assert "Validation Gate Report" in html_text
    assert "Temporary write pattern" in html_text
    assert_html_internal_links_resolve(outputs["html"])


def test_validation_gate_report_example_supports_common_cli(
    tmp_path: Path,
    capsys,
) -> None:
    example = _load_example_module()
    output_dir = tmp_path / "cli"

    outputs = example.build(output_dir / "programmatic", output_formats=("html",))
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()
    assert outputs.validation_json.exists()

    example.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "validation-gate-report.html").exists()
    assert (output_dir / "validation-result.json").exists()
