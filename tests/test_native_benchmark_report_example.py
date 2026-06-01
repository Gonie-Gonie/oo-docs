from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_native_benchmark_report_example_builds_outputs(tmp_path: Path) -> None:
    benchmark_example = _load_example_module("native_benchmark_report")

    payload = benchmark_example.generate_payload(24)
    results = benchmark_example.benchmark_normalizers(payload, repeat=2, inner_iterations=2)

    assert len(payload) == 24
    assert len(results) == 3
    assert len({result["checksum"] for result in results}) == 1
    assert results == sorted(results, key=lambda result: float(result["median_ms"]))

    docx_path, pdf_path = benchmark_example.build_native_benchmark_report(tmp_path)
    html_path = tmp_path / "native-python-benchmark.html"

    assert_rendered_bundle(docx_path, pdf_path, html_path)

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    normalized_html_text = _normalized_html_text(html_path)

    assert "Native Python Benchmark Report" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert "List of Tables" in paragraph_texts
    assert "1 Benchmark as a Document Workflow" in paragraph_texts
    assert "2 Workload" in paragraph_texts
    assert "3 Results" in paragraph_texts
    assert "4 Reusing the Pattern" in paragraph_texts
    assert any("The fastest candidate in this run was" in text for text in paragraph_texts)
    assert any(
        "Serialized Python-to-document flow used by this example." in text
        for text in paragraph_texts
    )
    assert any(
        "Native Python benchmark results converted directly from measured data." in text
        for text in paragraph_texts
    )
    assert "character loop" in table_text
    assert "replace + split" in table_text
    assert "translate + split" in table_text
    assert len(word_document.tables) == 3
    assert_docx_structure(
        docx_path,
        required_paragraphs=(
            "Native Python Benchmark Report",
            "Contents",
            "List of Tables",
            "1 Benchmark as a Document Workflow",
        ),
        table_count=3,
    )

    assert "Native Python Benchmark Report" in pdf_text
    assert "Benchmark as a Document Workflow" in pdf_text
    assert "Serialized Python-to-document flow" in pdf_text
    assert "NORMALIZERS" in pdf_text
    assert "document.save_all" in pdf_text
    assert len(pdf_reader.pages) >= 3
    assert_pdf_text_and_pages(
        pdf_path,
        required_text=(
            "Native Python Benchmark Report",
            "Benchmark as a Document Workflow",
            "NORMALIZERS",
        ),
        min_pages=3,
    )

    assert "Native Python Benchmark Report" in normalized_html_text
    assert "Documenting measured Python work without leaving Python" in normalized_html_text
    assert "results: list[dict]" in normalized_html_text
    assert "Native Python benchmark results converted directly from measured data." in normalized_html_text
    assert "save_all" in normalized_html_text
    assert_html_internal_links_resolve(
        html_path,
        required_hrefs=("#heading_1", "#table_1", "#table_3"),
        required_text=("Native Python Benchmark Report", "results: list[dict]"),
    )
