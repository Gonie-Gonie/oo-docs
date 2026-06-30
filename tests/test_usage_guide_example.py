from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import zipfile

from docx import Document as WordDocument
from pypdf import PdfReader

from example_regression import (
    assert_docx_structure,
    assert_html_internal_links_resolve,
    assert_pdf_text_and_pages,
    assert_rendered_bundle,
)


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_usage_guide_example_builds_outputs(tmp_path: Path) -> None:
    usage_guide = _load_example_module("usage_guide_example")
    outputs = usage_guide.build_usage_guide(tmp_path)
    docx_path = outputs["docx"]
    pdf_path = outputs["pdf"]
    html_path = outputs["html"]

    assert_rendered_bundle(docx_path, pdf_path, html_path)
    assert (Path(usage_guide.__file__).resolve().parent / "assets" / "oodocs-logo.png").exists()

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    table_text = "\n".join(
        cell.text
        for table in word_document.tables
        for row in table.rows
        for cell in row.cells
    )
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    assert "OODocs User Guide" in paragraph_texts
    assert "Reference-style guide for structured Python document authoring" in paragraph_texts
    assert "OODocs Contributors" in paragraph_texts
    assert "Hyeong-Gon Jo" in paragraph_texts
    assert "Open-source documentation workflow" not in paragraph_texts
    assert any(
        "Document credits separate names, affiliations, and document roles." in text
        for text in paragraph_texts
    )
    assert "Maintainers and release editors" in table_text
    assert "Repository steward" in table_text
    assert "Building Simulation LAB, Seoul National University" in table_text
    assert "Guide Cover" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert "List of Tables" in paragraph_texts
    assert "List of Figures" in paragraph_texts
    assert "Comments" in paragraph_texts
    assert "References" in paragraph_texts
    assert "1 Overview" in paragraph_texts
    assert "2 Metadata and Title Matter" in paragraph_texts
    assert "3 Document Model" in paragraph_texts
    assert "4 Tables, Figures, and Cross-References" in paragraph_texts
    assert "5 Notes, Comments, and References" in paragraph_texts
    assert "6 Layout and Pagination" in paragraph_texts
    assert "6.4.1 Subsection entries" in paragraph_texts
    assert "6.4.1.1 Fourth-level section entries" in paragraph_texts
    assert "7 Project Structure and Scaling Up" in paragraph_texts
    assert "Part III" in paragraph_texts
    assert "Presets and Templates" in paragraph_texts
    assert "8 Component Presets" in paragraph_texts
    assert "9 Template Presets" in paragraph_texts
    assert any("AuthorLayout(mode='stacked')" in text for text in paragraph_texts)
    assert any("Theme(blocks=BlockDefaults(footnote_placement='document'))" in text for text in paragraph_texts)
    assert any("PageMargins" in text for text in paragraph_texts)
    assert any("PageBreak()" in text for text in paragraph_texts)
    assert any("settings.get_text_width(0.75)" in text for text in paragraph_texts)
    assert any("placement='inline'" in text for text in paragraph_texts)
    assert any("Inline image example:" in text for text in paragraph_texts)
    assert any("Compact inline chips cover categories" in text for text in paragraph_texts)
    assert any("keyboard(\"Ctrl+Enter\")" in text for text in paragraph_texts)
    assert any("prescript(...)" in text for text in paragraph_texts)
    assert any("Inline text can carry ordinary scripts such as H2O and x2, plus front scripts like 146C." in text for text in paragraph_texts)
    assert any("Pygments highlighting" in text for text in paragraph_texts)
    assert any("JAVASCRIPT" in text for text in paragraph_texts)
    assert any("SQL" in text for text in paragraph_texts)
    assert any("YAML" in text for text in paragraph_texts)
    assert any("CodeBlock.from_file(...)" in text for text in paragraph_texts)
    assert any("line_numbers" in text for text in paragraph_texts)
    assert any("highlight_lines" in text for text in paragraph_texts)
    assert "ListOfListings" in table_text
    assert any("Table(split=False)" in text for text in paragraph_texts)
    assert any("Table(split=True)" in text for text in paragraph_texts)
    assert any("continuation_label" in text for text in paragraph_texts)
    assert any("continued_caption_template" in text for text in paragraph_texts)
    assert any("DOCX and PDF use repeated header rows" in text for text in paragraph_texts)
    assert any("HTML keeps a plain flow policy" in text for text in paragraph_texts)
    assert any("display: table-header-group" in text for text in paragraph_texts)
    assert any("CompactTable.grouped_headers" in text for text in paragraph_texts)
    assert any("TableCell(rowspan=...)" in text for text in paragraph_texts)
    assert any("TableCell(colspan=...)" in text for text in paragraph_texts)
    assert any("Table.grouped_headers(...)" in text for text in paragraph_texts)
    assert any("rowspan=2" in text for text in paragraph_texts)
    assert any('style="booktabs"' in text for text in paragraph_texts)
    assert "top_rule, header_rule, bottom_rule" in table_text
    assert "publication-style horizontal rules without vertical grid lines" in table_text
    assert any("ColumnSpec(key=" in text for text in paragraph_texts)
    assert any("visible=False" in text for text in paragraph_texts)
    assert any("full_matrix.excerpt" in text for text in paragraph_texts)
    assert any("full_matrix.save_csv" in text for text in paragraph_texts)
    assert any("record-column visibility" in text for text in paragraph_texts)
    assert any("'tbp'" in text or 'placement="tbp"' in text for text in paragraph_texts)
    assert any("TableOfContents" in text for text in paragraph_texts)
    assert any("TableOfContentsLevelStyle" in text for text in paragraph_texts)
    assert "Scoped lists" in table_text
    assert "scope='part'" in table_text
    assert "scope='chapter'" in table_text
    assert "scope='section'" in table_text
    assert "HTML generated-list page numbers" in table_text
    assert "show_page_numbers=False" in table_text
    assert "Appendix child chapters use A, B, C numbering" in table_text
    assert "Table and figure captions keep document-wide numbering." in table_text
    assert "validation warns if main chapters continue afterward." in table_text
    assert any("A reading map for the guide." in text for text in paragraph_texts)
    assert any("Example catalog" in text for text in paragraph_texts)
    assert any("Purpose-based entry points for the bundled examples." in text for text in paragraph_texts)
    assert any("LaTeX habits translated into oodocs's Python-first authoring model." in text for text in paragraph_texts)
    assert any("For authors coming from LaTeX" in text for text in paragraph_texts)
    assert any("Page layout controls shared across renderers." in text for text in paragraph_texts)
    assert any("Table-of-contents defaults and customization options." in text for text in paragraph_texts)
    assert any("Document-level configuration options." in text for text in paragraph_texts)
    assert any("Grouped Theme defaults" in text for text in paragraph_texts)
    assert any("Block-level option scope" in text for text in paragraph_texts)
    assert any("HeadingStyle" in text for text in paragraph_texts)
    assert "heading_styles" in table_text
    assert "heading_style" in table_text
    assert any("Component presets wrap ordinary blocks" in text for text in paragraph_texts)
    assert any("Template presets build full Document objects" in text for text in paragraph_texts)
    assert any("oodocs.presets.components" in text for text in paragraph_texts)
    assert any("oodocs.presets.templates" in text for text in paragraph_texts)
    assert any("JournalArticleTemplate" in text for text in paragraph_texts)
    assert any("Your Paper Your Way" in text for text in paragraph_texts)
    assert any("Instructions for Authors" in text for text in paragraph_texts)
    assert "Nomenclature" in table_text
    assert any("Figure sizing patterns for width, height, and document-relative sizing." in text for text in paragraph_texts)
    assert any("CropBox" in text for text in paragraph_texts)
    assert any("rotation=90" in text for text in paragraph_texts)
    assert any("alt_text" in text for text in paragraph_texts)
    assert any("Coordinate-based drawings can be page overlays or inline flow objects." in text for text in paragraph_texts)
    assert any("Advanced table and figure placement controls." in text for text in paragraph_texts)
    assert any("SubFigure" in text for text in paragraph_texts)
    assert any("SubFigureGroup" in text for text in paragraph_texts)
    assert any("SubTable" in text for text in paragraph_texts)
    assert any("SubTableGroup" in text for text in paragraph_texts)
    assert any("Table 1(a)" in text for text in paragraph_texts)
    assert any("PdfPages" in text for text in paragraph_texts)
    assert any("link-style placeholder" in text for text in paragraph_texts)
    assert any("Renderer-specific behavior for notes, review workflows, and cross-reference stability." in text for text in paragraph_texts)
    assert any("Numbered statements, proofs, and custom counters" in text for text in paragraph_texts)
    assert any("CountableBlock" in text for text in paragraph_texts)
    assert any("Theorem 3. Stable references" in text for text in paragraph_texts)
    assert any("Exercise 1." in text for text in paragraph_texts)
    assert any("Build and validate from the CLI" in text for text in paragraph_texts)
    assert any("oodocs build report.py --out artifacts" in text for text in paragraph_texts)
    assert any("Validation results are structured objects" in text for text in paragraph_texts)
    assert any("Command-line builds and validation all call the same high-level workflow API." in text for text in paragraph_texts)
    assert any("portable footnotes exactly where the text appears." in text for text in paragraph_texts)
    assert any("github.com/Gonie-Gonie/oo-docs" in text for text in paragraph_texts)
    assert any("The journal example at examples/journal_paper_example/main.py" in text for text in paragraph_texts)
    assert any("Document.from_markdown(...)" in text for text in paragraph_texts)
    assert any("parse_markdown(...)" in text for text in paragraph_texts)
    assert any("Document.from_notebook(...)" in text for text in paragraph_texts)
    assert any("parse_notebook(...)" in text for text in paragraph_texts)
    assert any("Notebook-backed report" in text for text in paragraph_texts)
    assert any("Release note digest" in text for text in paragraph_texts)
    assert "Document Python API objects" in table_text
    assert "api_objects_example" in table_text
    assert "style_cleanup_smoke" in table_text
    assert "OODocs Contributor Certificate" in table_text
    assert "Footnotes" not in [text for text in paragraph_texts if text == "Footnotes"]
    assert len(word_document.tables) == 27
    assert len(word_document.inline_shapes) == 11
    assert len(word_document.comments) == 2
    assert_docx_structure(
        docx_path,
        required_paragraphs=(
            "OODocs User Guide",
            "Contents",
            "List of Tables",
            "List of Figures",
            "Comments",
            "References",
        ),
        table_count=27,
        inline_shape_count=11,
        comment_count=2,
    )
    assert next(paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text == "Comments") == "Heading 2"
    assert next(paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text == "References") == "Heading 2"

    with zipfile.ZipFile(docx_path) as archive:
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
    assert "DOCX uses page footnotes by default." in footnotes_xml
    assert "Portable footnotes are authored inline" in footnotes_xml
    assert all("ListOfComments() collects these review notes onto a dedicated generated page." in "\n".join(p.text for p in comment.paragraphs) or "This note will show up again on the generated comments page." in "\n".join(p.text for p in comment.paragraphs) for comment in word_document.comments)

    assert "OODocs User Guide" in pdf_text
    assert "Contents" in pdf_text
    assert "List of Tables" in pdf_text
    assert "List of Figures" in pdf_text
    assert "Comments" in pdf_text
    assert "References" in pdf_text
    assert "Guide Cover" in pdf_text
    assert "AuthorLayout(mode='stacked')" in pdf_text
    assert "Theme(blocks=BlockDefaults(footnote_placement='document'))" in pdf_text
    assert "PageMargins" in pdf_text
    assert "PageBreak()" in pdf_text
    assert "settings.get_text_width(0.75)" in pdf_text
    assert "placement='inline'" in pdf_text
    assert "Inline image example:" in pdf_text
    assert "Compact inline chips cover categories" in pdf_text
    assert "Ctrl+Enter" in pdf_text
    assert "prescript(...)" in pdf_text
    assert "Inline text can carry ordinary scripts such as H2O and x2" in pdf_text
    assert "READY" in pdf_text
    assert "Pygments highlighting" in pdf_text
    assert "JAVASCRIPT" in pdf_text
    assert "SQL" in pdf_text
    assert "YAML" in pdf_text
    assert "CodeBlock.from_file(...)" in pdf_text
    assert "line_numbers" in pdf_text
    assert "highlight_lines" in pdf_text
    assert "ListOfListings" in pdf_text
    assert "Table(split=False)" in pdf_text
    assert "Table(split=True)" in pdf_text
    assert "continuation_label" in pdf_text
    assert "continued_caption_template" in pdf_text
    assert "DOCX and PDF use repeated header rows" in pdf_text
    assert "plain flow policy" in pdf_text
    assert "display: table-header-group" in pdf_text
    assert "CompactTable.grouped_headers" in pdf_text
    assert "TableCell(rowspan=...)" in pdf_text
    assert "TableCell(colspan=...)" in pdf_text
    assert "Table.grouped_headers(...)" in pdf_text
    assert "rowspan=2" in pdf_text
    assert 'style="booktabs"' in pdf_text
    assert "top_rule" in pdf_text
    assert "vertical grid" in pdf_text
    assert "ColumnSpec(key=" in pdf_text
    assert "visible=False" in pdf_text
    assert "full_matrix.excerpt" in pdf_text
    assert "full_matrix.save_csv" in pdf_text
    assert "record-column visibility" in pdf_text
    assert ("'tbp'" in pdf_text) or ('placement="tbp"' in pdf_text)
    assert "TableOfContents" in pdf_text
    assert "TableOfContentsLevelStyle" in pdf_text
    assert "Scoped lists" in pdf_text
    assert "scope='part'" in pdf_text
    assert "scope='chapter'" in pdf_text
    assert "scope='section'" in pdf_text
    assert "Appendix child chapters use A, B, C" in pdf_text
    assert "Table and figure captions keep" in pdf_text
    assert "document-wide numbering" in pdf_text
    assert "A reading map for the guide." in pdf_text
    assert "Example catalog" in pdf_text
    assert "api_objects_example" in pdf_text
    assert "style_cleanup_smoke" in pdf_text
    assert "LaTeX habits translated into oodocs's Python-first authoring model." in pdf_text
    assert "For authors coming from LaTeX" in pdf_text
    assert "Page layout controls shared across renderers." in pdf_text
    assert "Table-of-contents defaults and customization options." in pdf_text
    assert "Document-level configuration options." in pdf_text
    assert "Grouped Theme defaults" in pdf_text
    assert "Block-level option scope" in pdf_text
    assert "HeadingStyle" in pdf_text
    assert "heading_styles" in pdf_text
    assert "heading_style" in pdf_text
    assert "Component presets wrap ordinary blocks" in pdf_text
    assert "Template presets build full Document objects" in pdf_text
    assert "oodocs.presets.components" in pdf_text
    assert "oodocs.presets.templates" in pdf_text
    assert "JournalArticleTemplate" in pdf_text
    assert "Your Paper Your Way" in pdf_text
    assert "Instructions for Authors" in pdf_text
    assert "6.4.1 Subsection entries" in pdf_text
    assert "6.4.1.1 Fourth-level section entries" in pdf_text
    assert "Figure sizing patterns for width, height, and document-relative sizing." in pdf_text
    assert "CropBox" in pdf_text
    assert "rotation=90" in pdf_text
    assert "alt_text" in pdf_text
    assert "Coordinate-based drawings can be page overlays or inline flow objects." in pdf_text
    assert "Advanced table and figure placement controls." in pdf_text
    assert "SubFigure" in pdf_text
    assert "SubFigureGroup" in pdf_text
    assert "SubTable" in pdf_text
    assert "SubTableGroup" in pdf_text
    assert "Table 1(a)" in pdf_text
    assert "PdfPages" in pdf_text
    assert "link-style placeholder" in pdf_text
    assert "OODocs Contributor Certificate" in pdf_text
    assert "Renderer-specific behavior for notes, review workflows, and cross-reference stability." in pdf_text
    assert "Numbered statements, proofs, and custom counters" in pdf_text
    assert "CountableBlock" in pdf_text
    assert "Theorem 3. Stable references" in pdf_text
    assert "Exercise 1." in pdf_text
    assert "Build and validate from the CLI" in pdf_text
    assert "oodocs build report.py --out artifacts" in pdf_text
    assert "Validation results are structured objects" in pdf_text
    assert "Command-line builds and validation all call the same high-level workflow API." in pdf_text
    assert "Portable footnotes are authored inline" in pdf_text
    assert "github.com/Gonie-Gonie/oo-docs" in pdf_text
    assert "Document.from_markdown(...)" in pdf_text
    assert "parse_markdown(...)" in pdf_text
    assert "Document.from_notebook(...)" in pdf_text
    assert "parse_notebook(...)" in pdf_text
    assert "Notebook-backed report" in pdf_text
    assert "Release note digest" in pdf_text
    assert "Footnotes" in pdf_text
    assert len(pdf_reader.pages) >= 14
    assert _pdf_image_draw_count(pdf_path) == 7
    assert_pdf_text_and_pages(
        pdf_path,
        required_text=(
            "OODocs User Guide",
            "Contents",
            "List of Tables",
            "List of Figures",
            "References",
        ),
        min_pages=14,
    )

    assert "OODocs User Guide" in normalized_html_text
    assert "Guide Cover" in normalized_html_text
    assert "List of Tables" in normalized_html_text
    assert "List of Figures" in normalized_html_text
    assert "Comments" in normalized_html_text
    assert "References" in normalized_html_text
    assert "AuthorLayout(mode='stacked')" in normalized_html_text
    assert "Theme(blocks=BlockDefaults(footnote_placement='document'))" in normalized_html_text
    assert "PageMargins" in normalized_html_text
    assert "PageBreak()" in normalized_html_text
    assert "settings.get_text_width(0.75)" in normalized_html_text
    assert "placement='inline'" in normalized_html_text
    assert "Inline image example:" in normalized_html_text
    assert "Compact inline chips cover categories" in normalized_html_text
    assert "Ctrl+Enter" in normalized_html_text
    assert "prescript(...)" in normalized_html_text
    assert "Inline text can carry ordinary scripts such as H 2 O and x 2" in normalized_html_text
    assert "READY" in normalized_html_text
    assert "Pygments highlighting" in normalized_html_text
    assert "JAVASCRIPT" in normalized_html_text
    assert "SQL" in normalized_html_text
    assert "YAML" in normalized_html_text
    assert "CodeBlock.from_file(...)" in normalized_html_text
    assert "line_numbers" in normalized_html_text
    assert "highlight_lines" in normalized_html_text
    assert "ListOfListings" in normalized_html_text
    assert "Table(split=False)" in normalized_html_text
    assert "Table(split=True)" in normalized_html_text
    assert "continuation_label" in normalized_html_text
    assert "continued_caption_template" in normalized_html_text
    assert "DOCX and PDF use repeated header rows" in normalized_html_text
    assert "plain flow policy" in normalized_html_text
    assert "display: table-header-group" in normalized_html_text
    assert "CompactTable" in normalized_html_text
    assert "grouped_headers" in normalized_html_text
    assert "rowspan" in normalized_html_text
    assert "colspan" in normalized_html_text
    assert 'style="booktabs"' in normalized_html_text
    assert "top_rule, header_rule, bottom_rule" in normalized_html_text
    assert "publication-style horizontal rules without vertical grid lines" in normalized_html_text
    assert "ColumnSpec" in normalized_html_text
    assert "visible = False" in normalized_html_text
    assert "full_matrix . excerpt" in normalized_html_text
    assert "save_csv" in normalized_html_text
    assert "record-column visibility" in normalized_html_text
    assert ("'tbp'" in normalized_html_text) or ('placement="tbp"' in normalized_html_text)
    assert "TableOfContents" in normalized_html_text
    assert "TableOfContentsLevelStyle" in normalized_html_text
    assert "Scoped lists" in normalized_html_text
    assert "scope='part'" in normalized_html_text
    assert "scope='chapter'" in normalized_html_text
    assert "scope='section'" in normalized_html_text
    assert "HTML generated-list page numbers" in normalized_html_text
    assert "show_page_numbers=False" in normalized_html_text
    assert "Appendix child chapters use A, B, C numbering" in normalized_html_text
    assert "Table and figure captions keep document-wide numbering." in normalized_html_text
    assert "validation warns if main chapters continue afterward." in normalized_html_text
    assert "ListOfComments() collects these review notes onto a dedicated generated page." in normalized_html_text
    assert "Example catalog" in normalized_html_text
    assert "api_objects_example" in normalized_html_text
    assert "style_cleanup_smoke" in normalized_html_text
    assert "LaTeX habits translated into oodocs's Python-first authoring model." in normalized_html_text
    assert "For authors coming from LaTeX" in normalized_html_text
    assert "Page layout controls shared across renderers." in normalized_html_text
    assert "Table-of-contents defaults and customization options." in normalized_html_text
    assert "Document-level configuration options." in normalized_html_text
    assert "Grouped Theme defaults" in normalized_html_text
    assert "Block-level option scope" in normalized_html_text
    assert "HeadingStyle" in normalized_html_text
    assert "heading_styles" in normalized_html_text
    assert "heading_style" in normalized_html_text
    assert "Component presets wrap ordinary blocks" in normalized_html_text
    assert "Template presets build full Document objects" in normalized_html_text
    assert "oodocs.presets.components" in normalized_html_text
    assert "oodocs.presets.templates" in normalized_html_text
    assert "JournalArticleTemplate" in normalized_html_text
    assert "Your Paper Your Way" in normalized_html_text
    assert "Instructions for Authors" in normalized_html_text
    assert "6.4.1 Subsection entries" in normalized_html_text
    assert "6.4.1.1 Fourth-level section entries" in normalized_html_text
    assert "Figure sizing patterns for width, height, and document-relative sizing." in normalized_html_text
    assert "CropBox" in normalized_html_text
    assert "rotation" in normalized_html_text
    assert "alt_text" in normalized_html_text
    assert "Coordinate-based drawings can be page overlays or inline flow objects." in normalized_html_text
    assert "Advanced table and figure placement controls." in normalized_html_text
    assert "SubFigure" in normalized_html_text
    assert "SubFigureGroup" in normalized_html_text
    assert "SubTable" in normalized_html_text
    assert "SubTableGroup" in normalized_html_text
    assert "Table 1(a)" in normalized_html_text
    assert "PdfPages" in normalized_html_text
    assert "link-style placeholder" in normalized_html_text
    assert "OODocs Contributor Certificate" in normalized_html_text
    assert "Numbered statements, proofs, and custom counters" in normalized_html_text
    assert "CountableBlock" in normalized_html_text
    assert "Theorem 3. Stable references" in normalized_html_text
    assert "Exercise 1." in normalized_html_text
    assert "Build and validate from the CLI" in normalized_html_text
    assert "oodocs build report.py --out artifacts" in normalized_html_text
    assert "Validation results are structured objects" in normalized_html_text
    assert "Command-line builds and validation all call the same high-level workflow API." in normalized_html_text
    assert "Portable footnotes are authored inline" in normalized_html_text
    assert "github.com/Gonie-Gonie/oo-docs" in normalized_html_text
    assert "Document.from_markdown(...)" in normalized_html_text
    assert "parse_markdown(...)" in normalized_html_text
    assert "Document.from_notebook(...)" in normalized_html_text
    assert "parse_notebook(...)" in normalized_html_text
    assert "Notebook-backed report" in normalized_html_text
    assert "Release note digest" in normalized_html_text
    assert "Footnotes" in normalized_html_text
    assert html_text.count("data:image/png;base64,") == 7
    assert 'href="#table_1"' in html_text
    assert 'href="#figure_1"' in html_text
    assert_html_internal_links_resolve(
        html_path,
        required_hrefs=("#table_1", "#figure_1"),
        required_text=("OODocs User Guide", "Build and validate from the CLI"),
    )
    assert 'class="oodocs-toc-entry oodocs-toc-entry-no-page oodocs-toc-entry-level-1"' in html_text
    assert 'class="oodocs-toc-entry oodocs-toc-entry-no-page oodocs-toc-entry-level-2"' in html_text
    assert 'class="oodocs-toc-entry oodocs-toc-entry-no-page oodocs-toc-entry-level-3"' in html_text
    assert 'class="oodocs-toc-entry oodocs-toc-entry-no-page oodocs-toc-entry-level-4"' in html_text


def test_usage_guide_example_supports_common_cli(tmp_path: Path, capsys) -> None:
    usage_guide = _load_example_module("usage_guide_example")
    output_dir = tmp_path / "cli"

    outputs = usage_guide.build_usage_guide(
        output_dir / "programmatic",
        output_formats=("html",),
    )
    assert set(outputs.keys()) == {"html"}
    assert outputs["html"].exists()

    usage_guide.main(
        [
            "--output-dir",
            str(output_dir),
            "--outputs",
            "html",
            "--quiet",
        ]
    )

    captured = capsys.readouterr()
    assert captured.out == ""
    assert (output_dir / "oodocs-user-guide.html").exists()
