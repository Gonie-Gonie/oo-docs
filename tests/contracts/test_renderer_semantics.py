"""Cross-renderer semantic contract built from one compact manual fixture.

The assertions intentionally target structure, anchors, visible text, links,
and table dimensions.  They avoid pixel-perfect snapshots so renderer styling
can evolve without weakening the document semantics contract.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
import re
from typing import Iterator
from zipfile import ZipFile

from docx import Document as WordDocument
from pypdf import PdfReader
import pytest

from oodocs import (
    Chapter,
    Document,
    DocumentSettings,
    Paragraph,
    Part,
    Section,
    SubFigure,
    SubFigureGroup,
    Table,
    Theme,
)
from oodocs.components.blocks import AlignedEquation, Appendix
from oodocs.components.descriptions import DescriptionList
from oodocs.components.matter import BackMatter, FrontMatter, MainMatter
from oodocs.layout.indexing import build_render_index
from oodocs.renderers import pdf as pdf_renderer_module
from tests.fixtures.manual_suite import (
    ASSET_NAMES,
    LONG_TABLE_ROW_COUNT,
    MANUAL_NAME,
    build_cli_application,
    build_schema_catalog,
    build_suite,
)


@dataclass
class _ParsedTable:
    rows: list[int] = field(default_factory=list)
    text_parts: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return " ".join(" ".join(self.text_parts).split())


class _SemanticTableParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.tables: list[_ParsedTable] = []
        self._active: _ParsedTable | None = None
        self._cell_count: int | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag == "table":
            assert self._active is None, "nested tables are outside this fixture contract"
            self._active = _ParsedTable()
        elif tag == "tr" and self._active is not None:
            self._cell_count = 0
        elif tag in {"th", "td"} and self._cell_count is not None:
            self._cell_count += 1

    def handle_endtag(self, tag: str) -> None:
        if tag == "tr" and self._active is not None and self._cell_count is not None:
            self._active.rows.append(self._cell_count)
            self._cell_count = None
        elif tag == "table" and self._active is not None:
            self.tables.append(self._active)
            self._active = None

    def handle_data(self, data: str) -> None:
        if self._active is not None and data.strip():
            self._active.text_parts.append(data)


@pytest.fixture(scope="module")
def manual_suite(tmp_path_factory: pytest.TempPathFactory):
    return build_suite(tmp_path_factory.mktemp("renderer-semantics"))


@pytest.fixture(scope="module")
def manual_document(manual_suite):
    return manual_suite.build(MANUAL_NAME)


@pytest.fixture(scope="module")
def html_path(manual_document, tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("renderer-html") / "manual.html"
    return manual_document.save_html(path)


@pytest.fixture(scope="module")
def docx_path(manual_document, tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("renderer-docx") / "manual.docx"
    return manual_document.save_docx(path)


@pytest.fixture(scope="module")
def pdf_path(manual_document, tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("renderer-pdf") / "manual.pdf"
    return manual_document.save_pdf(path)


def _walk_blocks(blocks: object) -> Iterator[object]:
    if not isinstance(blocks, (list, tuple)):
        blocks = (blocks,)
    for block in blocks:
        yield block
        children = getattr(block, "children", ())
        if isinstance(children, (list, tuple)):
            yield from _walk_blocks(children)


def _word_text(document: WordDocument) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(values)


def _word_xml(path: Path) -> tuple[str, str]:
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        all_word_xml = "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
        relationships = "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/_rels/") and name.endswith(".rels")
        )
    return document_xml, all_word_xml + "\n" + relationships


def _pdf_uri_targets(reader: PdfReader) -> set[str]:
    targets: set[str] = set()
    for page in reader.pages:
        for annotation_ref in page.get("/Annots", ()):
            annotation = annotation_ref.get_object()
            action = annotation.get("/A")
            if action is not None and action.get("/URI") is not None:
                targets.add(str(action["/URI"]))
    return targets


def _pdf_internal_destinations(reader: PdfReader) -> list[object]:
    destinations: list[object] = list(reader.named_destinations)
    for page in reader.pages:
        for annotation_ref in page.get("/Annots", ()):
            annotation = annotation_ref.get_object()
            if annotation.get("/Dest") is not None:
                destinations.append(annotation["/Dest"])
    return destinations


def test_manual_suite_model_is_compact_complete_and_valid(manual_suite, manual_document) -> None:
    assert tuple(manual_suite.context.variables) == ("organization", "version", "date")
    assert manual_suite.context.variables["organization"] == "Example Documentation Cooperative"
    assert all((manual_suite.context.root / "assets" / name).is_file() for name in ASSET_NAMES)
    assert {manual_suite.context.citations.resolve(key).entry_type for key in (
        "article2026", "thesis2024", "manual2025", "standard2023"
    )} == {"article", "phdthesis", "manual", "standard"}

    front, main, back = manual_document.body.children
    assert isinstance(front, FrontMatter)
    assert isinstance(main, MainMatter)
    assert isinstance(back, BackMatter)
    assert len(manual_document.matter_layout().front.children) == 4
    assert isinstance(main.children[0], Part)
    assert any(isinstance(block, Appendix) for block in back.children)
    assert manual_document.settings.title_matter.cover is not None
    assert manual_document.settings.theme.resolve_language_tag() == "ko-KR"

    blocks = list(_walk_blocks(manual_document.body.children))
    assert any(isinstance(block, Chapter) for block in blocks)
    assert any(isinstance(block, Section) for block in blocks)
    assert any(isinstance(block, DescriptionList) for block in blocks)
    subfigure_group = next(block for block in blocks if isinstance(block, SubFigureGroup))
    assert len(subfigure_group.subfigures) == 2
    assert all(isinstance(subfigure, SubFigure) for subfigure in subfigure_group.subfigures)
    assert any(
        isinstance(block, Table) and block.identifier == "grouped-header-table"
        for block in blocks
    )
    assert any(isinstance(block, Table) and block.identifier == "wide-table" for block in blocks)
    long_table = next(
        block for block in blocks
        if isinstance(block, Table) and block.identifier == "long-table"
    )
    assert len(long_table.rows) == LONG_TABLE_ROW_COUNT
    assert long_table._resolve_split() is True
    assert long_table.style.repeat_header_rows is True

    aligned = next(block for block in blocks if isinstance(block, AlignedEquation))
    render_index = build_render_index(manual_document)
    assert [render_index.anchor_for(line) for line in aligned.lines] == [
        "equation-total-line",
        "equation-rate-line",
    ]
    assert [render_index.equation_number(line) for line in aligned.lines] == [1, 2]

    page_numbers = manual_document.settings.theme.page_numbers
    assert page_numbers.show_page_numbers is True
    assert manual_document.settings.theme.format_page_number(4, matter="front") == "iv"
    assert manual_document.settings.theme.format_page_number(4, matter="main") == "4"
    assert manual_document.settings.theme.format_page_number(4, matter="back") == "4"
    assert page_numbers.restart_main_matter is True
    assert page_numbers.restart_back_matter is False
    assert not manual_document.validate().errors

    cli = build_cli_application()
    assert cli.root_command.options[0].required is True
    assert cli.root_command.subcommands[0].name == "inspect"
    schema = build_schema_catalog()
    assert schema.schema_for("manual-root").fields[0].target_schema == "manual-child"
    assert schema.schema_for("manual-child").fields[0].target_schema == "manual-root"
    assert schema.validate().ok


def test_html_preserves_manual_semantics_by_default(html_path: Path) -> None:
    html = html_path.read_text(encoding="utf-8")

    assert '<html lang="ko-KR">' in html
    assert html.count("oodocs-cover-page") >= 1
    assert "data:image/png;base64," in html
    cover_start = html.index('<section class="oodocs-title-matter oodocs-cover-page')
    front_start = html.index('<section class="oodocs-front-matter')
    main_start = html.index('<section class="oodocs-main-matter')
    back_start = html.index('<section class="oodocs-back-matter')
    assert cover_start < front_start < main_start < back_start
    assert "oodocs-page-break-before" in html

    assert 'id="manual-overview"' in html
    assert 'id="manual-navigation"' in html
    assert 'href="#manual-navigation">open the navigation target</a>' in html
    assert 'href="#manual-overview">return to the overview</a>' in html
    assert re.search(r'href="#manual-navigation">1\.2절</a>', html)
    assert 'href="#schema-manual-child">manual-child</a>' in html
    assert 'href="#schema-manual-root">manual-root</a>' in html

    assert '<dl class="oodocs-description-list' in html
    assert '<dt class="oodocs-description-term"' in html
    assert '<dd class="oodocs-description-definition"' in html
    assert 'data-numbering="each"' in html
    assert '<tr id="equation-total-line" class="oodocs-equation-row">' in html
    assert '<tr id="equation-rate-line" class="oodocs-equation-row">' in html
    assert 'href="#equation-total-line">Equation 1</a>' in html
    assert 'href="#equation-rate-line">Equation 2</a>' in html

    assert 'href="https://doi.org/10.5555/oodocs.acceptance"' in html
    assert 'href="https://example.org/articles/portable-semantics"' in html
    assert 'class="oodocs-section oodocs-section-level-2 oodocs-section-page-layout"' in html
    assert html.count('class="oodocs-subfigure"') == 2
    assert "doc-tool [OPTIONS] SOURCE" in html

    parser = _SemanticTableParser()
    parser.feed(html)
    long_table = next(table for table in parser.tables if "row-042" in table.text)
    assert len(long_table.rows) == LONG_TABLE_ROW_COUNT + 1
    assert long_table.rows == [3] * (LONG_TABLE_ROW_COUNT + 1)
    grouped_table = next(table for table in parser.tables if "Measured range" in table.text)
    assert grouped_table.rows == [2, 3, 3, 3]


@pytest.mark.render
def test_docx_preserves_bookmarks_page_sequence_links_and_table_shape(docx_path: Path) -> None:
    word = WordDocument(docx_path)
    text = _word_text(word)
    document_xml, all_word_xml = _word_xml(docx_path)

    assert len(word.sections) >= 5
    assert len(word.inline_shapes) >= 3
    assert "공용 기술 설명서" in text
    assert "Back-matter appendix marker." in text
    assert text.index("머리말") < text.index("Portable manual") < text.index("Acceptance notes")

    for anchor in (
        "manual-overview",
        "manual-navigation",
        "schema-manual-root",
        "schema-manual-child",
        "equation-total-line",
        "equation-rate-line",
    ):
        assert f'w:name="{anchor}"' in document_xml
        assert f'w:anchor="{anchor}"' in document_xml

    assert "1.2절" in text
    assert "Equation 1" in text and "Equation 2" in text
    assert "Q_total" in text and "1,234.5 kWh" in text
    assert "https://doi.org/10.5555/oodocs.acceptance" in all_word_xml
    assert "https://example.org/articles/portable-semantics" in all_word_xml

    long_table = next(
        table for table in word.tables
        if table.rows and "Long Table Header" in table.rows[0].cells[0].text
    )
    assert len(long_table.rows) == LONG_TABLE_ROW_COUNT + 1
    assert len(long_table.columns) == 3
    assert "<w:tblHeader" in document_xml

    assert 'w:fmt="lowerRoman"' in all_word_xml
    assert 'w:fmt="decimal"' in all_word_xml
    assert all_word_xml.count('w:start="1"') >= 2


@pytest.mark.render
def test_pdf_preserves_destinations_visible_text_links_and_long_table_header(
    pdf_path: Path,
) -> None:
    reader = PdfReader(pdf_path)
    page_text = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(page_text)
    normalized = " ".join(text.split())

    assert len(reader.pages) >= 8
    assert "공용 기술 설명서" in page_text[0]
    assert "Object link forward" not in page_text[0]
    assert "Object link forward" in normalized
    assert "Back-matter appendix marker." in normalized
    assert "Equation 1" in normalized and "Equation 2" in normalized
    assert "Manual root schema" in normalized and "Manual child schema" in normalized
    assert "Q_total" in normalized or "Qtotal" in normalized
    assert "row-001" in normalized and "row-042" in normalized
    assert normalized.count("Long Table Header") >= 2

    uri_targets = _pdf_uri_targets(reader)
    assert "https://doi.org/10.5555/oodocs.acceptance" in uri_targets
    assert "https://example.org/articles/portable-semantics" in uri_targets
    assert len(_pdf_internal_destinations(reader)) >= 8

    preface_page = next(value for value in page_text if "머리말" in value)
    main_page = next(
        value for value in page_text
        if value.startswith("1\n") and "Portable manual" in value
    )
    back_page = next(value for value in page_text if "Back-matter appendix marker." in value)
    # The unnumbered cover still occupies physical page one, so front matter
    # begins at ii; main matter restarts at 1 and back matter continues it.
    assert re.search(r"(?:^|\n)ii(?:\n|$)", preface_page)
    assert re.search(r"(?:^|\n)1(?:\n|$)", main_page)
    assert re.search(r"(?:^|\n)10(?:\n|$)", back_page)


@pytest.mark.render
def test_pdf_builtin_cid_fallback_preserves_unicode_without_system_fonts(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(pdf_renderer_module, "SYSTEM_FONT_VARIANTS", {})
    renderer = pdf_renderer_module.PdfRenderer()
    assert renderer._resolve_font("Times New Roman", False, False) == "Times-Roman"

    default_document = Document(
        "Unicode",
        Paragraph("ASCII body. ", "400 CO₂"),
    )
    korean_document = Document(
        "한국어",
        Paragraph("공용 기술 설명서. See 1장 and 1.1절."),
        settings=DocumentSettings(theme=Theme.from_locale("ko-KR")),
    )

    default_path = default_document.save_pdf(tmp_path / "default-unicode.pdf")
    korean_path = korean_document.save_pdf(tmp_path / "korean-unicode.pdf")
    default_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(default_path).pages
    )
    korean_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(korean_path).pages
    )

    assert "400 CO₂" in default_text
    assert "공용 기술 설명서" in korean_text
    assert "See 1장 and 1.1절." in korean_text
