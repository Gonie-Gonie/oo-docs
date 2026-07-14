"""DOCX renderer.

Attributes:
    ALIGNMENTS: Mapping from OODocs paragraph alignment names to DOCX values.
    TABLE_ALIGNMENTS: Mapping from OODocs table alignment names to DOCX values.
    CELL_VERTICAL_ALIGNMENTS: Mapping from OODocs cell vertical alignment names
        to DOCX values.
    DEFAULT_FOOTNOTES_XML: Minimal footnotes part XML used when creating native
        DOCX footnotes.
"""

from __future__ import annotations

from dataclasses import replace
from io import BytesIO
from pathlib import Path
import re
from xml.sax.saxutils import escape as xml_escape

from docx import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_BREAK
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.parts.story import StoryPart
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image, ImageDraw, ImageFont

from oodocs.components.blocks import (
    AlignedEquation,
    Box,
    BulletList,
    CodeBlock,
    ColumnSpan,
    CountableBlock,
    Divider,
    Equation,
    MultiColumn,
    NumberedList,
    PageBreak,
    Paragraph,
    Part,
    Section,
    VerticalSpace,
)
from oodocs.components.generated import (
    ListOfComments,
    ListOfGlossaryTerms,
    ListOfAlgorithms,
    ListOfFigures,
    ListOfFootnotes,
    ListOfListings,
    ListOfReferences,
    ListOfTables,
    TableOfContents,
)
from oodocs.components.inline import (
    _BlockReference,
    Citation,
    Comment,
    Footnote,
    Hyperlink,
    InlineChip,
    LineBreak,
    Math,
    ReferenceGroup,
    Text,
)
from oodocs.components.descriptions import DescriptionList
from oodocs.components.media import (
    Figure,
    PdfPages,
    SubFigure,
    SubFigureGroup,
    SubTable,
    SubTableGroup,
    Table,
    build_table_layout,
    image_source_to_buffer,
    processed_image_source_to_buffer,
)
from oodocs.components.people import AuthorTitleLine
from oodocs.components.positioning import (
    ImageBox,
    PositionedBox,
    PositionedItem,
    Shape,
    TextBox,
    resolve_positioned_boxes,
)
from oodocs.components.references import format_citation_label, reference_entry_marker
from oodocs.document import Document
from oodocs.components.equations import SUBSCRIPT, SUPERSCRIPT, parse_latex_segments
from oodocs.core import OODocsError, PathLike, length_to_inches
from oodocs.styles.generated import TableOfContentsLevelStyle
from oodocs.layout.indexing import (
    ReferenceTextPiece,
    RenderIndex,
    build_render_index,
    reference_text_pieces,
    resolve_block_reference,
)
from oodocs.settings import PageLayout
from oodocs.styles import BoxStyle, DescriptionListStyle, HeadingStyle, ListStyle, ParagraphStyle, TableStyle, TextStyle, Theme
from oodocs.renderers.context import DocxRenderContext, _settings_with_page_layout
from oodocs.renderers.syntax import SyntaxToken, _syntax_line_tokens, syntax_tokens


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

TABLE_ALIGNMENTS = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}

CELL_VERTICAL_ALIGNMENTS = {
    "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
    "middle": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
}

DEFAULT_FOOTNOTES_XML = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p>
      <w:r>
        <w:separator />
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p>
      <w:r>
        <w:continuationSeparator />
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>
"""


class FootnotesPart(StoryPart):
    """Container part for native DOCX footnotes.

    Attributes:
        package: DOCX package that owns the generated footnotes part.
    """

    @classmethod
    def default(cls, package: object) -> "FootnotesPart":
        """Create a default DOCX footnotes part.

        Args:
            package: DOCX package that will own the part.

        Returns:
            Initialized footnotes part containing Word's required separator
            entries.
        """

        return cls(
            PackURI("/word/footnotes.xml"),
            CT.WML_FOOTNOTES,
            parse_xml(DEFAULT_FOOTNOTES_XML),
            package,
        )

    def add_footnote_paragraph(self, footnote_id: int) -> DocxParagraph:
        """Add an empty native footnote paragraph.

        Args:
            footnote_id: Numeric Word footnote identifier.

        Returns:
            Paragraph object that can receive footnote runs.
        """

        footnote = OxmlElement("w:footnote")
        footnote.set(qn("w:id"), str(footnote_id))
        paragraph = OxmlElement("w:p")
        footnote.append(paragraph)
        self._element.append(footnote)
        return DocxParagraph(paragraph, self)


class DocxRenderer:
    """Render OODocs documents into DOCX files.

    The renderer exposes ``render_*`` methods so block classes and custom
    extensions can dispatch DOCX rendering through a shared context.

    Attributes:
        _bookmark_id: Internal per-render bookmark counter.
        _native_footnotes_part: Internal DOCX footnotes part when native
            footnotes are emitted.
    """

    def render(self, document: Document, output_path: PathLike) -> Path:
        """Render an OODocs document to a DOCX file.

        Args:
            document: Document to render.
            output_path: Destination ``.docx`` path.

        Returns:
            Resolved output path that was written.
        """

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        word_document = WordDocument()
        self._initialized_cells: set[int] = set()
        self._bookmark_id = 1
        self._native_footnotes_part: FootnotesPart | None = None
        self._rendered_native_footnotes: set[int] = set()
        self._native_footnotes_enabled = True
        render_index = build_render_index(document)
        self._configure_document(word_document, document)
        settings = document.settings
        self._native_footnotes_enabled = self._should_use_native_footnotes(
            settings.theme,
            render_index,
        )
        context = DocxRenderContext(
            theme=settings.theme,
            render_index=render_index,
            settings=settings,
            unit=settings.unit,
            word_document=word_document,
        )
        matter = document.matter_layout()
        front_children = matter.front.children
        main_children = matter.main.children
        back_children = matter.back.children
        title_matter = settings.title_matter
        has_front_matter = title_matter.cover is not None or bool(front_children)
        matter_sections: dict[str, int] = {"front": 0}
        self._render_page_items(
            word_document,
            document,
            context,
            has_front_matter=has_front_matter,
        )

        self._render_title_matter(
            word_document,
            document,
            context,
        )

        if title_matter.cover is not None and (
            front_children or matter.front.page_break_before
        ):
            self._ensure_page_break(word_document)

        if front_children:
            self._render_top_level_children(word_document, front_children, context)
        if main_children:
            if has_front_matter or matter.main.page_break_before:
                section = word_document.add_section(WD_SECTION.NEW_PAGE)
                self._configure_section_page_box(section, settings)
                self._render_main_page_items(section, document, context)
                matter_sections["main"] = len(word_document.sections) - 1
            else:
                matter_sections["main"] = 0
            self._render_top_level_children(word_document, main_children, context)

        if back_children:
            section = word_document.add_section(
                WD_SECTION.NEW_PAGE
                if matter.back.page_break_before
                else WD_SECTION.CONTINUOUS
            )
            self._configure_section_page_box(section, settings)
            self._render_main_page_items(section, document, context)
            matter_sections["back"] = len(word_document.sections) - 1
            self._render_top_level_children(word_document, back_children, context)

        if self._should_auto_render_footnote_list(document, render_index):
            self.render_footnote_list(ListOfFootnotes(), context)

        if settings.theme.uses_header_footer():
            self._configure_header_footer_sections(
                word_document,
                document,
                settings.theme,
                has_front_matter=has_front_matter,
                has_main_matter=bool(main_children) or not has_front_matter,
                has_back_matter=bool(back_children),
                matter_sections=matter_sections,
            )

        word_document.save(path)
        return path

    def add_heading(
        self,
        container: object,
        title: list[Text],
        level: int,
        context: DocxRenderContext,
        *,
        number_label: str | None = None,
        anchor: str | None = None,
        toc: bool = False,
        heading_style: HeadingStyle | None = None,
    ) -> None:
        """Render a heading into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            title: Inline title fragments.
            level: Heading level to render.
            context: Current DOCX render context.
            number_label: Optional generated heading number.
            anchor: Optional bookmark anchor.
            toc: Whether to append a Word table-of-contents entry.
            heading_style: Optional direct heading style override.
        """

        self._add_heading(
            container,
            title,
            level,
            context.theme,
            number_label=number_label,
            anchor=anchor,
            render_index=context.render_index,
            toc=toc,
            heading_style=heading_style,
        )

    def render_part(
        self,
        container: object,
        block: Part,
        context: DocxRenderContext,
    ) -> None:
        """Render a part separator page and its child blocks into DOCX.

        Args:
            container: Active DOCX document container.
            block: Part block to render.
            context: Current DOCX render context.
        """

        self._assert_document_container(container, "Part")
        word_document = context.word_document
        if word_document.paragraphs or word_document.tables:
            self._ensure_page_break(word_document)

        number_label = context.render_index.heading_number(block) if block.numbered else None
        anchor = context.render_index.heading_anchor(block)
        if number_label:
            self._add_title_line(
                word_document,
                [Text(number_label)],
                font_size=max(context.theme.typography.body_font_size + 3, 14),
                alignment="center",
                bold=True,
                space_after=18,
                anchor=anchor,
            )
            anchor = None
        title_paragraph = self._add_title_line(
            word_document,
            block.title,
            font_size=max(context.theme.typography.title_font_size, context.theme.resolve_heading_size(1) + 2),
            alignment="center",
            bold=True,
            space_after=0,
            anchor=anchor,
        )
        if block.toc:
            self._append_toc_entry_field(
                title_paragraph,
                self._flatten_fragments(
                    self._heading_fragments(block.title, number_label),
                    context.theme,
                    context.render_index,
                ),
                level=block.level,
            )

        if block.children:
            self._ensure_page_break(word_document)
            self._render_docx_flow_children(
                word_document,
                block.children,
                context,
                restore_at_end=True,
            )

    def render_section(
        self,
        container: object,
        block: Section,
        context: DocxRenderContext,
    ) -> None:
        """Render a section heading and child blocks into DOCX.

        Args:
            container: Active DOCX document container.
            block: Section block to render.
            context: Current DOCX render context.
        """

        section_context = context
        if block.page_layout is not None:
            section_context = self._ensure_docx_page_layout(
                container,
                context,
                block.page_layout,
            )
        self.add_heading(
            container,
            block.title,
            block.level,
            section_context,
            number_label=(
                section_context.render_index.heading_number(block)
                if block.numbered
                else None
            ),
            anchor=section_context.render_index.heading_anchor(block),
            toc=block.toc,
            heading_style=block.heading_style,
        )
        child_context = section_context
        if block.run_in_title_style is not None:
            child_context = replace(
                section_context,
                run_in_title_style=block.run_in_title_style,
            )
        self._render_docx_flow_children(
            container,
            block.children,
            child_context,
            restore_at_end=True,
        )

    def render_paragraph(
        self,
        container: object,
        paragraph_block: Paragraph,
        context: DocxRenderContext,
    ) -> None:
        """Render a paragraph block into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            paragraph_block: Paragraph block to render.
            context: Current DOCX render context.
        """

        paragraph = self._add_paragraph(container)
        paragraph_style = context.stylesheet.resolve(
            "paragraph",
            paragraph_block.style,
            ParagraphStyle(),
        )
        title_style = context.theme.resolve_run_in_title_style(
            paragraph_block.title_style,
            context.run_in_title_style,
        )
        self._apply_paragraph_style(
            paragraph,
            paragraph_style,
            context.theme,
            context.unit,
        )
        self._append_runs(
            paragraph,
            paragraph_block.render_content(title_style),
            theme=context.theme,
            render_index=context.render_index,
            word_document=context.word_document,
            unit=context.unit,
        )
        anchor = context.render_index.block_anchor(paragraph_block)
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)

    def render_list(
        self,
        container: object,
        list_block: BulletList | NumberedList,
        context: DocxRenderContext,
    ) -> None:
        """Render a list block into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            list_block: Bullet or numbered list block to render.
            context: Current DOCX render context.
        """

        self._render_list(
            container,
            list_block,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
        )

    def render_description_list(
        self,
        container: object,
        block: DescriptionList,
        context: DocxRenderContext,
    ) -> None:
        """Render a description list as a stable borderless DOCX table."""

        if not block.items:
            return
        style = context.stylesheet.resolve(
            "description_list",
            block.style,
            DescriptionListStyle(),
        )
        table = container.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = style.layout != "hanging"
        available_width = context.settings.text_width_in_inches()
        term_width = style.term_width_in_inches(context.unit)
        if term_width is None:
            term_width = min(max(available_width * 0.25, 0.75), 2.0)
        term_width = min(term_width, max(available_width * 0.5, 0.25))
        gap_points = style.term_gap_in_inches(context.unit) * 72.0

        for item in block.items:
            if style.layout == "stacked":
                term_row = table.add_row()
                term_cell = term_row.cells[0].merge(term_row.cells[1])
                self._prepare_description_cell(
                    term_cell,
                    bottom=0.0,
                    left=0.0,
                    right=0.0,
                )
                self._render_description_term(term_cell, item.term, style, context)

                definition_row = table.add_row()
                definition_cell = definition_row.cells[0].merge(definition_row.cells[1])
                self._prepare_description_cell(
                    definition_cell,
                    bottom=style.item_spacing,
                    left=gap_points,
                    right=0.0,
                )
                self._render_description_children(
                    definition_cell,
                    item.children,
                    style.definition_style,
                    context,
                )
                if not definition_cell.paragraphs:
                    paragraph = self._add_paragraph(definition_cell)
                    self._apply_paragraph_style(
                        paragraph,
                        style.definition_style,
                        context.theme,
                        context.unit,
                    )
                continue

            row = table.add_row()
            term_cell, definition_cell = row.cells
            self._prepare_description_cell(
                term_cell,
                bottom=style.item_spacing,
                left=0.0,
                right=gap_points,
            )
            self._prepare_description_cell(
                definition_cell,
                bottom=style.item_spacing,
                left=0.0,
                right=0.0,
            )
            term_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            definition_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if style.layout == "hanging":
                table.columns[0].width = Inches(term_width)
                table.columns[1].width = Inches(max(available_width - term_width, 0.25))
                self._set_cell_width(term_cell, term_width)
                self._set_cell_width(
                    definition_cell,
                    max(available_width - term_width, 0.25),
                )
            self._render_description_term(term_cell, item.term, style, context)
            self._render_description_children(
                definition_cell,
                item.children,
                style.definition_style,
                context,
            )
            if not definition_cell.paragraphs:
                paragraph = self._add_paragraph(definition_cell)
                self._apply_paragraph_style(
                    paragraph,
                    style.definition_style,
                    context.theme,
                    context.unit,
                )

    def _prepare_description_cell(
        self,
        cell: object,
        *,
        bottom: float,
        left: float,
        right: float,
    ) -> None:
        cell._tc.clear_content()
        self._initialized_cells.discard(id(cell))
        self._set_cell_borders_none(cell)
        self._set_cell_margins(cell, 0.0, right, bottom, left)

    def _render_description_term(
        self,
        cell: object,
        fragments: list[Text],
        style: object,
        context: DocxRenderContext,
    ) -> None:
        paragraph = self._add_paragraph(cell)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        self._append_runs(
            paragraph,
            fragments,
            default_size=context.theme.typography.body_font_size,
            default_style=style.term_text_style,
            theme=context.theme,
            render_index=context.render_index,
            word_document=context.word_document,
            unit=context.unit,
        )

    def _render_description_children(
        self,
        container: object,
        children: list[object],
        definition_style: ParagraphStyle,
        context: DocxRenderContext,
    ) -> None:
        for child in children:
            original_style = None
            if (
                isinstance(child, Paragraph)
                and isinstance(child.style, ParagraphStyle)
                and child.style == ParagraphStyle()
            ):
                original_style = child.style
                child.style = definition_style
            try:
                self._render_block(container, child, context)
            finally:
                if original_style is not None:
                    child.style = original_style

    def render_code_block(
        self,
        container: object,
        code_block: CodeBlock,
        context: DocxRenderContext,
    ) -> None:
        """Render a code block into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            code_block: Code block to render.
            context: Current DOCX render context.
        """

        self._render_code_block(
            container,
            code_block,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
        )

    def render_equation(
        self,
        container: object,
        equation: Equation,
        context: DocxRenderContext,
    ) -> None:
        """Render a block equation into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            equation: Equation block to render.
            context: Current DOCX render context.
        """

        self._render_equation(
            container,
            equation,
            context.theme,
            context.render_index,
            context.unit,
        )

    def render_page_break(
        self,
        container: object,
        block: PageBreak,
        context: DocxRenderContext,
    ) -> None:
        """Render an explicit page break into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            block: Page break block to render.
            context: Current DOCX render context.
        """

        paragraph = self._add_paragraph(container)
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def render_vertical_space(
        self,
        container: object,
        block: VerticalSpace,
        context: DocxRenderContext,
    ) -> None:
        """Render a LaTeX-like vertical spacer into DOCX.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            block: Vertical space block to render.
            context: Current DOCX render context.
        """

        paragraph = self._add_paragraph(container)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(block.height_in_points())
        paragraph.paragraph_format.line_spacing = Pt(1)
        run = paragraph.add_run("")
        run.font.size = Pt(1)

    def render_divider(
        self,
        container: object,
        block: Divider,
        context: DocxRenderContext,
    ) -> None:
        """Render a horizontal divider into DOCX.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            block: Divider block to render.
            context: Current DOCX render context.
        """

        paragraph = self._add_paragraph(container)
        paragraph.alignment = ALIGNMENTS[block.alignment]
        paragraph.paragraph_format.space_before = Pt(block.space_before)
        paragraph.paragraph_format.space_after = Pt(block.space_after)
        paragraph.paragraph_format.line_spacing = Pt(1)
        self._apply_divider_width(paragraph, block, context)
        self._set_paragraph_bottom_border(paragraph, block.color, block.thickness)

    def render_box(
        self,
        container: object,
        box: Box,
        context: DocxRenderContext,
    ) -> None:
        """Render a box and its child blocks into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            box: Box block to render.
            context: Current DOCX render context.
        """

        self._render_box(
            container,
            box,
            context.theme,
            context.render_index,
            context.settings,
            context.unit,
            word_document=context.word_document,
        )

    def render_countable_block(
        self,
        container: object,
        block: CountableBlock,
        context: DocxRenderContext,
    ) -> None:
        """Render a theorem-like countable block into DOCX.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            block: Countable block to render.
            context: Current DOCX render context.
        """

        if block.box_style is not None:
            boxed_block = Box(
                *block.children,
                title=block.heading_fragments(context.render_index.countable_number(block)),
                style=block.box_style,
            )
            anchor = context.render_index.block_anchor(block)
            if anchor is not None:
                context.render_index.block_anchors[id(boxed_block)] = anchor
            self.render_box(container, boxed_block, context)
            return

        paragraph = self._add_paragraph(container)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_with_next = True
        anchor = context.render_index.block_anchor(block)
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)
        label_run = paragraph.add_run(
            block.heading_label(context.render_index.countable_number(block))
        )
        self._apply_run_style(
            label_run,
            TextStyle(bold=True),
            default_size=context.theme.typography.body_font_size,
        )
        if block.title is not None:
            paragraph.add_run(" ")
            self._append_runs(
                paragraph,
                block.title,
                default_size=context.theme.typography.body_font_size,
                default_style=TextStyle(italic=True),
                theme=context.theme,
                render_index=context.render_index,
                word_document=context.word_document,
                unit=context.unit,
            )
        for child in block.children:
            child._render_to_docx(self, container, context)

    def render_column_span(
        self,
        container: object,
        block: ColumnSpan,
        context: DocxRenderContext,
    ) -> None:
        """Render full-width content from a multicolumn flow.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            block: Column-span block to render.
            context: Current DOCX render context.
        """

        for child in block.children:
            child._render_to_docx(self, container, context)

    def render_multi_column(
        self,
        container: object,
        block: MultiColumn,
        context: DocxRenderContext,
    ) -> None:
        """Render a multicolumn flow into DOCX.

        Args:
            container: Active DOCX document container.
            block: Multi-column block to render.
            context: Current DOCX render context.
        """

        self._render_multi_column(container, block, context)

    def render_shape(
        self,
        container: object,
        shape: Shape,
        context: DocxRenderContext,
    ) -> None:
        """Render a shape either inline or as a page-positioned drawing.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            shape: Shape to render.
            context: Current DOCX render context.
        """

        self._render_positioned_item(container, shape, context)

    def render_text_box(
        self,
        container: object,
        text_box: TextBox,
        context: DocxRenderContext,
    ) -> None:
        """Render a textbox either inline or as a page-positioned drawing.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            text_box: Text box to render.
            context: Current DOCX render context.
        """

        self._render_positioned_item(container, text_box, context)

    def render_image_box(
        self,
        container: object,
        image_box: ImageBox,
        context: DocxRenderContext,
    ) -> None:
        """Render an image either inline or as a page-positioned drawing.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            image_box: Image box to render.
            context: Current DOCX render context.
        """

        self._render_positioned_item(container, image_box, context)

    def render_comment_list(
        self,
        block: ListOfComments,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated comments page into the DOCX document.

        Args:
            block: Generated comments page block.
            context: Current DOCX render context.
        """

        self._render_comment_list(
            context.word_document,
            block.title,
            context.theme,
            context.render_index,
        )

    def render_footnote_list(
        self,
        block: ListOfFootnotes,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated footnotes page into the DOCX document.

        Args:
            block: Generated footnotes page block.
            context: Current DOCX render context.
        """

        self._render_footnote_list(
            context.word_document,
            block.title,
            context.theme,
            context.render_index,
        )

    def render_reference_list(
        self,
        block: ListOfReferences,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated references page into the DOCX document.

        Args:
            block: Generated references page block.
            context: Current DOCX render context.
        """

        self._render_reference_list(
            context.word_document,
            block,
            context.theme,
            context.render_index,
        )

    def render_table_of_contents(
        self,
        block: TableOfContents,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated table of contents into the DOCX document.

        Args:
            block: Generated table-of-contents block.
            context: Current DOCX render context.
        """

        self._render_table_of_contents(
            context.word_document,
            block,
            context,
        )

    def render_list_of_tables(
        self,
        block: ListOfTables,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated list of tables into the DOCX document.

        Args:
            block: Generated table-list block.
            context: Current DOCX render context.
        """

        self._render_caption_list(
            context.word_document,
            block.title,
            context.render_index.scoped_tables(block),
            context.theme,
            context.render_index,
            context.theme.resolve_generated_page_title("list_of_tables"),
            context.theme.resolve_caption_label("table", "caption"),
            show_page_numbers=block.show_page_numbers,
            leader=block.leader,
            text_width=context.settings.text_width_in_inches(),
        )

    def render_list_of_figures(
        self,
        block: ListOfFigures,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated list of figures into the DOCX document.

        Args:
            block: Generated figure-list block.
            context: Current DOCX render context.
        """

        self._render_caption_list(
            context.word_document,
            block.title,
            context.render_index.scoped_figures(block),
            context.theme,
            context.render_index,
            context.theme.resolve_generated_page_title("list_of_figures"),
            context.theme.resolve_caption_label("figure", "caption"),
            show_page_numbers=block.show_page_numbers,
            leader=block.leader,
            text_width=context.settings.text_width_in_inches(),
        )

    def render_list_of_algorithms(
        self,
        block: ListOfAlgorithms,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated list of algorithms into the DOCX document.

        Args:
            block: Generated algorithm-list block.
            context: Current DOCX render context.
        """

        self._render_countable_list(
            context.word_document,
            block.title,
            context.render_index.scoped_algorithms(block),
            context.theme,
            context.render_index,
            context.theme.resolve_generated_page_title("list_of_algorithms"),
            show_page_numbers=block.show_page_numbers,
            leader=block.leader,
            text_width=context.settings.text_width_in_inches(),
        )

    def render_list_of_listings(
        self,
        block: ListOfListings,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated list of listings into the DOCX document.

        Args:
            block: Generated listing-list block.
            context: Current DOCX render context.
        """

        self._render_caption_list(
            context.word_document,
            block.title,
            context.render_index.scoped_listings(block),
            context.theme,
            context.render_index,
            context.theme.resolve_generated_page_title("list_of_listings"),
            context.theme.resolve_reference_template("code_block").singular_label,
            show_page_numbers=block.show_page_numbers,
            leader=block.leader,
            text_width=context.settings.text_width_in_inches(),
        )

    def render_table(
        self,
        container: object,
        table_block: Table,
        context: DocxRenderContext,
    ) -> None:
        """Render a table block into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other table owner.
            table_block: Table block to render.
            context: Current DOCX render context.
        """

        self._render_table(
            container,
            table_block,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
            text_width=context.settings.text_width_in_inches(),
        )

    def render_pdf_pages(
        self,
        container: object,
        block: PdfPages,
        context: DocxRenderContext,
    ) -> None:
        """Render a DOCX fallback for external PDF pages.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            block: External PDF page block to render.
            context: Current DOCX render context.
        """

        paragraph = self._add_paragraph(container)
        paragraph.alignment = ALIGNMENTS[context.theme.blocks.paragraph_text_alignment]
        label = block.title or block.source.name
        paragraph.add_run(f"PDF pages: {label} ({block.page_label()})")

    def render_figure(
        self,
        container: object,
        figure: Figure,
        context: DocxRenderContext,
    ) -> None:
        """Render a figure block into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            figure: Figure block to render.
            context: Current DOCX render context.
        """

        self._render_figure(
            container,
            figure,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
            in_box=context.in_box,
        )

    def render_subfigure_group(
        self,
        container: object,
        group: SubFigureGroup,
        context: DocxRenderContext,
    ) -> None:
        """Render a subfigure group into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            group: Subfigure group to render.
            context: Current DOCX render context.
        """

        self._render_subfigure_group(
            container,
            group,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
            in_box=context.in_box,
        )

    def render_subtable_group(
        self,
        container: object,
        group: SubTableGroup,
        context: DocxRenderContext,
    ) -> None:
        """Render a subtable group into the current DOCX container.

        Args:
            container: Active DOCX document, cell, or other paragraph owner.
            group: Subtable group to render.
            context: Current DOCX render context.
        """

        self._render_subtable_group(
            container,
            group,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
            text_width=context.settings.text_width_in_inches(),
            in_box=context.in_box,
        )

    def _configure_document(self, word_document: WordDocument, document: Document) -> None:
        settings = document.settings
        theme = settings.theme
        properties = word_document.core_properties
        properties.title = settings.resolved_metadata_title(document.title)
        author = settings.resolved_author()
        if author:
            properties.author = author
        subject = settings.resolved_metadata_subject()
        if subject:
            properties.subject = subject
        keywords = settings.resolved_metadata_keywords_text()
        if keywords:
            properties.keywords = keywords
        description = settings.metadata.description
        if description:
            properties.comments = description
        self._enable_field_updates(word_document)
        self._configure_document_language(word_document, theme)

        normal_style = word_document.styles["Normal"]
        self._set_style_font_family(normal_style, theme.resolve_body_font())
        normal_style.font.size = Pt(theme.typography.body_font_size)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
        footer_style = word_document.styles["Footer"]
        self._set_style_font_family(footer_style, theme.resolve_body_font())
        footer_style.font.size = Pt(theme.page_numbers.page_number_font_size)
        footer_style.font.color.rgb = RGBColor(0, 0, 0)
        for section in word_document.sections:
            self._configure_section_page_box(section, settings)

        self._configure_named_style(
            word_document,
            "Title",
            font_name=theme.resolve_body_font(),
            font_size=theme.typography.title_font_size,
            bold=True,
            italic=False,
        )
        for level in range(1, 5):
            heading_style = theme.resolve_heading_style(level)
            heading_text_style = heading_style.text_style
            self._configure_named_style(
                word_document,
                f"Heading {level}",
                font_name=heading_text_style.font_name or theme.resolve_body_font(),
                font_size=heading_text_style.font_size or theme.resolve_heading_size(level),
                bold=bool(heading_text_style.bold),
                italic=bool(heading_text_style.italic),
                text_color=heading_text_style.text_color,
            )
        self._set_page_background(word_document, theme.blocks.page_background_color)

    def _enable_field_updates(self, word_document: WordDocument) -> None:
        settings = word_document.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def _configure_document_language(
        self,
        word_document: WordDocument,
        theme: Theme,
    ) -> None:
        language_tag = theme.resolve_language_tag()
        styles_element = word_document.styles.element
        doc_defaults = styles_element.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_element.insert(0, doc_defaults)
        run_properties_default = doc_defaults.find(qn("w:rPrDefault"))
        if run_properties_default is None:
            run_properties_default = OxmlElement("w:rPrDefault")
            doc_defaults.append(run_properties_default)
        run_properties = run_properties_default.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            run_properties_default.append(run_properties)
        language = run_properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            run_properties.append(language)
        language.set(qn("w:val"), language_tag)
        language.set(qn("w:eastAsia"), language_tag)
        language.set(qn("w:bidi"), language_tag)

    def _render_block(
        self,
        container: object,
        block: object,
        context: DocxRenderContext,
    ) -> None:
        """Delegate block rendering back to the block instance itself."""

        block._render_to_docx(self, container, context)

    def _configure_section_page_box(self, section: object, settings: object) -> None:
        top, right, bottom, left = settings.page_margin_inches()
        section.page_width = Inches(settings.page_width_in_inches())
        section.page_height = Inches(settings.page_height_in_inches())
        section.top_margin = Inches(top)
        section.right_margin = Inches(right)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)

    def _context_with_page_layout(
        self,
        context: DocxRenderContext,
        page_layout: PageLayout,
    ) -> DocxRenderContext:
        return replace(
            context,
            settings=_settings_with_page_layout(context.settings, page_layout),
        )

    def _block_page_layout(self, block: object) -> PageLayout | None:
        if isinstance(block, Section):
            return block.page_layout
        return None

    def _ensure_docx_page_layout(
        self,
        container: object,
        context: DocxRenderContext,
        page_layout: PageLayout,
    ) -> DocxRenderContext:
        if context.settings.page_layout == page_layout:
            return context
        if self._is_cell_container(container) or not hasattr(container, "add_section"):
            return context
        scoped_context = self._context_with_page_layout(context, page_layout)
        section = container.add_section(WD_SECTION.NEW_PAGE)
        self._configure_section_page_box(section, scoped_context.settings)
        return scoped_context

    def _render_docx_flow_children(
        self,
        container: object,
        children: list[object],
        context: DocxRenderContext,
        *,
        restore_at_end: bool,
    ) -> DocxRenderContext:
        active_context = context
        for child in children:
            target_layout = self._block_page_layout(child) or context.settings.page_layout
            active_context = self._ensure_docx_page_layout(
                container,
                active_context,
                target_layout,
            )
            child._render_to_docx(self, container, active_context)
        if restore_at_end and active_context.settings.page_layout != context.settings.page_layout:
            active_context = self._ensure_docx_page_layout(
                container,
                active_context,
                context.settings.page_layout,
            )
        return active_context

    def _set_section_columns(self, section: object, columns: int, gap_inches: float) -> None:
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), str(max(columns, 1)))
        cols.set(qn("w:space"), str(max(int(round(gap_inches * 1440)), 0)))

    def _render_top_level_children(
        self,
        word_document: WordDocument,
        children: list[object],
        context: DocxRenderContext,
    ) -> None:
        active_context = context
        for index, child in enumerate(children):
            target_layout = self._block_page_layout(child) or context.settings.page_layout
            active_context = self._ensure_docx_page_layout(
                word_document,
                active_context,
                target_layout,
            )
            if self._is_paginated_generated_page(child) and context.theme.generated_content.generated_content_page_breaks:
                if word_document.paragraphs and not self._ends_with_page_break(word_document):
                    self._ensure_page_break(word_document)
                child._render_to_docx(self, word_document, active_context)
                if index < len(children) - 1:
                    self._ensure_page_break(word_document)
                continue
            child._render_to_docx(self, word_document, active_context)
            if isinstance(child, Part) and not child.children and index < len(children) - 1:
                self._ensure_page_break(word_document)

    def _render_title_matter(
        self,
        word_document: WordDocument,
        document: Document,
        context: DocxRenderContext,
    ) -> None:
        title_matter = document.settings.title_matter
        cover = title_matter.cover
        cover_style = cover.resolved_style() if cover is not None else None
        title_alignment = (
            cover_style.text_alignment
            if cover_style is not None
            else context.theme.title_matter.title_text_alignment
        )
        if cover is not None:
            if cover_style.top_spacing:
                spacer = self._add_paragraph(word_document)
                spacer.paragraph_format.space_after = Inches(cover_style.top_spacing)
            for child in cover.extra_top:
                child._render_to_docx(self, word_document, context)
            if cover.eyebrow is not None:
                self._add_title_line(
                    word_document,
                    cover.eyebrow,
                    font_size=cover_style.eyebrow_style.font_size or 10,
                    alignment=title_alignment,
                    default_style=cover_style.eyebrow_style,
                    space_after=cover_style.section_spacing * 72,
                )
            if cover.organization is not None:
                self._add_title_line(
                    word_document,
                    cover.organization,
                    font_size=cover_style.organization_style.font_size or 11,
                    alignment=title_alignment,
                    default_style=cover_style.organization_style,
                    space_after=cover_style.section_spacing * 72,
                )
            logo = cover.logo_figure()
            if logo is not None:
                logo._render_to_docx(self, word_document, context)
        self._add_title_line(
            word_document,
            [Text(document.title)],
            font_size=context.theme.typography.title_font_size,
            alignment=title_alignment,
            bold=True,
            space_after=12,
        )
        if title_matter.subtitle is not None:
            self._add_title_line(
                word_document,
                title_matter.subtitle,
                font_size=max(context.theme.typography.body_font_size + 1, 12),
                alignment=title_alignment
                if cover is not None
                else context.theme.title_matter.subtitle_text_alignment,
                italic=True,
                space_after=10,
            )
        author_lines = list(title_matter.iter_author_title_lines())
        for index, (line, _is_last_for_author) in enumerate(author_lines):
            self._add_title_line(
                word_document,
                list(line.fragments),
                font_size=self._title_line_font_size(line, context.theme),
                alignment=title_alignment
                if cover is not None
                else self._title_line_alignment(line, context.theme),
                italic=line.kind == "affiliation",
                space_after=self._author_title_line_space_after(author_lines, index, last_space=10),
            )
        if cover is not None:
            if cover.date is not None:
                self._add_title_line(
                    word_document,
                    cover.date,
                    font_size=cover_style.date_style.font_size or 10,
                    alignment=title_alignment,
                    default_style=cover_style.date_style,
                    space_after=cover_style.section_spacing * 72,
                )
            for child in (*cover.note, *cover.extra_bottom):
                child._render_to_docx(self, word_document, context)
            if cover.footer is not None:
                self._add_title_line(
                    word_document,
                    cover.footer,
                    font_size=cover_style.footer_style.font_size or 9,
                    alignment=title_alignment,
                    default_style=cover_style.footer_style,
                )

    def _add_title_line(
        self,
        container: object,
        fragments: list[Text],
        *,
        font_size: float,
        alignment: str,
        bold: bool = False,
        italic: bool = False,
        space_after: float = 0,
        anchor: str | None = None,
        default_style: TextStyle | None = None,
    ) -> object:
        paragraph = self._add_paragraph(container)
        paragraph.alignment = ALIGNMENTS[alignment]
        paragraph.paragraph_format.space_after = Pt(space_after)
        base_style = (default_style or TextStyle()).merged(
            TextStyle(font_size=font_size, bold=bold, italic=italic)
        )
        for fragment in fragments:
            style = base_style.merged(
                TextStyle(
                    font_name=fragment.style.font_name,
                    font_size=fragment.style.font_size,
                    text_color=fragment.style.text_color,
                    bold=fragment.style.bold,
                    italic=fragment.style.italic,
                    underline=fragment.style.underline,
                )
            )
            if isinstance(fragment, Hyperlink):
                self._append_hyperlink_runs(
                    paragraph,
                    fragment.target,
                    fragment.label,
                    internal=fragment.internal,
                    style=style,
                    default_size=font_size,
                )
                continue
            if isinstance(fragment, InlineChip):
                self._append_inline_chip_image(
                    paragraph,
                    fragment,
                    style=style,
                    default_size=font_size,
                    theme=Theme(),
                )
                continue
            run = paragraph.add_run(self._resolve_fragment_text(fragment, None, None))
            self._apply_run_style(run, style, default_size=font_size)
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)
        return paragraph

    def _title_line_alignment(self, line: AuthorTitleLine, theme: Theme) -> str:
        if line.kind == "name":
            return theme.title_matter.author_text_alignment
        if line.kind == "affiliation":
            return theme.title_matter.affiliation_text_alignment
        return theme.title_matter.author_detail_text_alignment

    def _title_line_font_size(self, line: AuthorTitleLine, theme: Theme) -> float:
        if line.kind == "name":
            return theme.typography.body_font_size
        if line.kind == "affiliation":
            return max(theme.typography.body_font_size - 0.5, 9)
        return max(theme.typography.body_font_size - 1, 9)

    def _author_title_line_space_after(
        self,
        lines: list[tuple[AuthorTitleLine, bool]],
        index: int,
        *,
        last_space: float,
    ) -> float:
        line, is_last = lines[index]
        if is_last:
            return last_space
        next_line = lines[index + 1][0] if index + 1 < len(lines) else None
        if line.kind == "name":
            return 8
        if line.kind == "affiliation" and next_line is not None and next_line.kind == "detail":
            return 7
        return 3

    def _is_paginated_generated_page(self, block: object) -> bool:
        return isinstance(
            block,
            (
                ListOfGlossaryTerms,
                ListOfTables,
                ListOfFigures,
                ListOfAlgorithms,
                ListOfListings,
                TableOfContents,
            ),
        )

    def _should_auto_render_footnote_list(
        self,
        document: Document,
        render_index: RenderIndex,
    ) -> bool:
        return (
            not self._native_footnotes_enabled
            and document.settings.theme.blocks.auto_footnotes_page
            and bool(render_index.footnotes)
            and not any(isinstance(child, ListOfFootnotes) for child in document.body.children)
        )

    def _should_use_native_footnotes(
        self,
        theme: Theme,
        render_index: RenderIndex,
    ) -> bool:
        return (
            theme.blocks.footnote_placement == "page"
            and all(
                theme.footnotes.is_native_docx_compatible(entry.stream)
                for entry in render_index.footnotes
            )
        )

    def _keep_with_next(self, paragraph: object) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        paragraph_properties.append(OxmlElement("w:keepNext"))

    def _keep_lines_together(self, paragraph: object) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        paragraph_properties.append(OxmlElement("w:keepLines"))

    def _prevent_table_row_split(self, table: object) -> None:
        for row in table.rows:
            properties = row._tr.get_or_add_trPr()
            properties.append(OxmlElement("w:cantSplit"))

    def _repeat_table_header_rows(self, table: object, header_row_count: int) -> None:
        for row in table.rows[:header_row_count]:
            properties = row._tr.get_or_add_trPr()
            properties.append(OxmlElement("w:tblHeader"))

    def _keep_table_together(self, table: object) -> None:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._keep_with_next(paragraph)

    def _apply_media_placement_before(
        self,
        container: object,
        word_document: WordDocument,
        placement: str,
    ) -> None:
        if self._is_cell_container(container):
            return
        if placement in {"top", "page"} and (word_document.paragraphs or word_document.tables):
            self._ensure_page_break(word_document)

    def _apply_media_placement_after(
        self,
        container: object,
        word_document: WordDocument,
        placement: str,
    ) -> None:
        if self._is_cell_container(container):
            return
        if placement == "page":
            self._ensure_page_break(word_document)

    def _docx_footnotes_part(self, word_document: WordDocument) -> FootnotesPart:
        if self._native_footnotes_part is not None:
            return self._native_footnotes_part
        try:
            footnotes_part = word_document.part.part_related_by(RT.FOOTNOTES)
        except KeyError:
            footnotes_part = FootnotesPart.default(word_document.part.package)
            word_document.part.relate_to(footnotes_part, RT.FOOTNOTES)
        self._native_footnotes_part = footnotes_part
        return footnotes_part

    def _ensure_native_footnote(
        self,
        fragment: Footnote,
        *,
        theme: Theme,
        render_index: RenderIndex,
        word_document: WordDocument,
    ) -> int:
        footnote_id = render_index.footnote_number(fragment)
        if footnote_id in self._rendered_native_footnotes:
            return footnote_id

        footnotes_part = self._docx_footnotes_part(word_document)
        paragraph = footnotes_part.add_footnote_paragraph(footnote_id)
        paragraph.paragraph_format.space_after = Pt(0)
        reference_run = paragraph.add_run()
        reference_run.font.superscript = True
        reference_run._r.append(OxmlElement("w:footnoteRef"))
        spacer_run = paragraph.add_run(" ")
        self._apply_run_style(
            spacer_run,
            Text("").style,
            default_size=max(theme.typography.body_font_size - 1, 8),
        )
        self._append_runs(
            paragraph,
            fragment.note,
            default_size=max(theme.typography.body_font_size - 1, 8),
            theme=theme,
            render_index=render_index,
            word_document=word_document,
        )
        self._rendered_native_footnotes.add(footnote_id)
        return footnote_id

    def _append_native_footnote_reference(self, run: object, footnote_id: int) -> None:
        reference = OxmlElement("w:footnoteReference")
        reference.set(qn("w:id"), str(footnote_id))
        run._r.append(reference)

    def _ensure_page_break(self, word_document: WordDocument) -> None:
        if self._ends_with_page_break(word_document):
            return
        word_document.add_page_break()

    def _ends_with_page_break(self, word_document: WordDocument) -> bool:
        if not word_document.paragraphs:
            return False
        last_paragraph_xml = word_document.paragraphs[-1]._p.xml
        return 'w:type="page"' in last_paragraph_xml or "<w:sectPr" in last_paragraph_xml

    def _add_heading(
        self,
        container: object,
        title: list[Text],
        level: int,
        theme: Theme,
        *,
        number_label: str | None = None,
        anchor: str | None = None,
        render_index: RenderIndex | None = None,
        toc: bool = False,
        heading_style: HeadingStyle | None = None,
    ) -> None:
        paragraph = self._add_paragraph(container)
        paragraph.style = "Title" if level == 0 else f"Heading {min(level, 9)}"
        default_size = theme.typography.title_font_size
        default_style: TextStyle | None = None
        if level == 0:
            paragraph.alignment = ALIGNMENTS[theme.title_matter.title_text_alignment]
        else:
            resolved_style = theme.resolve_heading_style(level, heading_style)
            paragraph.alignment = ALIGNMENTS[resolved_style.text_alignment or "left"]
            if resolved_style.space_before is not None:
                paragraph.paragraph_format.space_before = Pt(resolved_style.space_before)
            if resolved_style.space_after is not None:
                paragraph.paragraph_format.space_after = Pt(resolved_style.space_after)
            if resolved_style.leading is not None:
                paragraph.paragraph_format.line_spacing = Pt(resolved_style.leading)
            default_size = resolved_style.text_style.font_size or theme.resolve_heading_size(level)
            default_style = resolved_style.text_style
        heading_fragments = self._heading_fragments(title, number_label)
        self._append_runs(
            paragraph,
            heading_fragments,
            default_size=default_size,
            default_style=default_style,
            theme=theme,
            render_index=render_index,
        )
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)
        if toc:
            self._append_toc_entry_field(
                paragraph,
                self._flatten_fragments(heading_fragments, theme, render_index),
                level=level,
            )

    def _add_paragraph(self, container: object) -> object:
        if self._is_cell_container(container):
            container_id = id(container)
            if container_id not in self._initialized_cells and container.paragraphs:
                self._initialized_cells.add(container_id)
                return container.paragraphs[0]
            self._initialized_cells.add(container_id)
        return container.add_paragraph()

    def _is_cell_container(self, container: object) -> bool:
        return hasattr(container, "_tc") and hasattr(container, "add_table")

    def _assert_document_container(self, container: object, block_name: str) -> None:
        if self._is_cell_container(container):
            raise OODocsError(f"{block_name} cannot be rendered inside a Box")

    def _assert_box_child_supported(self, child: object) -> None:
        if isinstance(
            child,
            (
                ListOfComments,
                ListOfFootnotes,
                ListOfGlossaryTerms,
                ListOfReferences,
                TableOfContents,
                ListOfTables,
                ListOfFigures,
                ListOfAlgorithms,
                ListOfListings,
                Part,
            ),
        ):
            raise OODocsError(
                f"{type(child).__name__} cannot be rendered inside a Box"
            )

    def _heading_fragments(self, title: list[Text], number_label: str | None) -> list[Text]:
        if not number_label:
            return title
        return [Text(f"{number_label} ")] + title

    def _apply_paragraph_style(
        self,
        paragraph: object,
        style: ParagraphStyle,
        theme: Theme,
        default_unit: str,
    ) -> None:
        paragraph.alignment = ALIGNMENTS[theme.resolve_paragraph_text_alignment(style)]
        if style.space_before is not None:
            paragraph.paragraph_format.space_before = Pt(style.space_before)
        if style.space_after is not None:
            paragraph.paragraph_format.space_after = Pt(style.space_after)
        if style.leading is not None:
            paragraph.paragraph_format.line_spacing = Pt(style.leading)
        if style.keep_together is not None:
            paragraph.paragraph_format.keep_together = style.keep_together
        if style.keep_with_next is not None:
            paragraph.paragraph_format.keep_with_next = style.keep_with_next
        if style.page_break_before is not None:
            paragraph.paragraph_format.page_break_before = style.page_break_before
        if style.widow_control is not None:
            paragraph.paragraph_format.widow_control = style.widow_control
        left_indent = style.left_indent_in_inches(default_unit)
        right_indent = style.right_indent_in_inches(default_unit)
        first_line_indent = style.first_line_indent_in_inches(default_unit)
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = Inches(left_indent)
        if right_indent is not None:
            paragraph.paragraph_format.right_indent = Inches(right_indent)
        if first_line_indent is not None:
            paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)

    def _apply_divider_width(
        self,
        paragraph: object,
        block: Divider,
        context: DocxRenderContext,
    ) -> None:
        width = block.width_in_inches(context.unit)
        if width is None:
            return
        text_width = context.settings.text_width_in_inches()
        extra = max(text_width - width, 0)
        if block.alignment == "left":
            paragraph.paragraph_format.right_indent = Inches(extra)
        elif block.alignment == "right":
            paragraph.paragraph_format.left_indent = Inches(extra)
        else:
            paragraph.paragraph_format.left_indent = Inches(extra / 2)
            paragraph.paragraph_format.right_indent = Inches(extra / 2)

    def _append_runs(
        self,
        paragraph: object,
        fragments: list[Text],
        default_size: float | None = None,
        *,
        default_style: TextStyle | None = None,
        theme: Theme | None = None,
        render_index: RenderIndex | None = None,
        word_document: WordDocument | None = None,
        unit: str = "in",
    ) -> None:
        for fragment in fragments:
            if isinstance(fragment, (ImageBox, Shape, TextBox)):
                self._append_inline_positioned_fragment(
                    paragraph,
                    fragment,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                    unit=unit,
                )
                continue
            fragment_style = default_style.merged(fragment.style) if default_style is not None else fragment.style
            if isinstance(fragment, InlineChip):
                self._append_inline_chip_image(
                    paragraph,
                    fragment,
                    style=fragment_style,
                    default_size=default_size,
                    theme=theme or Theme(),
                )
                continue
            if isinstance(fragment, LineBreak):
                paragraph.add_run().add_break()
                continue
            if isinstance(fragment, Hyperlink):
                link_style = (
                    theme.resolve_link_text_style(fragment_style)
                    if theme is not None
                    else fragment_style
                )
                self._append_hyperlink_runs(
                    paragraph,
                    fragment.target,
                    fragment.label,
                    internal=fragment.internal,
                    style=link_style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, ReferenceGroup) and theme is not None and render_index is not None:
                self._append_reference_pieces(
                    paragraph,
                    reference_text_pieces(
                        fragment.targets,
                        fragment.reference_format,
                        theme,
                        render_index,
                        range_reference=fragment.range_reference,
                    ),
                    style=fragment_style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, _BlockReference) and theme is not None and render_index is not None:
                anchor = self._block_reference_anchor(fragment.target, render_index)
                if fragment.label is None:
                    self._append_reference_pieces(
                        paragraph,
                        reference_text_pieces(
                            (fragment.target,),
                            fragment.reference_format,
                            theme,
                            render_index,
                        ),
                        style=fragment_style,
                        default_size=default_size,
                    )
                    continue
                self._append_hyperlink_runs(
                    paragraph,
                    anchor,
                    fragment.label,
                    internal=True,
                    style=fragment_style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, Citation) and render_index is not None:
                citation_entry = render_index.citation_entry(fragment.target)
                citation_label = format_citation_label(
                    citation_entry.source,
                    citation_entry.number,
                    theme.citations.citation_style if theme is not None else "numeric",
                )
                self._append_hyperlink_runs(
                    paragraph,
                    citation_entry.anchor,
                    [Text(citation_label, style=fragment_style)],
                    internal=True,
                    style=fragment_style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, Comment):
                self._append_comment_runs(
                    paragraph,
                    fragment,
                    default_size=default_size,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )
                continue
            if isinstance(fragment, Footnote):
                self._append_footnote_runs(
                    paragraph,
                    fragment,
                    default_size=default_size,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )
                continue
            if isinstance(fragment, Math):
                self._append_math_runs(paragraph, fragment, default_size=default_size)
                continue
            run = paragraph.add_run(self._resolve_fragment_text(fragment, theme, render_index))
            self._apply_run_style(run, fragment_style, default_size=default_size)

    def _append_inline_positioned_fragment(
        self,
        paragraph: object,
        fragment: PositionedItem,
        *,
        theme: Theme | None,
        render_index: RenderIndex | None,
        word_document: WordDocument | None,
        unit: str,
    ) -> None:
        if isinstance(fragment, ImageBox):
            paragraph.add_run().add_picture(
                self._image_box_picture_source(fragment),
                width=Inches(length_to_inches(fragment.width, fragment.unit or unit)),
                height=Inches(length_to_inches(fragment.height, fragment.unit or unit)),
            )
            return
        box = PositionedBox(
            item=fragment,
            x=0.0,
            y=0.0,
            width=length_to_inches(fragment.width, fragment.unit or unit),
            height=length_to_inches(fragment.height, fragment.unit or unit),
        )
        if isinstance(fragment, Shape):
            pict = self._shape_pict(fragment, box, absolute=False)
        else:
            pict = self._text_box_pict(
                fragment,
                box,
                theme=theme or Theme(),
                absolute=False,
            )
        paragraph.add_run()._r.append(pict)

    def _append_inline_chip_image(
        self,
        paragraph: object,
        fragment: InlineChip,
        *,
        style: TextStyle,
        default_size: float | None,
        theme: Theme,
    ) -> None:
        image, width_inches, height_inches = self._inline_chip_image(
            fragment,
            style=style,
            default_size=default_size,
            theme=theme,
        )
        paragraph.add_run().add_picture(
            image,
            width=Inches(width_inches),
            height=Inches(height_inches),
        )

    def _inline_chip_image(
        self,
        fragment: InlineChip,
        *,
        style: TextStyle,
        default_size: float | None,
        theme: Theme,
    ) -> tuple[BytesIO, float, float]:
        dpi = 192
        chip_style = theme.stylesheet.resolve("chip", fragment.chip_style, None)
        base_size = style.font_size or default_size or theme.typography.body_font_size
        font_size = max(base_size + chip_style.font_size_delta, 6.0)
        font_px = max(int(round(font_size * dpi / 72)), 8)
        font = self._inline_chip_font(
            chip_style.font_name or style.font_name or theme.resolve_body_font(),
            chip_style.bold,
            chip_style.italic,
            font_px,
        )
        text = fragment.display_text(chip_style)
        sample = Image.new("RGBA", (1, 1), (255, 255, 255, 0))
        draw = ImageDraw.Draw(sample)
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = max(text_bbox[2] - text_bbox[0], 1)
        text_height = max(text_bbox[3] - text_bbox[1], 1)
        border_px = max(int(round(chip_style.border.width_points() * dpi / 72)), 0)
        top_padding, right_padding, bottom_padding, left_padding = chip_style.padding.as_tuple()
        if chip_style.padding.unit == "em":
            top_padding_px = max(int(round(font_px * top_padding)), 0)
            right_padding_px = max(int(round(font_px * right_padding)), 0)
            bottom_padding_px = max(int(round(font_px * bottom_padding)), 0)
            left_padding_px = max(int(round(font_px * left_padding)), 0)
        else:
            top_pt, right_pt, bottom_pt, left_pt = chip_style.padding.to_points()
            top_padding_px = max(int(round(top_pt * dpi / 72)), 0)
            right_padding_px = max(int(round(right_pt * dpi / 72)), 0)
            bottom_padding_px = max(int(round(bottom_pt * dpi / 72)), 0)
            left_padding_px = max(int(round(left_pt * dpi / 72)), 0)
        width = text_width + left_padding_px + right_padding_px + border_px * 2 + 2
        height = text_height + top_padding_px + bottom_padding_px + border_px * 2 + 2
        if chip_style.border.radius_unit == "em":
            radius = max(int(round(font_px * chip_style.border.radius_em())), 0)
        else:
            radius = max(int(round(chip_style.border.radius_points() * dpi / 72)), 0)
        image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
        draw = ImageDraw.Draw(image)
        rectangle = (0, 0, width - 1, height - 1)
        outline = self._inline_chip_rgba(chip_style.border.color) if border_px and chip_style.border.color else None
        draw.rounded_rectangle(
            rectangle,
            radius=radius,
            fill=self._inline_chip_rgba(chip_style.background_color),
            outline=outline,
            width=border_px or 1,
        )
        text_x = left_padding_px + border_px + 1 - text_bbox[0]
        text_y = top_padding_px + border_px + 1 - text_bbox[1]
        draw.text(
            (text_x, text_y),
            text,
            font=font,
            fill=self._inline_chip_rgba(chip_style.text_color),
        )
        buffer = BytesIO()
        image.save(buffer, format="PNG", dpi=(dpi, dpi))
        buffer.seek(0)
        return buffer, width / dpi, height / dpi

    def _inline_chip_font(
        self,
        font_name: str,
        bold: bool,
        italic: bool,
        font_px: int,
    ) -> object:
        candidates = self._inline_chip_font_candidates(font_name, bold, italic)
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, font_px)
            except OSError:
                continue
        try:
            return ImageFont.truetype(font_name, font_px)
        except OSError:
            return ImageFont.load_default()

    def _inline_chip_font_candidates(self, font_name: str, bold: bool, italic: bool) -> list[str]:
        key = font_name.lower()
        windows_fonts = Path("C:/Windows/Fonts")
        font_files = {
            "arial": {
                (False, False): "arial.ttf",
                (True, False): "arialbd.ttf",
                (False, True): "ariali.ttf",
                (True, True): "arialbi.ttf",
            },
            "courier new": {
                (False, False): "cour.ttf",
                (True, False): "courbd.ttf",
                (False, True): "couri.ttf",
                (True, True): "courbi.ttf",
            },
            "times new roman": {
                (False, False): "times.ttf",
                (True, False): "timesbd.ttf",
                (False, True): "timesi.ttf",
                (True, True): "timesbi.ttf",
            },
            "malgun gothic": {
                (False, False): "malgun.ttf",
                (True, False): "malgunbd.ttf",
                (False, True): "malgun.ttf",
                (True, True): "malgunbd.ttf",
            },
            "d2coding": {
                (False, False): "D2Coding.ttf",
                (True, False): "D2CodingBold.ttf",
                (False, True): "D2Coding.ttf",
                (True, True): "D2CodingBold.ttf",
            },
        }
        candidates: list[str] = []
        if key in font_files:
            candidates.append(str(windows_fonts / font_files[key][(bold, italic)]))
        candidates.extend(
            [
                font_name,
                "DejaVuSans-BoldOblique.ttf" if bold and italic else "",
                "DejaVuSans-Bold.ttf" if bold else "",
                "DejaVuSans-Oblique.ttf" if italic else "",
                "DejaVuSans.ttf",
            ]
        )
        return [candidate for candidate in candidates if candidate]

    def _inline_chip_rgba(self, color: str) -> tuple[int, int, int, int]:
        return (
            int(color[0:2], 16),
            int(color[2:4], 16),
            int(color[4:6], 16),
            255,
        )

    def _append_hyperlink_runs(
        self,
        paragraph: object,
        target: str | None,
        label_fragments: list[Text],
        *,
        internal: bool,
        style: TextStyle,
        default_size: float | None,
    ) -> None:
        if not target:
            self._append_runs(
                paragraph,
                label_fragments,
                default_size=default_size,
            )
            return

        hyperlink = OxmlElement("w:hyperlink")
        if internal:
            hyperlink.set(qn("w:anchor"), target)
        else:
            relationship_id = paragraph.part.relate_to(
                target,
                RT.HYPERLINK,
                is_external=True,
            )
            hyperlink.set(qn("r:id"), relationship_id)

        label_text = "".join(fragment.plain_text() for fragment in label_fragments)
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        if style.font_name is not None:
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), style.font_name)
            fonts.set(qn("w:hAnsi"), style.font_name)
            fonts.set(qn("w:eastAsia"), style.font_name)
            fonts.set(qn("w:cs"), style.font_name)
            run_properties.append(fonts)
        font_size = style.font_size if style.font_size is not None else default_size
        if font_size is not None:
            size = OxmlElement("w:sz")
            size.set(qn("w:val"), str(int(round(font_size * 2))))
            run_properties.append(size)
        if style.bold is not None:
            bold = OxmlElement("w:b")
            if not style.bold:
                bold.set(qn("w:val"), "0")
            run_properties.append(bold)
        if style.italic is not None:
            italic = OxmlElement("w:i")
            if not style.italic:
                italic.set(qn("w:val"), "0")
            run_properties.append(italic)
        if style.text_color is not None:
            color = OxmlElement("w:color")
            color.set(qn("w:val"), style.text_color)
            run_properties.append(color)
        if style.underline:
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            run_properties.append(underline)
        if style.strikethrough:
            strike = OxmlElement("w:strike")
            run_properties.append(strike)
        if style.small_caps is not None:
            small_caps = OxmlElement("w:smallCaps")
            if not style.small_caps:
                small_caps.set(qn("w:val"), "0")
            run_properties.append(small_caps)
        if style.uppercase is not None:
            caps = OxmlElement("w:caps")
            if not style.uppercase:
                caps.set(qn("w:val"), "0")
            run_properties.append(caps)
        if style.subscript:
            vert_align = OxmlElement("w:vertAlign")
            vert_align.set(qn("w:val"), "subscript")
            run_properties.append(vert_align)
        if style.superscript:
            vert_align = OxmlElement("w:vertAlign")
            vert_align.set(qn("w:val"), "superscript")
            run_properties.append(vert_align)
        if style.highlight_color is not None:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), style.highlight_color)
            run_properties.append(shading)
        run.append(run_properties)
        text = OxmlElement("w:t")
        text.text = label_text
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _append_reference_pieces(
        self,
        paragraph: object,
        pieces: list[ReferenceTextPiece],
        *,
        style: TextStyle,
        default_size: float | None,
    ) -> None:
        for piece in pieces:
            if piece.anchor is not None:
                self._append_hyperlink_runs(
                    paragraph,
                    piece.anchor,
                    [Text(piece.text, style=style)],
                    internal=True,
                    style=style,
                    default_size=default_size,
                )
                continue
            run = paragraph.add_run(piece.text)
            self._apply_run_style(run, style, default_size=default_size)

    def _add_bookmark(self, paragraph: object, anchor: str) -> None:
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(self._bookmark_id))
        bookmark_start.set(qn("w:name"), anchor)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(self._bookmark_id))
        insert_index = (
            1
            if len(paragraph._p) > 0 and paragraph._p[0].tag == qn("w:pPr")
            else 0
        )
        paragraph._p.insert(insert_index, bookmark_start)
        paragraph._p.append(bookmark_end)
        self._bookmark_id += 1

    def _block_reference_anchor(
        self,
        target: object,
        render_index: RenderIndex,
    ) -> str | None:
        return render_index.anchor_for(target)

    def _apply_run_style(self, run: object, style: object, *, default_size: float | None = None) -> None:
        font = run.font
        if style.font_name:
            self._set_run_font_family(run, style.font_name)
        if style.font_size is not None:
            font.size = Pt(style.font_size)
        elif default_size is not None:
            font.size = Pt(default_size)
        if style.bold is not None:
            font.bold = style.bold
        if style.italic is not None:
            font.italic = style.italic
        if style.underline is not None:
            font.underline = style.underline
        if style.text_color is not None:
            font.color.rgb = RGBColor.from_string(style.text_color)
        if style.strikethrough is not None:
            font.strike = style.strikethrough
        if style.small_caps is not None:
            font.small_caps = style.small_caps
        if style.uppercase is not None:
            run_properties = run._r.get_or_add_rPr()
            caps = OxmlElement("w:caps")
            if not style.uppercase:
                caps.set(qn("w:val"), "0")
            run_properties.append(caps)
        if style.subscript is not None:
            font.subscript = style.subscript
        if style.superscript is not None:
            font.superscript = style.superscript
        if style.highlight_color is not None:
            self._set_run_shading(run, style.highlight_color)

    def _set_run_shading(self, run: object, fill: str) -> None:
        run_properties = run._r.get_or_add_rPr()
        existing = run_properties.find(qn("w:shd"))
        if existing is not None:
            run_properties.remove(existing)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        run_properties.append(shading)

    def _append_comment_runs(
        self,
        paragraph: object,
        fragment: Comment,
        *,
        default_size: float | None,
        theme: Theme | None,
        render_index: RenderIndex | None,
        word_document: WordDocument | None,
    ) -> None:
        visible_runs: list[object] = []
        if fragment.value:
            visible_run = paragraph.add_run(fragment.value)
            self._apply_run_style(visible_run, fragment.style, default_size=default_size)
            visible_runs.append(visible_run)

        marker_run = paragraph.add_run(self._comment_marker(fragment, render_index))
        self._apply_run_style(marker_run, fragment.style, default_size=max((default_size or 10.0) - 2, 8))
        marker_run.font.superscript = True
        anchor_runs = visible_runs or [marker_run]
        if word_document is not None and render_index is not None:
            word_document.add_comment(
                anchor_runs,
                text=self._flatten_fragments(fragment.comment, theme, render_index),
                author=fragment.author or "",
                initials=fragment.initials,
            )

    def _append_footnote_runs(
        self,
        paragraph: object,
        fragment: Footnote,
        *,
        default_size: float | None,
        theme: Theme | None,
        render_index: RenderIndex | None,
        word_document: WordDocument | None,
    ) -> None:
        if fragment.value:
            visible_run = paragraph.add_run(fragment.value)
            self._apply_run_style(visible_run, fragment.style, default_size=default_size)

        if (
            theme is not None
            and word_document is not None
            and render_index is not None
            and self._native_footnotes_enabled
        ):
            marker_run = paragraph.add_run()
            self._apply_run_style(
                marker_run,
                fragment.style,
                default_size=max((default_size or 10.0) - 2, 8),
            )
            marker_run.font.superscript = True
            self._append_native_footnote_reference(
                marker_run,
                self._ensure_native_footnote(
                    fragment,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                ),
            )
            return

        marker_run = paragraph.add_run(self._footnote_marker(fragment, render_index, theme))
        self._apply_run_style(marker_run, fragment.style, default_size=max((default_size or 10.0) - 2, 8))
        marker_run.font.superscript = True

    def _append_math_runs(self, paragraph: object, fragment: Math, *, default_size: float | None = None) -> None:
        for segment in parse_latex_segments(fragment.value):
            run = paragraph.add_run(segment.text)
            self._apply_run_style(run, fragment.style, default_size=default_size)
            if segment.vertical_align == SUPERSCRIPT:
                run.font.superscript = True
            elif segment.vertical_align == SUBSCRIPT:
                run.font.subscript = True

    def _render_list(
        self,
        container: object,
        list_block: BulletList | NumberedList,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
        depth: int = 0,
        parent_style: ListStyle | None = None,
    ) -> None:
        list_style = theme.stylesheet.resolve(
            "list",
            list_block.style,
            theme.list_style(ordered=isinstance(list_block, NumberedList)),
        )
        if parent_style is not None and list_block.style is None:
            list_style = replace(
                list_style,
                marker_gap=parent_style.marker_gap,
                item_spacing=parent_style.item_spacing,
                block_spacing=parent_style.block_spacing,
            )
        list_style = list_block._effective_style(list_style)
        if depth:
            list_style = replace(list_style, indent=list_style.indent + depth * 0.25)
        for index, item in enumerate(list_block.items):
            paragraph = self._add_paragraph(container)
            item_style = theme.stylesheet.resolve("paragraph", item.style, ParagraphStyle())
            self._apply_paragraph_style(paragraph, item_style, theme, unit)
            spacing = list_style.item_spacing
            if index == len(list_block.items) - 1:
                spacing = max(spacing, list_style.block_spacing)
            paragraph.paragraph_format.space_after = Pt(spacing)
            paragraph.paragraph_format.left_indent = Inches(list_style.indent)
            paragraph.paragraph_format.first_line_indent = Inches(-list_style.indent)
            anchor = render_index.block_anchor(item)
            if anchor is not None:
                self._add_bookmark(paragraph, anchor)
            marker = list_style.marker_for(index)
            if marker:
                marker_run = paragraph.add_run(f"{marker} ")
                self._apply_run_style(marker_run, Text("").style, default_size=theme.typography.body_font_size)
            self._append_runs(paragraph, item.content, theme=theme, render_index=render_index, word_document=word_document)
            for child_list in list_block.item_children[index]:
                self._render_list(
                    container,
                    child_list,
                    theme,
                    render_index,
                    unit,
                    word_document=word_document,
                    depth=depth + 1,
                    parent_style=list_style,
                )

    def _render_code_block(
        self,
        container: object,
        code_block: CodeBlock,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
    ) -> None:
        anchor = render_index.block_anchor(code_block)
        show_label = bool(code_block.language and code_block.show_language)
        paragraph = self._add_paragraph(container)
        code_style = theme.stylesheet.resolve("paragraph", code_block.style, ParagraphStyle())
        self._apply_paragraph_style(paragraph, code_style, theme, unit)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.right_indent = Inches(0.1)
        paragraph.paragraph_format.space_before = Pt(6)
        self._set_paragraph_shading(paragraph, "F5F7FA")
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)
        if show_label and code_block.language_position.startswith("top"):
            self._append_code_language_label_run(paragraph, code_block, theme)

        self._append_code_lines(paragraph, code_block, theme=theme)
        if show_label and code_block.language_position.startswith("bottom"):
            paragraph.add_run().add_break()
            self._append_code_language_label_run(paragraph, code_block, theme, trailing=True)
        if code_block.caption is not None:
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
            caption.paragraph_format.space_after = Pt(code_style.space_after or 0)
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.resolve_reference_template("code_block").singular_label,
                    render_index.code_block_number(code_block),
                    code_block.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )

    def _append_code_language_label_run(
        self,
        paragraph: object,
        code_block: CodeBlock,
        theme: Theme,
        *,
        trailing: bool = False,
    ) -> None:
        if code_block.language_position.endswith("right"):
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
            paragraph.add_run("\t")
        run = paragraph.add_run(code_block.language.upper() if code_block.language else "")
        run.font.name = theme.resolve_monospace_font()
        run.font.size = Pt(max(theme.caption_size() - 1, 7))
        run.font.bold = False
        run.font.color.rgb = RGBColor(0x6F, 0x7D, 0x90)
        if not trailing:
            paragraph.add_run().add_break()

    def _append_code_lines(
        self,
        paragraph: object,
        code_block: CodeBlock,
        *,
        theme: Theme,
    ) -> None:
        default_size = max(theme.typography.body_font_size - 1, 8)
        for line_index, line in enumerate(code_block.normalized_lines(), start=1):
            if line_index > 1:
                paragraph.add_run().add_break()
            highlighted = line_index in code_block.highlight_lines
            prefix = code_block.line_prefix(line_index)
            if prefix:
                prefix_run = paragraph.add_run(prefix)
                prefix_run.font.name = theme.resolve_monospace_font()
                prefix_run.font.size = Pt(default_size)
                prefix_run.font.color.rgb = RGBColor(0x6F, 0x7D, 0x90)
                if highlighted:
                    self._set_run_shading(prefix_run, "FFF3B0")
            wrote_line = bool(prefix)
            for token in _syntax_line_tokens(line, code_block.language):
                if not token.text:
                    continue
                run = paragraph.add_run(token.text)
                self._apply_code_token_run(run, token, theme=theme, default_size=default_size)
                if highlighted:
                    self._set_run_shading(run, "FFF3B0")
                wrote_line = True
            if highlighted and not wrote_line:
                run = paragraph.add_run(" ")
                run.font.name = theme.resolve_monospace_font()
                run.font.size = Pt(default_size)
                self._set_run_shading(run, "FFF3B0")

    def _apply_code_token_run(
        self,
        run: object,
        token: SyntaxToken,
        *,
        theme: Theme,
        default_size: float,
    ) -> None:
        run.font.name = theme.resolve_monospace_font()
        run.font.size = Pt(default_size)
        if token.color is not None:
            run.font.color.rgb = RGBColor.from_string(token.color.upper())
        if token.bold:
            run.font.bold = True
        if token.italic:
            run.font.italic = True

    def _append_code_tokens(
        self,
        paragraph: object,
        tokens: list[SyntaxToken],
        *,
        theme: Theme,
    ) -> None:
        default_size = max(theme.typography.body_font_size - 1, 8)
        for token in tokens:
            parts = token.text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
            for index, part in enumerate(parts):
                if index:
                    paragraph.add_run().add_break()
                if not part:
                    continue
                run = paragraph.add_run(part)
                self._apply_code_token_run(run, token, theme=theme, default_size=default_size)

    def _render_equation(
        self,
        container: object,
        equation: Equation,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
    ) -> None:
        if isinstance(equation, AlignedEquation):
            self._render_aligned_equation(
                container,
                equation,
                theme,
                render_index,
                unit,
            )
            return

        paragraph = self._add_paragraph(container)
        equation_style = theme.stylesheet.resolve("paragraph", equation.style, ParagraphStyle())
        self._apply_paragraph_style(paragraph, equation_style, theme, unit)
        if paragraph.alignment is None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor = render_index.block_anchor(equation)
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)
        self._append_math_runs(
            paragraph,
            Math(equation.expression),
            default_size=max(theme.typography.body_font_size + 1, 12),
        )
        number = render_index.equation_number(equation)
        if number is not None:
            run = paragraph.add_run(f" ({number})")
            self._apply_run_style(run, TextStyle(), default_size=theme.typography.body_font_size)

    def _render_aligned_equation(
        self,
        container: object,
        equation: AlignedEquation,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
    ) -> None:
        """Render aligned equation rows in a borderless DOCX table."""

        equation_style = theme.stylesheet.resolve(
            "paragraph",
            equation.style,
            ParagraphStyle(),
        )
        maximum_columns = max(len(line.alignment_parts()) for line in equation.lines)
        table = container.add_table(rows=len(equation.lines), cols=maximum_columns + 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table_description = OxmlElement("w:tblDescription")
        table_description.set(
            qn("w:val"),
            " | ".join(" ".join(line.plain_text().split()) for line in equation.lines),
        )
        table._tbl.tblPr.append(table_description)
        group_anchor = render_index.anchor_for(equation)
        group_number = render_index.equation_number(equation)

        for row_index, line in enumerate(equation.lines):
            parts = line.alignment_parts()
            row = table.rows[row_index]
            for column_index, cell in enumerate(row.cells):
                cell._tc.clear_content()
                self._initialized_cells.discard(id(cell))
                self._set_cell_borders_none(cell)
                paragraph = self._add_paragraph(cell)
                self._apply_paragraph_style(paragraph, equation_style, theme, unit)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(
                    equation_style.space_after or 0
                    if row_index == len(equation.lines) - 1
                    else 0
                )

                if column_index == 0:
                    if row_index == 0 and group_anchor is not None:
                        self._add_bookmark(paragraph, group_anchor)
                    line_anchor = render_index.anchor_for(line)
                    if line_anchor is not None:
                        self._add_bookmark(paragraph, line_anchor)

                if column_index < maximum_columns:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT
                        if maximum_columns > 1 and column_index % 2 == 0
                        else WD_ALIGN_PARAGRAPH.LEFT
                        if maximum_columns > 1
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    if column_index < len(parts) and parts[column_index]:
                        self._append_math_runs(
                            paragraph,
                            Math(parts[column_index]),
                            default_size=max(theme.typography.body_font_size + 1, 12),
                        )
                    continue

                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                number = render_index.equation_number(line)
                if number is None and row_index == len(equation.lines) - 1:
                    number = group_number
                if number is not None:
                    run = paragraph.add_run(f"({number})")
                    self._apply_run_style(
                        run,
                        TextStyle(),
                        default_size=theme.typography.body_font_size,
                    )

    def _effective_box_style(self, box: Box, theme: Theme) -> BoxStyle:
        box_style = theme.stylesheet.resolve("box", box.style, None)
        if box.title_position is None and box.shadow is None:
            return box_style
        return replace(
            box_style,
            title_position=box.title_position or box_style.title_position,
            shadow=box_style.shadow if box.shadow is None else box.shadow,
        )

    def _box_side_title_width(self, box_width: float) -> float:
        if box_width <= 1.5:
            return max(box_width / 2, 0.25)
        return min(max(box_width * 0.24, 0.75), box_width - 0.75)

    def _render_box(
        self,
        container: object,
        box: Box,
        theme: Theme,
        render_index: RenderIndex,
        settings: object,
        unit: str,
        *,
        word_document: WordDocument,
    ) -> None:
        box_style = self._effective_box_style(box, theme)
        alignment = box_style.block_alignment or theme.blocks.box_block_alignment
        title_fragments = box.title_fragments()
        title_is_side = title_fragments is not None and box_style.title_position == "side"
        anchor = render_index.block_anchor(box)
        if anchor is not None:
            anchor_paragraph = self._add_paragraph(container)
            anchor_paragraph.paragraph_format.space_before = Pt(0)
            anchor_paragraph.paragraph_format.space_after = Pt(0)
            self._add_bookmark(anchor_paragraph, anchor)
        outer_table = container.add_table(rows=1, cols=2 if title_is_side else 1)
        outer_table.alignment = TABLE_ALIGNMENTS[alignment]
        box_width = None
        if box_style.width is not None:
            box_width = length_to_inches(box_style.width, box_style.unit or unit)
            outer_table.autofit = False
            self._set_table_width(outer_table, box_width)

        cells = outer_table.rows[0].cells
        body_cell = cells[-1]
        for cell in cells:
            cell._tc.clear_content()
            self._initialized_cells.discard(id(cell))
            self._set_cell_shading(cell, box_style.background_color)
            if box_style.border.color is not None and box_style.border.width > 0:
                self._set_cell_borders(cell, box_style.border.color, box_style.border.width_points())
            else:
                self._set_cell_borders_none(cell)
            self._set_cell_margins(cell, *box_style.padding.to_points())

        if title_is_side:
            title_cell = cells[0]
            if box_style.title_background_color is not None:
                self._set_cell_shading(title_cell, box_style.title_background_color)
            if box_width is not None:
                title_width = self._box_side_title_width(box_width)
                body_width = max(box_width - title_width, 0.25)
                outer_table.columns[0].width = Inches(title_width)
                outer_table.columns[1].width = Inches(body_width)
                self._set_cell_width(title_cell, title_width)
                self._set_cell_width(body_cell, body_width)
            title_paragraph = self._add_paragraph(title_cell)
            title_paragraph.paragraph_format.space_after = Pt(0)
            title_style = TextStyle(text_color=box_style.title_text_color, bold=True)
            self._append_runs(
                title_paragraph,
                title_fragments,
                default_size=theme.typography.body_font_size,
                default_style=title_style,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
        else:
            if box_width is not None:
                outer_table.columns[0].width = Inches(box_width)
                self._set_cell_width(body_cell, box_width)
            if title_fragments is not None:
                title_paragraph = self._add_paragraph(body_cell)
                title_paragraph.paragraph_format.space_after = Pt(6)
                if box_style.title_background_color is not None:
                    self._set_paragraph_shading(title_paragraph, box_style.title_background_color)
                title_style = TextStyle(text_color=box_style.title_text_color, bold=True)
                self._append_runs(
                    title_paragraph,
                    title_fragments,
                    default_size=theme.typography.body_font_size,
                    default_style=title_style,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )

        context = DocxRenderContext(
            theme=theme,
            render_index=render_index,
            settings=settings,
            unit=unit,
            word_document=word_document,
            in_box=True,
        )
        for child in box.children:
            self._assert_box_child_supported(child)
            self._render_block(body_cell, child, context)

        if not body_cell.paragraphs:
            body_cell.add_paragraph()

    def _render_multi_column(
        self,
        container: object,
        block: MultiColumn,
        context: DocxRenderContext,
    ) -> None:
        if block.columns == 1:
            for child in block.children:
                child._render_to_docx(self, container, context)
            return

        if not self._is_cell_container(container) and hasattr(container, "add_section"):
            self._render_multi_column_sections(container, block, context)
            return

        self._render_multi_column_layout_table(container, block, context)

    def _render_multi_column_sections(
        self,
        word_document: WordDocument,
        block: MultiColumn,
        context: DocxRenderContext,
    ) -> None:
        current_group: list[object] = []
        available_width = context.settings.text_width_in_inches()
        self._start_column_section(word_document, block.columns, block, context)

        def flush_group() -> None:
            if not current_group:
                return
            for group_child in current_group:
                group_child._render_to_docx(self, word_document, context)
            current_group.clear()

        for child in block.children:
            if block._child_spans_columns(
                child,
                available_width=available_width,
                default_unit=context.unit,
            ):
                flush_group()
                self._start_column_section(word_document, 1, block, context)
                child._render_to_docx(self, word_document, context)
                self._start_column_section(word_document, block.columns, block, context)
                continue
            current_group.append(child)
        flush_group()
        self._start_column_section(word_document, 1, block, context)

    def _start_column_section(
        self,
        word_document: WordDocument,
        columns: int,
        block: MultiColumn,
        context: DocxRenderContext,
    ) -> None:
        section = word_document.add_section(WD_SECTION.CONTINUOUS)
        self._configure_section_page_box(section, context.settings)
        self._set_section_columns(section, columns, block.column_gap_in_inches(context.unit))

    def _render_multi_column_layout_table(
        self,
        container: object,
        block: MultiColumn,
        context: DocxRenderContext,
    ) -> None:
        current_group: list[object] = []
        available_width = context.settings.text_width_in_inches()

        def flush_group() -> None:
            if not current_group:
                return
            self._render_multi_column_group_table(container, block, current_group, context)
            current_group.clear()

        for child in block.children:
            if block._child_spans_columns(
                child,
                available_width=available_width,
                default_unit=context.unit,
            ):
                flush_group()
                child._render_to_docx(self, container, context)
                continue
            current_group.append(child)
        flush_group()

    def _render_multi_column_group_table(
        self,
        container: object,
        block: MultiColumn,
        children: list[object],
        context: DocxRenderContext,
    ) -> None:
        table = container.add_table(rows=1, cols=block.columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gap_inches = block.column_gap_in_inches(context.unit)
        gap_points = gap_inches * 72

        if not context.in_box:
            available_width = context.settings.text_width_in_inches()
            table.autofit = False
            self._set_table_width(table, available_width)
            column_width = block.column_width_in_inches(available_width, context.unit)
            for column in table.columns:
                column.width = Inches(column_width)

        chunk_size = max((len(children) + block.columns - 1) // block.columns, 1)
        for column_index, cell in enumerate(table.rows[0].cells):
            cell._tc.clear_content()
            self._initialized_cells.discard(id(cell))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if not context.in_box:
                self._set_cell_width(
                    cell,
                    block.column_width_in_inches(
                        context.settings.text_width_in_inches(),
                        context.unit,
                    ),
                )
            self._set_cell_margins(
                cell,
                0,
                gap_points / 2 if column_index < block.columns - 1 else 0,
                0,
                gap_points / 2 if column_index > 0 else 0,
            )
            self._set_cell_borders_none(cell)
            start = column_index * chunk_size
            end = start + chunk_size
            for child in children[start:end]:
                child._render_to_docx(self, cell, context)
            if not cell.paragraphs:
                cell.add_paragraph()

    def _render_page_items(
        self,
        word_document: WordDocument,
        document: Document,
        context: DocxRenderContext,
        *,
        has_front_matter: bool,
    ) -> None:
        if not document.settings.overlays:
            return
        section = word_document.sections[0]
        if document.settings.title_matter.cover is not None:
            section.different_first_page_header_footer = True
            self._render_section_page_items(
                section.first_page_header,
                document,
                context,
                phase="cover",
                page_number=1,
            )
            if has_front_matter:
                self._render_section_page_items(
                    section.header,
                    document,
                    context,
                    phase="front",
                    page_number=None,
                )
            return

        self._render_section_page_items(
            section.header,
            document,
            context,
            phase="front" if has_front_matter else "main",
            page_number=1,
        )

    def _render_main_page_items(
        self,
        section: object,
        document: Document,
        context: DocxRenderContext,
    ) -> None:
        if not document.settings.overlays:
            return
        self._unlink_page_item_header(section.header)
        self._render_section_page_items(
            section.header,
            document,
            context,
            phase="main",
            page_number=None,
        )

    def _render_section_page_items(
        self,
        header: object,
        document: Document,
        context: DocxRenderContext,
        *,
        phase: str,
        page_number: int | None,
    ) -> None:
        boxes = resolve_positioned_boxes(
            document.settings.overlays,
            document.settings,
            context.unit,
            page_number=page_number,
            phase=phase,
        )
        if not boxes:
            return
        self._unlink_page_item_header(header)
        paragraph = header.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for box in boxes:
            paragraph.add_run()._r.append(
                self._positioned_item_pict(
                    box,
                    context,
                    container=header,
                    absolute=True,
                )
            )

    def _unlink_page_item_header(self, header: object) -> None:
        if hasattr(header, "is_linked_to_previous"):
            header.is_linked_to_previous = False

    def _render_positioned_item(
        self,
        container: object,
        item: PositionedItem,
        context: DocxRenderContext,
    ) -> None:
        if item.placement == "inline":
            self._render_inline_positioned_item(container, item, context)
            return
        paragraph = self._positioned_paragraph(container)
        paragraph.add_run()._r.append(
            self._positioned_item_pict(
                self._single_positioned_box(item, context),
                context,
                container=container,
                absolute=True,
            )
        )

    def _render_inline_positioned_item(
        self,
        container: object,
        item: PositionedItem,
        context: DocxRenderContext,
    ) -> None:
        paragraph = self._add_paragraph(container)
        paragraph.paragraph_format.space_after = Pt(6)
        if isinstance(item, ImageBox):
            paragraph.add_run().add_picture(
                self._image_box_picture_source(item),
                width=Inches(length_to_inches(item.width, item.unit or context.unit)),
                height=Inches(length_to_inches(item.height, item.unit or context.unit)),
            )
            return
        box = PositionedBox(
            item=item,
            x=0.0,
            y=0.0,
            width=length_to_inches(item.width, item.unit or context.unit),
            height=length_to_inches(item.height, item.unit or context.unit),
        )
        paragraph.add_run()._r.append(
            self._positioned_item_pict(
                box,
                context,
                container=container,
                absolute=False,
            )
        )

    def _positioned_paragraph(self, container: object) -> object:
        paragraphs = getattr(container, "paragraphs", None)
        if paragraphs:
            return paragraphs[-1]
        paragraph = self._add_paragraph(container)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        return paragraph

    def _single_positioned_box(
        self,
        item: PositionedItem,
        context: DocxRenderContext,
    ) -> PositionedBox:
        return resolve_positioned_boxes([item], context.settings, context.unit)[0]

    def _positioned_item_pict(
        self,
        box: PositionedBox,
        context: DocxRenderContext,
        *,
        container: object,
        absolute: bool,
    ) -> object:
        item = box.item
        if isinstance(item, Shape):
            return self._shape_pict(item, box, absolute=absolute)
        if isinstance(item, TextBox):
            return self._text_box_pict(
                item,
                box,
                theme=context.theme,
                absolute=absolute,
            )
        return self._image_box_pict(item, box, container=container, absolute=absolute)

    def _shape_pict(self, shape: Shape, box: PositionedBox, *, absolute: bool) -> object:
        x = box.x * 72
        y = box.y * 72
        width = box.width * 72
        height = box.height * 72
        shape_id = f"oodocs_shape_{self._bookmark_id}"
        self._bookmark_id += 1
        stroke = self._vml_stroke_xml(shape)
        fill = self._vml_fill_xml(shape)

        if shape.kind == "line":
            if absolute:
                style = (
                    f"position:absolute;z-index:{shape.z_index};"
                    "mso-position-horizontal-relative:page;"
                    "mso-position-vertical-relative:page;"
                    "mso-wrap-style:none"
                )
                line_from = f'{x:.2f}pt,{y:.2f}pt'
                line_to = f'{x + width:.2f}pt,{y + height:.2f}pt'
            else:
                style = f"width:{abs(width):.2f}pt;height:{abs(height):.2f}pt"
                line_from = "0,0"
                line_to = f'{width:.2f}pt,{height:.2f}pt'
            return parse_xml(
                f'<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                f'xmlns:v="urn:schemas-microsoft-com:vml">'
                f'<v:line id="{shape_id}" '
                f'style="{xml_escape(style)}" '
                f'from="{line_from}" to="{line_to}">'
                f'{stroke}'
                f'</v:line>'
                f'</w:pict>'
            )

        tag = "v:oval" if shape.kind == "ellipse" else "v:rect"
        style = self._vml_box_style(
            x=x,
            y=y,
            width=width,
            height=height,
            z_index=shape.z_index,
            absolute=absolute,
        )
        return parse_xml(
            f'<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'xmlns:v="urn:schemas-microsoft-com:vml">'
            f'<{tag} id="{shape_id}" style="{xml_escape(style)}">'
            f'{fill}'
            f'{stroke}'
            f'</{tag}>'
            f'</w:pict>'
        )

    def _text_box_pict(
        self,
        text_box: TextBox,
        box: PositionedBox,
        *,
        theme: Theme,
        absolute: bool,
    ) -> object:
        x = box.x * 72
        y = box.y * 72
        width = box.width * 72
        height = box.height * 72
        shape_id = f"oodocs_textbox_{self._bookmark_id}"
        self._bookmark_id += 1
        style = self._vml_box_style(
            x=x,
            y=y,
            width=width,
            height=height,
            z_index=text_box.z_index,
            absolute=absolute,
        )
        style += f";v-text-anchor:{self._vml_text_anchor(text_box.vertical_alignment)}"
        font_size = text_box.font_size or theme.typography.body_font_size
        paragraph_xml = self._text_box_paragraph_xml(text_box, font_size)
        return parse_xml(
            f'<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'xmlns:v="urn:schemas-microsoft-com:vml">'
            f'<v:shape id="{shape_id}" type="#_x0000_t202" '
            f'style="{xml_escape(style)}" stroked="f" filled="f">'
            '<v:textbox inset="0,0,0,0">'
            f'<w:txbxContent>{paragraph_xml}</w:txbxContent>'
            '</v:textbox>'
            '</v:shape>'
            '</w:pict>'
        )

    def _image_box_pict(
        self,
        image_box: ImageBox,
        box: PositionedBox,
        *,
        container: object,
        absolute: bool,
    ) -> object:
        x = box.x * 72
        y = box.y * 72
        width = box.width * 72
        height = box.height * 72
        shape_id = f"oodocs_image_{self._bookmark_id}"
        self._bookmark_id += 1
        style = self._vml_box_style(
            x=x,
            y=y,
            width=width,
            height=height,
            z_index=image_box.z_index,
            absolute=absolute,
        )
        relationship_id = self._image_box_relationship_id(container, image_box)
        return parse_xml(
            f'<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'xmlns:v="urn:schemas-microsoft-com:vml" '
            f'xmlns:o="urn:schemas-microsoft-com:office:office" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<v:shape id="{shape_id}" type="#_x0000_t75" '
            f'style="{xml_escape(style)}" stroked="f">'
            f'<v:imagedata r:id="{relationship_id}" o:title=""/>'
            '</v:shape>'
            '</w:pict>'
        )

    def _vml_box_style(
        self,
        *,
        x: float,
        y: float,
        width: float,
        height: float,
        z_index: int,
        absolute: bool,
    ) -> str:
        if not absolute:
            return f"width:{width:.2f}pt;height:{height:.2f}pt;z-index:{z_index}"
        return (
            f"position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;"
            f"width:{width:.2f}pt;height:{height:.2f}pt;z-index:{z_index};"
            "mso-position-horizontal-relative:page;"
            "mso-position-vertical-relative:page;"
            "mso-wrap-style:none"
        )

    def _vml_text_anchor(self, vertical_alignment: str) -> str:
        return {"top": "top", "middle": "middle", "bottom": "bottom"}[
            vertical_alignment
        ]

    def _text_box_paragraph_xml(self, text_box: TextBox, font_size: float) -> str:
        runs = "".join(
            self._text_box_run_xml(fragment, font_size)
            for fragment in text_box.content
        )
        alignment = {"left": "left", "center": "center", "right": "right"}[
            text_box.text_alignment
        ]
        return (
            '<w:p>'
            f'<w:pPr><w:jc w:val="{alignment}"/></w:pPr>'
            f'{runs or self._text_box_run_xml(Text(""), font_size)}'
            '</w:p>'
        )

    def _text_box_run_xml(self, fragment: Text, font_size: float) -> str:
        style = fragment.style
        properties = [
            f'<w:sz w:val="{int(round((style.font_size or font_size) * 2))}"/>',
            f'<w:szCs w:val="{int(round((style.font_size or font_size) * 2))}"/>',
        ]
        if style.bold:
            properties.append("<w:b/>")
        if style.italic:
            properties.append("<w:i/>")
        if style.text_color is not None:
            properties.append(f'<w:color w:val="{style.text_color}"/>')
        text = xml_escape(fragment.value)
        return (
            '<w:r>'
            f'<w:rPr>{"".join(properties)}</w:rPr>'
            f'<w:t xml:space="preserve">{text}</w:t>'
            '</w:r>'
        )

    def _vml_stroke_xml(self, shape: Shape) -> str:
        if shape.stroke.color is None or shape.stroke.width <= 0:
            return '<v:stroke on="false"/>'
        return (
            f'<v:stroke color="#{shape.stroke.color}" '
            f'weight="{shape.stroke.width_points():.2f}pt"/>'
        )

    def _vml_fill_xml(self, shape: Shape) -> str:
        if shape.fill_color is None:
            return '<v:fill on="false"/>'
        return f'<v:fill color="#{shape.fill_color}"/>'

    def _render_table(
        self,
        container: object,
        table_block: Table,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        text_width: float,
        word_document: WordDocument,
    ) -> None:
        table_style = theme.stylesheet.resolve("table", table_block.style, TableStyle())
        split_table = table_block._resolve_split()
        media_placement = table_block._resolve_placement()
        self._apply_media_placement_before(container, word_document, media_placement)

        def render_caption() -> None:
            if table_block.caption is None:
                return
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.resolve_caption_label("table", "caption"),
                    render_index.table_number(table_block),
                    table_block.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            anchor = render_index.table_anchor(table_block)
            if anchor is not None:
                self._add_bookmark(caption, anchor)
            if theme.captions.table_caption_position == "above":
                self._keep_with_next(caption)

        if table_block.caption is not None and theme.captions.table_caption_position == "above":
            render_caption()

        layout = build_table_layout(table_block.header_rows, table_block.rows)
        table = container.add_table(rows=layout.row_count, cols=layout.column_count)
        table.style = "Table Grid"
        table.alignment = TABLE_ALIGNMENTS[theme.blocks.table_block_alignment]
        self._prevent_table_row_split(table)
        if split_table or table_style.repeat_header_rows:
            self._repeat_table_header_rows(table, layout.header_row_count)
        else:
            self._keep_table_together(table)
        column_widths = table_block._column_widths_in_inches(
            unit,
            available_width=text_width,
        )
        if column_widths is not None:
            for column_index, width in enumerate(column_widths):
                table.columns[column_index].width = Inches(width)

        table_cells = [row.cells for row in table.rows]
        if table_block.caption is None and table_cells and table_cells[0]:
            anchor = render_index.table_anchor(table_block)
            if anchor is not None:
                self._add_bookmark(table_cells[0][0].paragraphs[0], anchor)
        for cell_placement in layout.placements:
            start_cell = table_cells[cell_placement.row][cell_placement.column]
            target_cell = start_cell
            if cell_placement.cell.colspan > 1 or cell_placement.cell.rowspan > 1:
                end_cell = table_cells[
                    cell_placement.row + cell_placement.cell.rowspan - 1
                ][cell_placement.column + cell_placement.cell.colspan - 1]
                target_cell = start_cell.merge(end_cell)

            paragraph = target_cell.paragraphs[0]
            effective_style = table_block._effective_cell_style(
                cell_placement,
                stylesheet=theme.stylesheet,
                table_style=table_style,
            )
            cell_text_alignment = self._table_cell_text_alignment(
                cell_placement,
                table_block,
                theme.stylesheet,
                table_style,
            ) or "left"
            paragraph.alignment = ALIGNMENTS[cell_text_alignment]
            cell_vertical_alignment = self._table_cell_vertical_alignment(
                cell_placement,
                table_block,
                theme.stylesheet,
                table_style,
            )
            if cell_vertical_alignment is not None:
                target_cell.vertical_alignment = CELL_VERTICAL_ALIGNMENTS[cell_vertical_alignment]
            self._append_runs(
                paragraph,
                cell_placement.cell.content.content or [Text("")],
                default_style=effective_style.text_style(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            if effective_style.background_color is not None:
                self._set_cell_shading(target_cell, effective_style.background_color)
            self._set_cell_border_edges(
                target_cell,
                table_style._border_edges(
                    row=cell_placement.row,
                    rowspan=cell_placement.cell.rowspan,
                    row_count=layout.row_count,
                    header_row_count=layout.header_row_count,
                ),
            )
            self._set_cell_margins(target_cell, *table_style.cell_padding.to_points())

        if table_block.caption is not None and theme.captions.table_caption_position == "below":
            render_caption()
        self._apply_media_placement_after(container, word_document, media_placement)

    def _table_cell_text_alignment(
        self,
        placement: object,
        table_block: Table,
        stylesheet: object | None = None,
        table_style: TableStyle | None = None,
    ) -> str | None:
        return table_block._effective_cell_style(
            placement,
            stylesheet=stylesheet,
            table_style=table_style,
        ).text_alignment

    def _table_cell_vertical_alignment(
        self,
        placement: object,
        table_block: Table,
        stylesheet: object | None = None,
        table_style: TableStyle | None = None,
    ) -> str | None:
        return table_block._effective_cell_style(
            placement,
            stylesheet=stylesheet,
            table_style=table_style,
        ).vertical_alignment

    def _render_figure(
        self,
        container: object,
        figure: Figure,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
        in_box: bool = False,
    ) -> None:
        placement = figure.resolved_placement()
        self._apply_media_placement_before(container, word_document, placement)

        def render_caption() -> None:
            if figure.caption is None:
                return
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
            caption.paragraph_format.space_after = Pt(0 if in_box else 12)
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.resolve_caption_label("figure", "caption"),
                    render_index.figure_number(figure),
                    figure.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            anchor = render_index.figure_anchor(figure)
            if anchor is not None:
                self._add_bookmark(caption, anchor)
            if theme.captions.figure_caption_position == "above":
                self._keep_with_next(caption)

        if figure.caption is not None and theme.captions.figure_caption_position == "above":
            render_caption()

        paragraph = self._add_paragraph(container)
        paragraph.alignment = ALIGNMENTS[theme.blocks.figure_block_alignment]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0 if in_box else 12)
        if figure.caption is None:
            anchor = render_index.figure_anchor(figure)
            if anchor is not None:
                self._add_bookmark(paragraph, anchor)
        if figure.caption is not None and theme.captions.figure_caption_position == "below":
            self._keep_with_next(paragraph)
        run = paragraph.add_run()
        resolved_width = figure.width_in_inches(unit)
        resolved_height = figure.height_in_inches(unit)
        width = Inches(resolved_width) if resolved_width is not None else None
        height = Inches(resolved_height) if resolved_height is not None else None
        inline_shape = run.add_picture(
            self._figure_picture_source(figure, unit),
            width=width,
            height=height,
        )
        self._set_picture_alt_text(inline_shape, self._figure_alt_text(figure, "Figure"))

        if figure.caption is not None and theme.captions.figure_caption_position == "below":
            render_caption()
        self._apply_media_placement_after(container, word_document, placement)

    def _render_subfigure_group(
        self,
        container: object,
        group: SubFigureGroup,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
        in_box: bool = False,
    ) -> None:
        placement = group.resolved_placement()
        self._apply_media_placement_before(container, word_document, placement)

        def render_caption() -> None:
            if group.caption is None:
                return
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
            caption.paragraph_format.space_after = Pt(0 if in_box else 12)
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.resolve_caption_label("figure", "caption"),
                    render_index.figure_number(group),
                    group.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            anchor = render_index.figure_anchor(group)
            if anchor is not None:
                self._add_bookmark(caption, anchor)
            if theme.captions.figure_caption_position == "above":
                self._keep_with_next(caption)

        if group.caption is not None and theme.captions.figure_caption_position == "above":
            render_caption()

        row_count = (len(group.subfigures) + group.columns - 1) // group.columns
        table = container.add_table(rows=row_count, cols=group.columns)
        table.alignment = TABLE_ALIGNMENTS[theme.blocks.figure_block_alignment]
        gap_points = length_to_inches(group.column_gap, group.unit or unit) * 72

        for index, subfigure in enumerate(group.subfigures):
            row_index = index // group.columns
            column_index = index % group.columns
            cell = table.cell(row_index, column_index)
            self._set_cell_margins(
                cell,
                0,
                gap_points / 2 if column_index < group.columns - 1 else 0,
                0,
                gap_points / 2 if column_index > 0 else 0,
            )
            image_paragraph = cell.paragraphs[0]
            image_paragraph.alignment = ALIGNMENTS[theme.blocks.figure_block_alignment]
            image_paragraph.paragraph_format.space_after = Pt(2)
            run = image_paragraph.add_run()
            resolved_width = subfigure.width_in_inches(unit)
            resolved_height = subfigure.height_in_inches(unit)
            width = Inches(resolved_width) if resolved_width is not None else None
            height = Inches(resolved_height) if resolved_height is not None else None
            inline_shape = run.add_picture(
                self._figure_picture_source(subfigure, unit),
                width=width,
                height=height,
            )
            self._set_picture_alt_text(
                inline_shape,
                self._figure_alt_text(subfigure, "Subfigure"),
            )

            anchor_paragraph = image_paragraph
            if subfigure.caption is not None:
                subcaption = cell.add_paragraph()
                subcaption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
                subcaption.paragraph_format.space_after = Pt(0)
                self._append_runs(
                    subcaption,
                    self._subfigure_caption_fragments(
                        group.formatted_label_for_index(index),
                        subfigure.caption,
                    ),
                    default_size=theme.caption_size(),
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )
                anchor_paragraph = subcaption
            anchor = render_index.figure_anchor(subfigure)
            if anchor is not None:
                self._add_bookmark(anchor_paragraph, anchor)

        if group.caption is not None and theme.captions.figure_caption_position == "below":
            render_caption()
        self._apply_media_placement_after(container, word_document, placement)

    def _render_subtable_group(
        self,
        container: object,
        group: SubTableGroup,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
        text_width: float,
        in_box: bool = False,
    ) -> None:
        placement = group.resolved_placement()
        self._apply_media_placement_before(container, word_document, placement)

        def render_caption() -> None:
            if group.caption is None:
                return
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
            caption.paragraph_format.space_after = Pt(0 if in_box else 12)
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.resolve_caption_label("table", "caption"),
                    render_index.table_number(group),
                    group.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            anchor = render_index.table_anchor(group)
            if anchor is not None:
                self._add_bookmark(caption, anchor)
            if theme.captions.table_caption_position == "above":
                self._keep_with_next(caption)

        if group.caption is not None and theme.captions.table_caption_position == "above":
            render_caption()

        row_count = (len(group.subtables) + group.columns - 1) // group.columns
        table = container.add_table(rows=row_count, cols=group.columns)
        table.alignment = TABLE_ALIGNMENTS[theme.blocks.table_block_alignment]
        self._prevent_table_row_split(table)
        self._keep_table_together(table)
        gap_inches = length_to_inches(group.column_gap, group.unit or unit)
        gap_points = gap_inches * 72
        cell_text_width = max(
            (text_width - gap_inches * max(group.columns - 1, 0)) / group.columns,
            0,
        )

        for index, subtable in enumerate(group.subtables):
            row_index = index // group.columns
            column_index = index % group.columns
            cell = table.cell(row_index, column_index)
            self._set_cell_margins(
                cell,
                0,
                gap_points / 2 if column_index < group.columns - 1 else 0,
                0,
                gap_points / 2 if column_index > 0 else 0,
            )
            self._render_table(
                cell,
                subtable.table_without_caption(),
                theme,
                render_index,
                unit,
                text_width=cell_text_width,
                word_document=word_document,
            )

            anchor_paragraph = cell.paragraphs[0]
            if subtable.caption is not None:
                subcaption = cell.add_paragraph()
                subcaption.alignment = ALIGNMENTS[theme.captions.caption_text_alignment]
                subcaption.paragraph_format.space_after = Pt(0)
                self._append_runs(
                    subcaption,
                    self._subfigure_caption_fragments(
                        group.formatted_label_for_index(index),
                        subtable.caption,
                    ),
                    default_size=theme.caption_size(),
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )
                anchor_paragraph = subcaption
            anchor = render_index.table_anchor(subtable)
            if anchor is not None:
                self._add_bookmark(anchor_paragraph, anchor)

        if group.caption is not None and theme.captions.table_caption_position == "below":
            render_caption()
        self._apply_media_placement_after(container, word_document, placement)

    def _set_paragraph_shading(self, paragraph: object, fill: str) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        paragraph_properties.append(shading)

    def _set_page_background(self, word_document: WordDocument, fill: str) -> None:
        document_element = word_document._element
        existing = document_element.find(qn("w:background"))
        if existing is not None:
            document_element.remove(existing)
        background = OxmlElement("w:background")
        background.set(qn("w:color"), fill)
        document_element.insert(0, background)

    def _set_cell_shading(self, cell: object, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    def _set_cell_width(self, cell: object, width: float) -> None:
        properties = cell._tc.get_or_add_tcPr()
        for existing in list(properties.findall(qn("w:tcW"))):
            properties.remove(existing)
        width_element = OxmlElement("w:tcW")
        width_element.set(qn("w:w"), str(max(int(round(width * 1440)), 0)))
        width_element.set(qn("w:type"), "dxa")
        properties.append(width_element)

    def _set_table_width(self, table: object, width: float) -> None:
        properties = table._tbl.tblPr
        for existing in list(properties.findall(qn("w:tblW"))):
            properties.remove(existing)
        width_element = OxmlElement("w:tblW")
        width_element.set(qn("w:w"), str(max(int(round(width * 1440)), 0)))
        width_element.set(qn("w:type"), "dxa")
        properties.insert(0, width_element)

    def _set_cell_borders(self, cell: object, color: str, width: float) -> None:
        properties = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        size = str(max(int(round(width * 8)), 0))
        for edge_name in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), size)
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), color)
            borders.append(edge)
        properties.append(borders)

    def _set_cell_borders_none(self, cell: object) -> None:
        properties = cell._tc.get_or_add_tcPr()
        for existing in list(properties.findall(qn("w:tcBorders"))):
            properties.remove(existing)
        borders = OxmlElement("w:tcBorders")
        for edge_name in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "nil")
            borders.append(edge)
        properties.append(borders)

    def _set_cell_border_edges(self, cell: object, edges: dict[str, object]) -> None:
        properties = cell._tc.get_or_add_tcPr()
        for existing in list(properties.findall(qn("w:tcBorders"))):
            properties.remove(existing)
        borders = OxmlElement("w:tcBorders")
        for edge_name in ("top", "left", "bottom", "right"):
            border = edges[edge_name]
            edge = OxmlElement(f"w:{edge_name}")
            if border.color is not None and border.width > 0:
                edge.set(qn("w:val"), "single")
                edge.set(qn("w:sz"), str(max(int(round(border.width_points() * 8)), 0)))
                edge.set(qn("w:space"), "0")
                edge.set(qn("w:color"), border.color)
            else:
                edge.set(qn("w:val"), "nil")
            borders.append(edge)
        properties.append(borders)

    def _set_paragraph_bottom_border(self, paragraph: object, color: str, width: float) -> None:
        properties = paragraph._p.get_or_add_pPr()
        for existing in list(properties.findall(qn("w:pBdr"))):
            properties.remove(existing)
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(max(int(round(width * 8)), 1)))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        borders.append(bottom)
        properties.append(borders)

    def _set_cell_padding(self, cell: object, padding: float) -> None:
        self._set_cell_margins(cell, padding, padding, padding, padding)

    def _set_cell_margins(self, cell: object, top: float, right: float, bottom: float, left: float) -> None:
        properties = cell._tc.get_or_add_tcPr()
        for existing in list(properties.findall(qn("w:tcMar"))):
            properties.remove(existing)
        margins = OxmlElement("w:tcMar")
        for side, padding in (
            ("top", top),
            ("left", left),
            ("bottom", bottom),
            ("right", right),
        ):
            element = OxmlElement(f"w:{side}")
            element.set(qn("w:w"), str(max(int(round(padding * 20)), 0)))
            element.set(qn("w:type"), "dxa")
            margins.append(element)
        properties.append(margins)

    def _configure_named_style(
        self,
        word_document: WordDocument,
        style_name: str,
        *,
        font_name: str,
        font_size: float,
        bold: bool,
        italic: bool,
        text_color: str | None = None,
    ) -> None:
        style = word_document.styles[style_name]
        self._set_style_font_family(style, font_name)
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = (
            RGBColor.from_string(text_color) if text_color is not None else RGBColor(0, 0, 0)
        )

    def _set_style_font_family(self, style: object, font_name: str) -> None:
        style.font.name = font_name
        run_properties = style.element.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            style.element.append(run_properties)
        self._set_rfonts(run_properties, font_name)

    def _set_run_font_family(self, run: object, font_name: str) -> None:
        run.font.name = font_name
        run_properties = run._r.get_or_add_rPr()
        self._set_rfonts(run_properties, font_name)

    def _set_rfonts(self, run_properties: object, font_name: str) -> None:
        fonts = run_properties.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run_properties.append(fonts)
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{slot}"), font_name)

    def _resolve_fragment_text(self, fragment: Text, theme: Theme | None, render_index: RenderIndex | None) -> str:
        if isinstance(fragment, _BlockReference):
            if fragment.label is not None:
                return "".join(item.plain_text() for item in fragment.label)
            if theme is None or render_index is None:
                return fragment.plain_text()
            return self._resolve_block_reference(
                fragment.target,
                theme,
                render_index,
                fragment.reference_format,
            )
        if isinstance(fragment, ReferenceGroup):
            if theme is None or render_index is None:
                return fragment.plain_text()
            return "".join(
                piece.text
                for piece in reference_text_pieces(
                    fragment.targets,
                    fragment.reference_format,
                    theme,
                    render_index,
                    range_reference=fragment.range_reference,
                )
            )
        if isinstance(fragment, Citation):
            if render_index is None:
                return fragment.plain_text()
            citation_entry = render_index.citation_entry(fragment.target)
            return format_citation_label(
                citation_entry.source,
                citation_entry.number,
                theme.citations.citation_style if theme is not None else "numeric",
            )
        if isinstance(fragment, Hyperlink):
            return fragment.plain_text()
        if isinstance(fragment, InlineChip):
            return fragment.display_text()
        if isinstance(fragment, Comment):
            return fragment.value
        if isinstance(fragment, Footnote):
            return fragment.value
        if isinstance(fragment, Math):
            return fragment.plain_text()
        return fragment.plain_text()

    def _resolve_block_reference(
        self,
        target: object,
        theme: Theme,
        render_index: RenderIndex,
        reference_format: object | None = None,
    ) -> str:
        from oodocs.components.inline import ReferenceFormat

        return resolve_block_reference(
            target,
            theme,
            render_index,
        ).text(reference_format if isinstance(reference_format, ReferenceFormat) else None)

    def _caption_fragments(self, label: str, number: int | None, caption: Paragraph) -> list[Text]:
        if number is None:
            return caption.content
        return [Text(f"{label} {number}. ")] + caption.content

    def _subfigure_caption_fragments(self, label: str, caption: Paragraph) -> list[Text]:
        return [Text(f"{label} ")] + caption.content

    def _render_caption_list(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        entries: list[object],
        theme: Theme,
        render_index: RenderIndex,
        default_title: str,
        label: str,
        *,
        show_page_numbers: bool,
        leader: str,
        text_width: float,
    ) -> None:
        self._add_heading(word_document, title or [Text(default_title)], level=theme.generated_content.generated_heading_level, theme=theme, number_label=None)
        for entry in entries:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            if show_page_numbers:
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(text_width),
                    WD_TAB_ALIGNMENT.RIGHT,
                    WD_TAB_LEADER.DOTS if leader == "." else WD_TAB_LEADER.SPACES,
                )
            anchor = entry.anchor
            self._append_hyperlink_runs(
                paragraph,
                anchor,
                self._caption_fragments(label, entry.number, entry.block.caption),
                internal=True,
                style=TextStyle(),
                default_size=theme.caption_size(),
            )
            if show_page_numbers:
                paragraph.add_run("\t")
                self._append_pageref_field(paragraph, entry.anchor)

    def _render_countable_list(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        entries: list[object],
        theme: Theme,
        render_index: RenderIndex,
        default_title: str,
        *,
        show_page_numbers: bool,
        leader: str,
        text_width: float,
    ) -> None:
        self._add_heading(word_document, title or [Text(default_title)], level=theme.generated_content.generated_heading_level, theme=theme, number_label=None)
        for entry in entries:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            if show_page_numbers:
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(text_width),
                    WD_TAB_ALIGNMENT.RIGHT,
                    WD_TAB_LEADER.DOTS if leader == "." else WD_TAB_LEADER.SPACES,
                )
            label = entry.block.reference_text(entry.number)
            if entry.block.title is not None:
                fragments = [Text(f"{label}. ")] + entry.block.title
            else:
                fragments = [Text(label)]
            self._append_hyperlink_runs(
                paragraph,
                entry.anchor,
                fragments,
                internal=True,
                style=TextStyle(),
                default_size=theme.caption_size(),
            )
            if show_page_numbers:
                paragraph.add_run("\t")
                self._append_pageref_field(paragraph, entry.anchor)

    def _render_comment_list(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        theme: Theme,
        render_index: RenderIndex,
    ) -> None:
        word_document.add_page_break()
        self._add_heading(word_document, title or [Text(theme.resolve_generated_page_title("list_of_comments"))], level=theme.generated_content.generated_heading_level, theme=theme, number_label=None)
        for entry in render_index.comments:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            self._append_runs(
                paragraph,
                [Text(f"[{entry.number}] ")] + entry.comment.comment,
                default_size=theme.typography.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )

    def _figure_picture_source(self, figure: Figure | SubFigure, unit: str) -> str | BytesIO:
        source = figure.image_source
        if figure.crop is not None or figure.rotation:
            return processed_image_source_to_buffer(
                source,
                image_format=figure.image_format,
                image_dpi=figure.image_dpi,
                crop=figure.crop,
                rotation=figure.rotation,
                default_unit=figure.unit or unit,
                usage="DOCX rendering",
            )
        if isinstance(source, Path):
            return str(source)
        return image_source_to_buffer(
            source,
            image_format=figure.image_format,
            image_dpi=figure.image_dpi,
            usage="DOCX rendering",
        )

    def _figure_alt_text(self, figure: Figure | SubFigure, fallback: str) -> str:
        if figure.alt_text is not None:
            return figure.alt_text
        if figure.caption is not None:
            return figure.caption.plain_text()
        return fallback

    def _set_picture_alt_text(self, inline_shape: object, alt_text: str) -> None:
        doc_properties = getattr(getattr(inline_shape, "_inline", None), "docPr", None)
        if doc_properties is not None:
            doc_properties.set("descr", alt_text)

    def _image_box_picture_source(self, image_box: ImageBox) -> str | BytesIO:
        source = image_box.image_source
        if isinstance(source, Path):
            return str(source)
        return image_source_to_buffer(
            source,
            image_format=image_box.image_format,
            image_dpi=image_box.image_dpi,
            usage="DOCX positioned image rendering",
        )

    def _image_box_relationship_id(self, container: object, image_box: ImageBox) -> str:
        part = getattr(container, "part", None)
        if part is None:
            part = getattr(container, "_parent", None)
            part = getattr(part, "part", None)
        if part is None:
            raise TypeError("Cannot attach positioned image to this DOCX container")
        relationship_id, _ = part.get_or_add_image(self._image_box_picture_source(image_box))
        return relationship_id

    def _render_footnote_list(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        theme: Theme,
        render_index: RenderIndex,
    ) -> None:
        word_document.add_page_break()
        self._add_heading(word_document, title or [Text(theme.resolve_generated_page_title("list_of_footnotes"))], level=theme.generated_content.generated_heading_level, theme=theme, number_label=None)
        for entry in render_index.footnotes:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            marker = theme.footnotes.format_marker(entry.stream, entry.number)
            self._append_runs(
                paragraph,
                [Text(f"[{marker}] ")] + entry.footnote.note,
                default_size=theme.typography.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )

    def _render_reference_list(
        self,
        word_document: WordDocument,
        block: ListOfReferences,
        theme: Theme,
        render_index: RenderIndex,
    ) -> None:
        word_document.add_page_break()
        self._add_heading(word_document, block.title or [Text(theme.resolve_generated_page_title("list_of_references"))], level=theme.generated_content.generated_heading_level, theme=theme, number_label=None)
        for entry in render_index.reference_entries(block, reference_sort=theme.citations.reference_sort):
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            marker = reference_entry_marker(
                entry.number,
                citation_style=theme.citations.citation_style,
                reference_style=theme.citations.reference_style,
            )
            fragments = entry.source.reference_fragments(theme.citations.reference_style)
            if marker:
                fragments = [Text(f"{marker} ")] + fragments
            self._append_runs(
                paragraph,
                fragments,
                default_size=theme.typography.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            self._add_bookmark(paragraph, entry.anchor)

    def _render_table_of_contents(
        self,
        word_document: WordDocument,
        block: TableOfContents,
        context: DocxRenderContext,
    ) -> None:
        theme = context.theme
        render_index = context.render_index
        self._add_generated_page_title(
            word_document,
            block.title or [Text(theme.resolve_generated_page_title("table_of_contents"))],
            level=theme.generated_content.generated_heading_level,
            theme=theme,
        )
        entries = render_index.scoped_headings(block)
        if block.show_page_numbers and block.scope == "document":
            self._append_native_toc_field(word_document, block)
            return

        for entry in entries:
            toc_style = self._toc_level_style(block, entry.level)
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(toc_style.indent)
            paragraph.paragraph_format.space_before = Pt(toc_style.space_before)
            paragraph.paragraph_format.space_after = Pt(toc_style.space_after)
            if block.show_page_numbers:
                text_width = context.settings.text_width_in_inches()
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(text_width),
                    WD_TAB_ALIGNMENT.RIGHT,
                    WD_TAB_LEADER.DOTS if block.leader == "." else WD_TAB_LEADER.SPACES,
                )
            self._append_hyperlink_runs(
                paragraph,
                entry.anchor,
                self._heading_fragments(entry.title, entry.number),
                internal=True,
                style=TextStyle(
                    bold=toc_style.bold,
                    italic=toc_style.italic,
                ),
                default_size=theme.typography.body_font_size + toc_style.font_size_delta,
            )
            if block.show_page_numbers and entry.anchor is not None:
                paragraph.add_run("\t")
                self._append_pageref_field(paragraph, entry.anchor)

    def _add_generated_page_title(
        self,
        container: object,
        title: list[Text],
        *,
        level: int,
        theme: Theme,
    ) -> None:
        heading_style = theme.resolve_heading_style(level) if level > 0 else None
        heading_text_style = heading_style.text_style if heading_style is not None else None
        space_after = (
            10
            if level == 0
            else (
                heading_style.space_after
                if heading_style.space_after is not None
                else (10 if level == 1 else 6)
            )
        )
        font_size = (
            theme.typography.title_font_size
            if level == 0
            else (heading_text_style.font_size or theme.resolve_heading_size(level))
        )
        self._add_title_line(
            container,
            title,
            font_size=font_size,
            alignment=(
                theme.title_matter.title_text_alignment
                if level == 0
                else (heading_style.text_alignment or "left")
            ),
            bold=bool(heading_text_style.bold) if heading_text_style is not None else True,
            italic=bool(heading_text_style.italic) if heading_text_style is not None else False,
            space_after=space_after,
        )

    def _append_native_toc_field(
        self,
        word_document: WordDocument,
        block: TableOfContents,
    ) -> None:
        max_level = max(block.max_level or 9, 1)
        instruction = f'TOC \\f \\l "1-{max_level}" \\h \\z'
        paragraph = word_document.add_paragraph()
        self._append_field(paragraph, instruction)

    def _append_toc_entry_field(
        self,
        paragraph: object,
        text: str,
        *,
        level: int,
    ) -> None:
        toc_level = min(max(level, 1), 9)
        safe_text = text.replace('"', "'")
        self._append_field(paragraph, f'TC "{safe_text}" \\l {toc_level}')

    def _toc_level_style(self, block: TableOfContents, level: int) -> TableOfContentsLevelStyle:
        if level == 0:
            defaults = TableOfContentsLevelStyle(
                indent=0,
                space_before=16,
                space_after=8,
                font_size_delta=1.2,
                bold=True,
                italic=False,
            )
            override = block.style_for_level(level)
            return TableOfContentsLevelStyle(
                indent=defaults.indent if override.indent is None else override.indent,
                space_before=defaults.space_before if override.space_before is None else override.space_before,
                space_after=defaults.space_after if override.space_after is None else override.space_after,
                font_size_delta=defaults.font_size_delta if override.font_size_delta is None else override.font_size_delta,
                bold=defaults.bold if override.bold is None else override.bold,
                italic=defaults.italic if override.italic is None else override.italic,
            )
        defaults = TableOfContentsLevelStyle(
            indent=0.24 * max(level - 1, 0),
            space_before=12 if level == 1 else (3 if level == 2 else 0),
            space_after=7 if level == 1 else (3 if level == 2 else 2),
            font_size_delta=0.6 if level == 1 else 0,
            bold=True if level == 1 else False,
            italic=False,
        )
        override = block.style_for_level(level)
        return TableOfContentsLevelStyle(
            indent=defaults.indent if override.indent is None else override.indent,
            space_before=defaults.space_before if override.space_before is None else override.space_before,
            space_after=defaults.space_after if override.space_after is None else override.space_after,
            font_size_delta=defaults.font_size_delta if override.font_size_delta is None else override.font_size_delta,
            bold=defaults.bold if override.bold is None else override.bold,
            italic=defaults.italic if override.italic is None else override.italic,
        )

    def _configure_header_footer_sections(
        self,
        word_document: WordDocument,
        document: Document,
        theme: Theme,
        *,
        has_front_matter: bool,
        has_main_matter: bool,
        has_back_matter: bool,
        matter_sections: dict[str, int],
    ) -> None:
        sections = list(word_document.sections)
        if not sections:
            return

        self._set_even_and_odd_headers(
            word_document,
            theme.effective_header_footer().different_odd_even_pages,
        )
        first_section = sections[0]
        if has_front_matter:
            front_section = sections[matter_sections.get("front", 0)]
            self._set_section_page_counter_format(
                front_section,
                theme.page_numbers.front_matter_counter.counter_format,
                start=1,
            )
            self._add_section_header_footer(
                front_section,
                document,
                theme,
                front_matter=True,
                hide_first=document.settings.title_matter.cover is not None,
            )
            if has_main_matter and "main" in matter_sections:
                main_section = sections[matter_sections["main"]]
                self._unlink_section_header_footer(main_section)
                main_section.footer.is_linked_to_previous = False
                self._set_section_page_counter_format(
                    main_section,
                    theme.page_numbers.main_matter_counter.counter_format,
                    start=1 if theme.page_numbers.restart_main_matter else None,
                )
                self._add_section_header_footer(
                    main_section,
                    document,
                    theme,
                    front_matter=False,
                )
        else:
            main_section = sections[matter_sections.get("main", 0)]
            self._set_section_page_counter_format(
                main_section,
                theme.page_numbers.main_matter_counter.counter_format,
                start=1,
            )
            self._add_section_header_footer(
                main_section,
                document,
                theme,
                front_matter=False,
            )

        if has_back_matter and "back" in matter_sections:
            back_section = sections[matter_sections["back"]]
            self._unlink_section_header_footer(back_section)
            counter = (
                theme.page_numbers.back_matter_counter
                or theme.page_numbers.main_matter_counter
            )
            self._set_section_page_counter_format(
                back_section,
                counter.counter_format,
                start=1 if theme.page_numbers.restart_back_matter else None,
            )
            self._add_section_header_footer(
                back_section,
                document,
                theme,
                front_matter=False,
            )

    def _add_section_header_footer(
        self,
        section: object,
        document: Document,
        theme: Theme,
        *,
        front_matter: bool,
        hide_first: bool = False,
    ) -> None:
        header_footer = theme.effective_header_footer()
        section.different_first_page_header_footer = (
            header_footer.different_first_page or section.different_first_page_header_footer
        )
        page_kinds: list[tuple[str, object, object]] = [
            ("default", section.header, section.footer)
        ]
        if section.different_first_page_header_footer:
            page_kinds.append(("first", section.first_page_header, section.first_page_footer))
        if header_footer.different_odd_even_pages:
            page_kinds.append(("even", section.even_page_header, section.even_page_footer))
        text_width = max(
            section.page_width.inches
            - section.left_margin.inches
            - section.right_margin.inches,
            0.1,
        )
        for page_kind, header, footer in page_kinds:
            header.is_linked_to_previous = False
            footer.is_linked_to_previous = False
            if page_kind == "first" and hide_first:
                continue
            self._add_header_footer_part(
                header,
                document,
                theme,
                region="header",
                page_kind=page_kind,
                front_matter=front_matter,
                text_width=text_width,
            )
            self._add_header_footer_part(
                footer,
                document,
                theme,
                region="footer",
                page_kind=page_kind,
                front_matter=front_matter,
                text_width=text_width,
            )

    def _add_header_footer_part(
        self,
        part: object,
        document: Document,
        theme: Theme,
        *,
        region: str,
        page_kind: str,
        front_matter: bool,
        text_width: float,
    ) -> None:
        slots = [
            theme.resolve_header_footer_template(
                region,  # type: ignore[arg-type]
                "left",
                page_kind=page_kind,  # type: ignore[arg-type]
            ),
            theme.resolve_header_footer_template(
                region,  # type: ignore[arg-type]
                "center",
                page_kind=page_kind,  # type: ignore[arg-type]
            ),
            theme.resolve_header_footer_template(
                region,  # type: ignore[arg-type]
                "right",
                page_kind=page_kind,  # type: ignore[arg-type]
            ),
        ]
        if not any(slot for slot in slots):
            return
        paragraph = part.paragraphs[0] if part.paragraphs else part.add_paragraph()
        paragraph.style = "Footer"
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        non_empty_slots = [
            (position, template)
            for position, template in zip(("left", "center", "right"), slots)
            if template
        ]
        if len(non_empty_slots) == 1:
            position, template = non_empty_slots[0]
            paragraph.alignment = ALIGNMENTS[position]
            self._append_header_footer_template(
                paragraph,
                template,
                document,
                theme,
                front_matter=front_matter,
            )
            return
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._configure_header_footer_tabs(paragraph, text_width)
        if slots[0]:
            self._append_header_footer_template(
                paragraph,
                slots[0],
                document,
                theme,
                front_matter=front_matter,
            )
        if slots[1] or slots[2]:
            paragraph.add_run("\t")
        if slots[1]:
            self._append_header_footer_template(
                paragraph,
                slots[1],
                document,
                theme,
                front_matter=front_matter,
            )
        if slots[2]:
            paragraph.add_run("\t")
            self._append_header_footer_template(
                paragraph,
                slots[2],
                document,
                theme,
                front_matter=front_matter,
            )

    def _configure_header_footer_tabs(self, paragraph: object, text_width: float) -> None:
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.clear_all()
        tab_stops.add_tab_stop(Inches(text_width / 2), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Inches(text_width), WD_TAB_ALIGNMENT.RIGHT)

    def _append_header_footer_template(
        self,
        paragraph: object,
        template: str,
        document: Document,
        theme: Theme,
        *,
        front_matter: bool,
    ) -> None:
        font_size = theme.resolve_header_footer_font_size()
        token_pattern = re.compile(r"(\{page\}|\{title\}|\{chapter\}|\{section\})")
        for piece in token_pattern.split(template):
            if not piece:
                continue
            if piece == "{page}":
                self._append_page_number_field(paragraph)
            elif piece == "{title}":
                self._append_header_footer_text(paragraph, document.title, theme)
            elif piece == "{chapter}":
                self._append_field(paragraph, 'STYLEREF "Heading 1"')
            elif piece == "{section}":
                self._append_field(paragraph, 'STYLEREF "Heading 2"')
            else:
                self._append_header_footer_text(paragraph, piece, theme)
        for run in paragraph.runs:
            if run.text or run._r.findall(qn("w:fldChar")) or run._r.findall(qn("w:instrText")):
                self._apply_run_style(
                    run,
                    Text("").style,
                    default_size=font_size,
                )

    def _append_header_footer_text(
        self,
        paragraph: object,
        text: str,
        theme: Theme,
    ) -> None:
        run = paragraph.add_run(text)
        self._apply_run_style(
            run,
            Text(text).style,
            default_size=theme.resolve_header_footer_font_size(),
        )

    def _unlink_section_header_footer(self, section: object) -> None:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            part.is_linked_to_previous = False

    def _set_even_and_odd_headers(
        self,
        word_document: WordDocument,
        enabled: bool,
    ) -> None:
        settings = word_document.settings.element
        existing = settings.find(qn("w:evenAndOddHeaders"))
        if enabled and existing is None:
            settings.append(OxmlElement("w:evenAndOddHeaders"))
        elif not enabled and existing is not None:
            settings.remove(existing)

    def _set_section_page_counter_format(
        self,
        section: object,
        page_counter_format: str,
        *,
        start: int | None = 1,
    ) -> None:
        format_map = {
            "decimal": "decimal",
            "lower-roman": "lowerRoman",
            "upper-roman": "upperRoman",
            "lower-alpha": "lowerLetter",
            "upper-alpha": "upperLetter",
        }
        sect_pr = section._sectPr
        page_number_type = sect_pr.find(qn("w:pgNumType"))
        if page_number_type is None:
            page_number_type = OxmlElement("w:pgNumType")
            sect_pr.append(page_number_type)
        page_number_type.set(
            qn("w:fmt"),
            format_map.get(page_counter_format, "decimal"),
        )
        if start is None:
            page_number_type.attrib.pop(qn("w:start"), None)
        else:
            page_number_type.set(qn("w:start"), str(start))

    def _append_page_number_field(self, paragraph: object) -> None:
        self._append_field(paragraph, "PAGE", cached_result="1")

    def _append_pageref_field(self, paragraph: object, anchor: str) -> None:
        self._append_field(paragraph, f"PAGEREF {anchor} \\h")

    def _append_field(
        self,
        paragraph: object,
        instruction: str,
        *,
        cached_result: str = "",
    ) -> None:
        begin_run = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        begin_run.append(begin)
        paragraph._p.append(begin_run)

        instruction_run = OxmlElement("w:r")
        instruction_text = OxmlElement("w:instrText")
        instruction_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instruction_text.text = f" {instruction} "
        instruction_run.append(instruction_text)
        paragraph._p.append(instruction_run)

        separator_run = OxmlElement("w:r")
        separator = OxmlElement("w:fldChar")
        separator.set(qn("w:fldCharType"), "separate")
        separator_run.append(separator)
        paragraph._p.append(separator_run)

        if cached_result:
            result_run = OxmlElement("w:r")
            result_text = OxmlElement("w:t")
            result_text.text = cached_result
            result_run.append(result_text)
            paragraph._p.append(result_run)

        end_run = OxmlElement("w:r")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run.append(end)
        paragraph._p.append(end_run)

    def _comment_marker(self, fragment: Comment, render_index: RenderIndex | None) -> str:
        if render_index is None:
            return "[?]"
        return f"[{render_index.comment_number(fragment)}]"

    def _footnote_marker(
        self,
        fragment: Footnote,
        render_index: RenderIndex | None,
        theme: Theme | None,
    ) -> str:
        if render_index is None:
            return "?"
        number = render_index.footnote_number(fragment)
        if theme is None:
            return str(number)
        return theme.footnotes.format_marker(fragment.stream, number)

    def _flatten_fragments(self, fragments: list[Text], theme: Theme | None, render_index: RenderIndex | None) -> str:
        parts: list[str] = []
        for fragment in fragments:
            if isinstance(fragment, Comment):
                parts.append(fragment.value)
                parts.append(self._comment_marker(fragment, render_index))
                continue
            if isinstance(fragment, Footnote):
                parts.append(fragment.value)
                parts.append(self._footnote_marker(fragment, render_index, theme))
                continue
            parts.append(self._resolve_fragment_text(fragment, theme, render_index))
        return "".join(parts)
